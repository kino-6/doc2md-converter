"""
Integration tests for metadata control feature.

Tests verify that metadata inclusion/exclusion works correctly through
the entire conversion pipeline including CLI, config, and serializer.

Validates: Requirements 8.3
"""

import pytest
import tempfile
from pathlib import Path
from docx import Document
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.conversion_orchestrator import ConversionOrchestrator


class TestMetadataControlIntegration:
    """Integration tests for metadata control through the conversion pipeline."""
    
    def test_metadata_excluded_in_conversion(self, tmp_path):
        """Test that metadata is excluded when include_metadata=False."""
        # Create a sample Word document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test content.")
        doc.save(str(docx_path))
        
        # Create output path
        output_path = tmp_path / "output.md"
        
        # Create config with include_metadata=False
        config = ConversionConfig(
            input_path=str(docx_path),
            output_path=str(output_path),
            include_metadata=False
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.INFO)
        
        # Create orchestrator and convert
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(docx_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify metadata is NOT in output
        assert "---" not in result.markdown_content
        assert "title:" not in result.markdown_content
        assert "author:" not in result.markdown_content
        
        # Verify content IS in output
        assert "# Test Heading" in result.markdown_content
        assert "Test content." in result.markdown_content
    
    def test_metadata_included_in_conversion(self, tmp_path):
        """Test that metadata is included when include_metadata=True."""
        # Create a sample Word document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test content.")
        doc.save(str(docx_path))
        
        # Create output path
        output_path = tmp_path / "output.md"
        
        # Create config with include_metadata=True
        config = ConversionConfig(
            input_path=str(docx_path),
            output_path=str(output_path),
            include_metadata=True
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.INFO)
        
        # Create orchestrator and convert
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(docx_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify metadata IS in output
        assert "---" in result.markdown_content
        assert "title: Test Document" in result.markdown_content
        assert "author: Test Author" in result.markdown_content
        assert "source_format: docx" in result.markdown_content
        
        # Verify content IS in output
        assert "# Test Heading" in result.markdown_content
        assert "Test content." in result.markdown_content
    
    def test_metadata_control_with_config_file(self, tmp_path):
        """Test metadata control through configuration file."""
        # Create a sample Word document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test content.")
        doc.save(str(docx_path))
        
        # Create config file with include_metadata=True
        config_path = tmp_path / "config.yaml"
        with open(config_path, 'w') as f:
            f.write("include_metadata: true\n")
        
        # Create output path
        output_path = tmp_path / "output.md"
        
        # Load config from file
        from src.config import ConfigManager
        config_manager = ConfigManager()
        file_config = config_manager.load_config(str(config_path))
        
        # Create CLI config
        cli_config = ConversionConfig(
            input_path=str(docx_path),
            output_path=str(output_path)
        )
        
        # Merge configs
        merged_config = config_manager.merge_configs(file_config, cli_config)
        
        # Verify include_metadata is True from config file
        assert merged_config.include_metadata is True
        
        # Create logger
        logger = Logger(log_level=LogLevel.INFO)
        
        # Create orchestrator and convert
        orchestrator = ConversionOrchestrator(config=merged_config, logger=logger)
        result = orchestrator.convert(str(docx_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify metadata IS in output (from config file)
        assert "---" in result.markdown_content
        assert "title: Test Document" in result.markdown_content
    
    def test_metadata_control_cli_overrides_config(self, tmp_path):
        """Test that CLI argument overrides config file for metadata control."""
        # Create a sample Word document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test content.")
        doc.save(str(docx_path))
        
        # Create config file with include_metadata=False
        config_path = tmp_path / "config.yaml"
        with open(config_path, 'w') as f:
            f.write("include_metadata: false\n")
        
        # Create output path
        output_path = tmp_path / "output.md"
        
        # Load config from file
        from src.config import ConfigManager
        config_manager = ConfigManager()
        file_config = config_manager.load_config(str(config_path))
        
        # Create CLI config with include_metadata=True (overrides file)
        cli_config = ConversionConfig(
            input_path=str(docx_path),
            output_path=str(output_path),
            include_metadata=True  # CLI override
        )
        
        # Merge configs (CLI should take precedence)
        merged_config = config_manager.merge_configs(file_config, cli_config)
        
        # Verify include_metadata is True from CLI
        assert merged_config.include_metadata is True
        
        # Create logger
        logger = Logger(log_level=LogLevel.INFO)
        
        # Create orchestrator and convert
        orchestrator = ConversionOrchestrator(config=merged_config, logger=logger)
        result = orchestrator.convert(str(docx_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify metadata IS in output (CLI override worked)
        assert "---" in result.markdown_content
        assert "title: Test Document" in result.markdown_content
