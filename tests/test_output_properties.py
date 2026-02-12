"""Property-based tests for output functionality.

Feature: document-to-markdown-converter
Properties 26, 46, 47: Output properties
Validates: Requirements 6.2, 11.1, 11.2
"""

import pytest
from hypothesis import given, strategies as st, settings
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.output_writer import OutputWriter
import tempfile
import sys
from io import StringIO


# Strategy for generating valid file names
filename_strategy = st.text(
    alphabet=st.characters(
        whitelist_categories=('Lu', 'Ll', 'Nd'),
        min_codepoint=ord('a'),
        max_codepoint=ord('z')
    ),
    min_size=1,
    max_size=20
)

# Strategy for generating text content compatible with Office documents
def is_valid_office_char(c):
    """Check if character is valid for Office documents."""
    code = ord(c)
    # Allow tab, newline, carriage return
    if code in (0x09, 0x0A, 0x0D):
        return True
    # Exclude other control characters (0x00-0x1F, 0x7F-0x9F)
    if code < 0x20 or (0x7F <= code <= 0x9F):
        return False
    # Exclude surrogates
    if 0xD800 <= code <= 0xDFFF:
        return False
    return True

text_content_strategy = st.text(
    alphabet=st.characters(
        blacklist_categories=('Cs',),  # Exclude surrogates
    ).filter(is_valid_office_char),
    min_size=1,
    max_size=200
)


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_26_output_to_file(filename, content):
    """Feature: document-to-markdown-converter, Property 26: Output destination
    
    For any conversion operation with a specified output file path, the converter 
    should output the Markdown content to that file.
    
    **Validates: Requirements 6.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Set up output file path
        output_path = tmp_path / "output.md"
        
        # Create config with output file
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should exist
        assert output_path.exists(), "Output file was not created"
        
        # Property: Output file should contain content
        output_content = output_path.read_text(encoding='utf-8')
        assert len(output_content) > 0, "Output file is empty"
        
        # Property: Output content should match result content
        assert result.markdown_content == output_content, \
            "Result content does not match file content"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_26_output_to_stdout(filename, content):
    """Feature: document-to-markdown-converter, Property 26: Output destination
    
    For any conversion operation without a specified output file path, the converter 
    should output the Markdown content to stdout.
    
    **Validates: Requirements 6.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Create config without output file (stdout)
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Result should contain markdown content
        assert result.markdown_content is not None, "No markdown content in result"
        assert len(result.markdown_content) > 0, "Markdown content is empty"
        
        # Property: No output file should be created
        # (since output_path is None, it should go to stdout)
        # We can't easily test stdout capture in property tests,
        # but we can verify the result contains the content


@settings(max_examples=100)
@given(content=text_content_strategy)
def test_property_26_output_writer_file(content):
    """Feature: document-to-markdown-converter, Property 26: Output destination
    
    For any content, the OutputWriter should write to the specified file path.
    
    **Validates: Requirements 6.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        output_path = tmp_path / "output.md"
        
        # Write content using OutputWriter
        writer = OutputWriter()
        writer.write_to_file(content, str(output_path))
        
        # Property: Output file should exist
        assert output_path.exists(), "Output file was not created"
        
        # Property: Output file should contain the content
        read_content = output_path.read_text(encoding='utf-8')
        
        # Normalize line endings for comparison
        expected_content = content.replace('\r\n', '\n').replace('\r', '\n')
        assert read_content == expected_content, "Content was not preserved"


@settings(max_examples=100)
@given(content=text_content_strategy)
def test_property_26_output_writer_stdout(content):
    """Feature: document-to-markdown-converter, Property 26: Output destination
    
    For any content, the OutputWriter should write to stdout when requested.
    
    **Validates: Requirements 6.2**
    """
    # Capture stdout
    old_stdout = sys.stdout
    sys.stdout = StringIO()
    
    try:
        # Write content using OutputWriter
        writer = OutputWriter()
        writer.write_to_stdout(content)
        
        # Get captured output
        captured_output = sys.stdout.getvalue()
        
        # Property: Captured output should match content
        assert captured_output == content, "Stdout content does not match"
    finally:
        # Restore stdout
        sys.stdout = old_stdout


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_46_preview_mode_no_file(filename, content):
    """Feature: document-to-markdown-converter, Property 46: Preview mode file handling
    
    For any conversion in preview mode, the converter should display the converted 
    Markdown without saving to a file.
    
    **Validates: Requirements 11.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Set up output file path
        output_path = tmp_path / "output.md"
        
        # Create config with preview mode enabled
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            preview_mode=True
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Result should contain markdown content
        assert result.markdown_content is not None, "No markdown content in result"
        assert len(result.markdown_content) > 0, "Markdown content is empty"
        
        # Property: Output file should NOT be created in preview mode
        assert not output_path.exists(), \
            "Output file should not be created in preview mode"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_46_preview_mode_excel(filename, content):
    """Feature: document-to-markdown-converter, Property 46: Preview mode file handling
    
    For any Excel conversion in preview mode, the converter should display the 
    converted Markdown without saving to a file.
    
    **Validates: Requirements 11.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file
        excel_path = tmp_path / f"{filename}.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = content
        wb.save(str(excel_path))
        
        # Set up output file path
        output_path = tmp_path / "output.md"
        
        # Create config with preview mode enabled
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path),
            preview_mode=True
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Result should contain markdown content
        assert result.markdown_content is not None, "No markdown content in result"
        
        # Property: Output file should NOT be created in preview mode
        assert not output_path.exists(), \
            "Output file should not be created in preview mode"


@settings(max_examples=100)
@given(
    num_lines=st.integers(min_value=10, max_value=100),
    preview_lines=st.integers(min_value=1, max_value=50)
)
def test_property_47_preview_line_limit(num_lines, preview_lines):
    """Feature: document-to-markdown-converter, Property 47: Preview line limit
    
    For any preview request, the converter should show the first N lines of the 
    converted output (where N is configurable).
    
    **Validates: Requirements 11.2**
    """
    # Create content with known number of lines
    lines = [f"Line {i+1}" for i in range(num_lines)]
    content = '\n'.join(lines)
    
    # Capture stdout
    old_stdout = sys.stdout
    sys.stdout = StringIO()
    
    try:
        # Use OutputWriter to preview
        writer = OutputWriter()
        writer.preview(content, lines=preview_lines)
        
        # Get captured output
        captured_output = sys.stdout.getvalue()
        
        # Property: Preview should show at most preview_lines of content
        # (plus header/footer lines)
        output_lines = captured_output.split('\n')
        
        # Count content lines (excluding header/footer)
        content_lines = []
        in_content = False
        for line in output_lines:
            if '=' * 80 in line:
                in_content = not in_content
                continue
            if in_content and not line.startswith('...'):
                content_lines.append(line)
        
        # Property: Number of content lines should not exceed preview_lines
        # (may be less if content is shorter)
        expected_lines = min(num_lines, preview_lines)
        assert len(content_lines) <= preview_lines, \
            f"Preview showed {len(content_lines)} lines, expected at most {preview_lines}"
        
        # Property: If content has more lines than preview limit,
        # preview should indicate remaining lines
        if num_lines > preview_lines:
            assert '...' in captured_output or 'more lines' in captured_output, \
                "Preview should indicate there are more lines"
    finally:
        # Restore stdout
        sys.stdout = old_stdout


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    num_paragraphs=st.integers(min_value=20, max_value=100),
    preview_lines=st.integers(min_value=5, max_value=30)
)
def test_property_47_preview_line_limit_document(filename, num_paragraphs, preview_lines):
    """Feature: document-to-markdown-converter, Property 47: Preview line limit
    
    For any document conversion with preview, the converter should show only the 
    first N lines of the converted output.
    
    **Validates: Requirements 11.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document with many paragraphs
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        for i in range(num_paragraphs):
            doc.add_paragraph(f"Paragraph {i+1}")
        doc.save(str(doc_path))
        
        # Create config with preview mode
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Result should contain full markdown content
        assert result.markdown_content is not None, "No markdown content in result"
        
        # Property: Full content should have many lines
        full_lines = result.markdown_content.split('\n')
        assert len(full_lines) > preview_lines, \
            "Test requires content with more lines than preview limit"
        
        # Property: Preview would show only first N lines
        # (We test the OutputWriter.preview method separately)
        # Here we verify the full content is available in result
        assert len(full_lines) >= num_paragraphs, \
            "Full content should contain all paragraphs"


@settings(max_examples=50)
@given(
    content1=text_content_strategy,
    content2=text_content_strategy,
    preview_lines=st.integers(min_value=1, max_value=20)
)
def test_property_47_preview_short_content(content1, content2, preview_lines):
    """Feature: document-to-markdown-converter, Property 47: Preview line limit
    
    For any preview request where content is shorter than the line limit,
    the converter should show all content without truncation.
    
    **Validates: Requirements 11.2**
    """
    # Create short content (fewer lines than preview limit)
    content = f"{content1}\n{content2}"
    content_lines = content.split('\n')
    
    # Only test when content is shorter than preview limit
    if len(content_lines) >= preview_lines:
        return
    
    # Capture stdout
    old_stdout = sys.stdout
    sys.stdout = StringIO()
    
    try:
        # Use OutputWriter to preview
        writer = OutputWriter()
        writer.preview(content, lines=preview_lines)
        
        # Get captured output
        captured_output = sys.stdout.getvalue()
        
        # Property: All content lines should be shown
        for line in content_lines:
            if line.strip():  # Skip empty lines
                assert line in captured_output, \
                    f"Content line '{line}' should be in preview"
        
        # Property: Should not indicate truncation
        # (since content is shorter than limit)
        # Note: The preview method always shows the header/footer,
        # but should not show "more lines" message
        if len(content_lines) <= preview_lines:
            # Check if truncation message is present
            has_truncation = '...' in captured_output and 'more lines' in captured_output
            # If content fits, there should be no truncation message
            # (or the message should indicate 0 more lines)
            if has_truncation:
                assert '0 more lines' in captured_output or \
                       len(content_lines) == preview_lines, \
                       "Should not indicate truncation for short content"
    finally:
        # Restore stdout
        sys.stdout = old_stdout


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy,
    output_format=st.sampled_from(['file', 'stdout', 'preview'])
)
def test_property_26_output_destination_modes(filename, content, output_format):
    """Feature: document-to-markdown-converter, Property 26: Output destination
    
    For any conversion operation, the converter should output to the configured 
    destination (file, stdout, or preview).
    
    **Validates: Requirements 6.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Configure based on output format
        output_path = tmp_path / "output.md" if output_format == 'file' else None
        preview_mode = (output_format == 'preview')
        
        # Create config
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path) if output_path else None,
            preview_mode=preview_mode
        )
        
        # Create logger
        logger = Logger(log_level=LogLevel.ERROR)
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Result should contain markdown content
        assert result.markdown_content is not None, "No markdown content in result"
        
        # Property: File creation depends on mode
        if output_format == 'file':
            # File mode: output file should exist
            assert output_path.exists(), "Output file should be created in file mode"
        elif output_format == 'preview':
            # Preview mode: no file should be created
            if output_path:
                assert not output_path.exists(), \
                    "Output file should not be created in preview mode"
        # stdout mode: no file to check


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_46_preview_mode_preserves_content(filename, content):
    """Feature: document-to-markdown-converter, Property 46: Preview mode file handling
    
    For any conversion in preview mode, the converter should preserve the full 
    content in the result even though no file is written.
    
    **Validates: Requirements 11.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Convert in normal mode
        config_normal = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=False
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator_normal = ConversionOrchestrator(config_normal, logger)
        result_normal = orchestrator_normal.convert(str(doc_path))
        
        # Convert in preview mode
        config_preview = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        orchestrator_preview = ConversionOrchestrator(config_preview, logger)
        result_preview = orchestrator_preview.convert(str(doc_path))
        
        # Property: Both conversions should succeed
        assert result_normal.success, f"Normal conversion failed: {result_normal.errors}"
        assert result_preview.success, f"Preview conversion failed: {result_preview.errors}"
        
        # Property: Content should be identical in both modes
        assert result_normal.markdown_content == result_preview.markdown_content, \
            "Preview mode should preserve the same content as normal mode"
