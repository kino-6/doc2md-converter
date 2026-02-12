"""
End-to-end test for Word to Markdown conversion.

This test verifies that the complete conversion pipeline works from
Word document input to Markdown output.
"""

import pytest
from docx import Document
from src.parsers import WordParser
from src.markdown_serializer import MarkdownSerializer
from src.pretty_printer import PrettyPrinter


class TestWordToMarkdownE2E:
    """End-to-end tests for Word to Markdown conversion."""
    
    def test_complete_word_to_markdown_conversion(self, tmp_path):
        """Test complete conversion from Word document to Markdown."""
        # Create a Word document with various elements
        doc = Document()
        
        # Add title
        doc.add_heading("Test Document", level=1)
        
        # Add introduction paragraph
        doc.add_paragraph("This is a test document with various elements.")
        
        # Add section with heading
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("This section contains a list:")
        
        # Add unordered list
        doc.add_paragraph("Item 1", style='List Bullet')
        doc.add_paragraph("Item 2", style='List Bullet')
        doc.add_paragraph("Item 3", style='List Bullet')
        
        # Add another section
        doc.add_heading("Section 2", level=2)
        
        # Add formatted text
        para = doc.add_paragraph()
        para.add_run("This text is ")
        bold_run = para.add_run("bold")
        bold_run.bold = True
        para.add_run(" and this is ")
        italic_run = para.add_run("italic")
        italic_run.italic = True
        para.add_run(".")
        
        # Add table
        doc.add_heading("Section 3: Table", level=2)
        table = doc.add_table(rows=3, cols=3)
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Name"
        header_cells[1].text = "Age"
        header_cells[2].text = "City"
        
        # Add data
        row1_cells = table.rows[1].cells
        row1_cells[0].text = "Alice"
        row1_cells[1].text = "30"
        row1_cells[2].text = "New York"
        
        row2_cells = table.rows[2].cells
        row2_cells[0].text = "Bob"
        row2_cells[1].text = "25"
        row2_cells[2].text = "London"
        
        # Save to temporary file
        file_path = tmp_path / "test_complete.docx"
        doc.save(str(file_path))
        
        # Step 1: Parse Word document
        parser = WordParser()
        internal_doc = parser.parse(str(file_path))
        
        # Verify parsing succeeded
        assert internal_doc is not None
        assert len(internal_doc.sections) > 0
        
        # Step 2: Serialize to Markdown
        serializer = MarkdownSerializer()
        markdown_output = serializer.serialize(internal_doc)
        
        # Verify serialization succeeded
        assert markdown_output is not None
        assert len(markdown_output) > 0
        
        # Step 3: Pretty print
        printer = PrettyPrinter()
        formatted_markdown = printer.format(markdown_output)
        
        # Verify formatting succeeded
        assert formatted_markdown is not None
        assert len(formatted_markdown) > 0
        
        # Verify content is present in the output
        assert "Test Document" in formatted_markdown
        assert "Section 1" in formatted_markdown
        assert "Section 2" in formatted_markdown
        assert "Section 3: Table" in formatted_markdown
        assert "Item 1" in formatted_markdown
        assert "Item 2" in formatted_markdown
        assert "Item 3" in formatted_markdown
        assert "Alice" in formatted_markdown
        assert "Bob" in formatted_markdown
        
        # Verify Markdown syntax is present
        assert "# Test Document" in formatted_markdown
        assert "## Section 1" in formatted_markdown
        assert "## Section 2" in formatted_markdown
        assert "- Item 1" in formatted_markdown or "* Item 1" in formatted_markdown
        assert "|" in formatted_markdown  # Table syntax
        
        # Print the output for manual verification
        print("\n" + "="*60)
        print("Generated Markdown Output:")
        print("="*60)
        print(formatted_markdown)
        print("="*60)
        
        return formatted_markdown
    
    def test_simple_word_to_markdown(self, tmp_path):
        """Test simple Word to Markdown conversion with minimal content."""
        # Create a simple Word document
        doc = Document()
        doc.add_heading("Hello World", level=1)
        doc.add_paragraph("This is a simple test.")
        
        # Save to temporary file
        file_path = tmp_path / "test_simple.docx"
        doc.save(str(file_path))
        
        # Parse, serialize, and format
        parser = WordParser()
        internal_doc = parser.parse(str(file_path))
        
        serializer = MarkdownSerializer()
        markdown_output = serializer.serialize(internal_doc)
        
        printer = PrettyPrinter()
        formatted_markdown = printer.format(markdown_output)
        
        # Verify output
        assert "# Hello World" in formatted_markdown
        assert "This is a simple test." in formatted_markdown
        
        print("\n" + "="*60)
        print("Simple Markdown Output:")
        print("="*60)
        print(formatted_markdown)
        print("="*60)
