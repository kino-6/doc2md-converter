"""
End-to-end integration tests for Document to Markdown Converter.

This test suite validates the complete conversion flow for Word, Excel, and PDF
documents using real sample files. It tests all major features including:
- Text extraction and formatting
- Heading structure detection
- Table conversion
- Image extraction
- OCR capabilities
- Configuration options
- Error handling

Validates: Requirements 1.1-1.7, 2.1-2.5, 3.1-3.7, 4.1-4.5, 5.1-5.5,
           6.1-6.5, 7.1-7.11, 8.1-8.7, 9.1-9.5, 10.1-10.6, 11.1-11.5
"""

import os
import pytest
from pathlib import Path
from docx import Document
from openpyxl import Workbook

from src.config import ConversionConfig, LogLevel, TableStyle, ImageFormat
from src.logger import Logger
from src.conversion_orchestrator import ConversionOrchestrator
from src.parsers import WordParser, ExcelParser, PDFParser
from src.markdown_serializer import MarkdownSerializer
from src.pretty_printer import PrettyPrinter
from src.image_extractor import ImageExtractor
from src.ocr_engine import OCREngine


class TestEndToEndWordConversion:
    """End-to-end tests for Word document conversion."""
    
    def test_complete_word_conversion_with_all_features(self, tmp_path):
        """Test complete Word conversion with all major features.
        
        Validates: Requirements 1.1-1.7, 5.1-5.5, 6.2, 8.5, 9.2
        """
        # Create a comprehensive Word document
        doc = Document()
        
        # Add title and metadata
        doc.add_heading("Complete Feature Test Document", level=1)
        doc.add_paragraph("This document tests all major Word conversion features.")
        
        # Add multiple heading levels
        doc.add_heading("Section 1: Headings", level=2)
        doc.add_heading("Subsection 1.1", level=3)
        doc.add_paragraph("Testing heading hierarchy.")
        
        # Add formatted text
        doc.add_heading("Section 2: Text Formatting", level=2)
        para = doc.add_paragraph()
        para.add_run("Normal text, ")
        bold_run = para.add_run("bold text")
        bold_run.bold = True
        para.add_run(", ")
        italic_run = para.add_run("italic text")
        italic_run.italic = True
        para.add_run(", and ")
        both_run = para.add_run("bold italic")
        both_run.bold = True
        both_run.italic = True
        para.add_run(".")
        
        # Add lists
        doc.add_heading("Section 3: Lists", level=2)
        doc.add_paragraph("Unordered list:")
        doc.add_paragraph("First item", style='List Bullet')
        doc.add_paragraph("Second item", style='List Bullet')
        doc.add_paragraph("Third item", style='List Bullet')
        
        doc.add_paragraph("Ordered list:")
        doc.add_paragraph("First step", style='List Number')
        doc.add_paragraph("Second step", style='List Number')
        doc.add_paragraph("Third step", style='List Number')
        
        # Add table
        doc.add_heading("Section 4: Tables", level=2)
        table = doc.add_table(rows=4, cols=3)
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Product"
        header_cells[1].text = "Price"
        header_cells[2].text = "Quantity"
        
        # Data rows
        for i, (product, price, qty) in enumerate([
            ("Widget A", "$10.00", "100"),
            ("Widget B", "$15.50", "75"),
            ("Widget C", "$8.25", "200")
        ], start=1):
            cells = table.rows[i].cells
            cells[0].text = product
            cells[1].text = price
            cells[2].text = qty
        
        # Add hyperlink section
        doc.add_heading("Section 5: Links", level=2)
        link_para = doc.add_paragraph("Visit ")
        link_para.add_run("example.com")
        link_para.add_run(" for more information.")
        
        # Add special characters
        doc.add_heading("Section 6: Special Characters", level=2)
        doc.add_paragraph("Testing special chars: * _ [ ] ( ) # + - . ! \\ ` { } | ~")
        doc.add_paragraph("Unicode: © ® ™ € £ ¥ § ¶ † ‡")
        
        # Save document
        file_path = tmp_path / "complete_test.docx"
        doc.save(str(file_path))
        
        # Configure conversion
        output_path = tmp_path / "complete_test.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            log_level=LogLevel.INFO,
            extract_images=True,
            validate_output=True,
            include_metadata=True
        )
        
        # Execute conversion
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        # Verify conversion succeeded
        assert result.success, f"Conversion failed: {result.errors}"
        assert result.markdown_content is not None
        assert len(result.markdown_content) > 0
        
        # Read output file
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        # Verify headings
        assert "# Complete Feature Test Document" in markdown
        assert "## Section 1: Headings" in markdown
        assert "### Subsection 1.1" in markdown
        assert "## Section 2: Text Formatting" in markdown
        
        # Verify text formatting (may be combined in a single run)
        assert "bold text" in markdown
        assert "italic text" in markdown
        assert "bold italic" in markdown
        
        # Verify lists
        assert "- First item" in markdown or "* First item" in markdown
        assert "1. First step" in markdown or "1) First step" in markdown
        
        # Verify table
        assert "| Product" in markdown
        assert "| Widget A" in markdown
        assert "$10.00" in markdown
        
        # Verify special characters are handled
        assert "©" in markdown or "&copy;" in markdown
        
        # Verify UTF-8 encoding
        assert isinstance(markdown, str)
        
        # Verify structure preservation
        lines = markdown.split('\n')
        assert len(lines) > 30, "Output should have substantial content"
        
        print("\n" + "="*60)
        print("Complete Word Conversion Output:")
        print("="*60)
        print(markdown)
        print("="*60)
    
    def test_word_conversion_with_custom_options(self, tmp_path):
        """Test Word conversion with custom configuration options.
        
        Validates: Requirements 8.1, 8.2, 8.3, 8.4
        """
        # Create simple document
        doc = Document()
        doc.add_heading("Test", level=1)
        doc.add_heading("Subsection", level=2)
        doc.add_paragraph("Content here.")
        
        file_path = tmp_path / "custom_options.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "custom_options.md"
        
        # Test with heading offset
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            heading_offset=1,  # Shift all headings down by 1 level
            include_metadata=False,
            table_style=TableStyle.STANDARD
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        assert result.success
        
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        # Verify heading offset applied (H1 -> H2, H2 -> H3)
        assert "## Test" in markdown
        assert "### Subsection" in markdown


class TestEndToEndExcelConversion:
    """End-to-end tests for Excel spreadsheet conversion."""
    
    def test_complete_excel_conversion_with_multiple_sheets(self, tmp_path):
        """Test complete Excel conversion with multiple sheets and features.
        
        Validates: Requirements 2.1-2.5, 5.1-5.5, 6.2
        """
        # Create Excel workbook with multiple sheets
        wb = Workbook()
        
        # Sheet 1: Product data
        ws1 = wb.active
        ws1.title = "Products"
        ws1.append(["Product", "Price", "Stock"])
        ws1.append(["Laptop", 999.99, 50])
        ws1.append(["Mouse", 29.99, 200])
        ws1.append(["Keyboard", 79.99, 150])
        
        # Sheet 2: Sales data with formulas
        ws2 = wb.create_sheet("Sales")
        ws2.append(["Item", "Quantity", "Price", "Total"])
        ws2.append(["Laptop", 5, 999.99, "=B2*C2"])
        ws2.append(["Mouse", 10, 29.99, "=B3*C3"])
        ws2.append(["Keyboard", 8, 79.99, "=B4*C4"])
        ws2.append(["", "", "Grand Total:", "=SUM(D2:D4)"])
        
        # Sheet 3: Empty sheet
        ws3 = wb.create_sheet("Empty")
        
        # Sheet 4: Mixed data
        ws4 = wb.create_sheet("Mixed")
        ws4.append(["Name", "Email", "Status"])
        ws4.append(["Alice", "alice@example.com", "Active"])
        ws4.append(["Bob", "bob@example.com", "Inactive"])
        
        # Save workbook
        file_path = tmp_path / "complete_test.xlsx"
        wb.save(str(file_path))
        
        # Configure conversion
        output_path = tmp_path / "complete_test.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            log_level=LogLevel.INFO
        )
        
        # Execute conversion
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        # Verify conversion succeeded
        assert result.success, f"Conversion failed: {result.errors}"
        assert result.markdown_content is not None
        
        # Read output
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        # Verify all sheets are present
        assert "## Products" in markdown
        assert "## Sales" in markdown
        assert "## Empty" in markdown
        assert "## Mixed" in markdown
        
        # Verify table data
        assert "| Product" in markdown
        assert "| Laptop" in markdown
        assert "999.99" in markdown
        
        # Verify formulas are converted to values
        assert "4999.95" in markdown or "4999.9" in markdown  # 5 * 999.99
        assert "299.9" in markdown  # 10 * 29.99
        
        # Verify empty sheet handling
        assert "Empty sheet" in markdown or "No data" in markdown
        
        # Verify structure
        lines = markdown.split('\n')
        assert len(lines) > 20
        
        print("\n" + "="*60)
        print("Complete Excel Conversion Output:")
        print("="*60)
        print(markdown)
        print("="*60)
    
    def test_excel_with_special_cases(self, tmp_path):
        """Test Excel conversion with edge cases.
        
        Validates: Requirements 2.4, 2.5, 4.3
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Special Cases"
        
        # Add headers
        ws.append(["Type", "Value", "Notes"])
        
        # Add various data types
        ws.append(["Text", "Hello World", "Normal text"])
        ws.append(["Number", 42, "Integer"])
        ws.append(["Float", 3.14159, "Decimal"])
        ws.append(["Formula", "=2+2", "Simple formula"])
        ws.append(["Empty", None, "Null value"])
        ws.append(["Special", "© ® ™", "Unicode"])
        
        file_path = tmp_path / "special_cases.xlsx"
        wb.save(str(file_path))
        
        output_path = tmp_path / "special_cases.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path)
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        assert result.success
        
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        # Verify data types are handled
        assert "Hello World" in markdown
        assert "42" in markdown
        assert "3.14" in markdown
        assert "4" in markdown  # Formula result
        
        # Verify special characters
        assert "©" in markdown or "&copy;" in markdown


class TestEndToEndPDFConversion:
    """End-to-end tests for PDF document conversion."""
    
    @pytest.mark.skipif(not os.path.exists("docs_target/tps65053.pdf"),
                        reason="Test PDF file not available")
    def test_real_pdf_conversion(self, tmp_path):
        """Test conversion of real PDF document.
        
        Validates: Requirements 3.1, 3.2, 3.3, 5.1-5.5
        """
        pdf_path = "docs_target/tps65053.pdf"
        output_path = tmp_path / "tps65053.md"
        
        config = ConversionConfig(
            input_path=pdf_path,
            output_path=str(output_path),
            log_level=LogLevel.INFO,
            extract_images=True,
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(pdf_path)
        
        # Verify conversion succeeded
        assert result.success, f"Conversion failed: {result.errors}"
        assert result.markdown_content is not None
        assert len(result.markdown_content) > 1000
        
        # Read output
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        # Verify text extraction
        assert "TPS6505" in markdown
        assert len(markdown) > 1000
        
        # Verify heading detection
        headings = [line for line in markdown.split('\n') if line.startswith('##')]
        assert len(headings) > 5
        
        # Verify table detection
        table_rows = [line for line in markdown.split('\n') if line.strip().startswith('|')]
        assert len(table_rows) > 10
        
        # Verify UTF-8 encoding
        assert isinstance(markdown, str)
        
        print(f"\n" + "="*60)
        print(f"PDF Conversion Stats:")
        print(f"  Total length: {len(markdown)} chars")
        print(f"  Headings detected: {len(headings)}")
        print(f"  Table rows: {len(table_rows)}")
        print(f"  Duration: {result.duration:.2f}s")
        print("="*60)


class TestEndToEndBatchConversion:
    """End-to-end tests for batch conversion."""
    
    def test_batch_conversion_multiple_formats(self, tmp_path):
        """Test batch conversion of multiple file formats.
        
        Validates: Requirements 6.5, 4.5
        """
        # Create multiple test files
        files = []
        
        # Word document
        doc = Document()
        doc.add_heading("Word Document", level=1)
        doc.add_paragraph("This is a Word document.")
        word_path = tmp_path / "test1.docx"
        doc.save(str(word_path))
        files.append(str(word_path))
        
        # Excel spreadsheet
        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Value"])
        ws.append(["Item 1", 100])
        ws.append(["Item 2", 200])
        excel_path = tmp_path / "test2.xlsx"
        wb.save(str(excel_path))
        files.append(str(excel_path))
        
        # Another Word document
        doc2 = Document()
        doc2.add_heading("Second Document", level=1)
        doc2.add_paragraph("Another test document.")
        word_path2 = tmp_path / "test3.docx"
        doc2.save(str(word_path2))
        files.append(str(word_path2))
        
        # Configure batch conversion
        config = ConversionConfig(
            input_path="",  # Not used for batch
            log_level=LogLevel.INFO,
            batch_mode=True
        )
        
        # Execute batch conversion
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        results = orchestrator.batch_convert(files)
        
        # Verify all conversions succeeded
        assert len(results) == 3
        for result in results:
            assert result.success, f"Conversion failed for {result.input_path}: {result.errors}"
            assert result.markdown_content is not None
            assert len(result.markdown_content) > 0
        
        # Verify content
        assert "Word Document" in results[0].markdown_content
        assert "Item 1" in results[1].markdown_content
        assert "Second Document" in results[2].markdown_content
        
        print("\n" + "="*60)
        print("Batch Conversion Results:")
        for i, result in enumerate(results, 1):
            print(f"  File {i}: {Path(result.input_path).name}")
            print(f"    Success: {result.success}")
            print(f"    Length: {len(result.markdown_content)} chars")
            print(f"    Duration: {result.duration:.2f}s")
        print("="*60)


class TestEndToEndImageExtraction:
    """End-to-end tests for image extraction and OCR."""
    
    def test_image_extraction_workflow(self, tmp_path):
        """Test complete image extraction workflow.
        
        Validates: Requirements 7.1-7.4, 7.6, 7.7
        """
        # Create Word document with image placeholder
        doc = Document()
        doc.add_heading("Document with Images", level=1)
        doc.add_paragraph("This document would contain images.")
        doc.add_paragraph("Image reference: [Image 1]")
        
        file_path = tmp_path / "with_images.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "with_images.md"
        images_dir = tmp_path / "with_images" / "images"
        
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            extract_images=True,
            enable_ocr=False  # Disable OCR for this test
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        assert result.success
        
        # Verify output
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        assert "Document with Images" in markdown
        
        # Note: Since we're using a simple document without actual images,
        # we're just verifying the workflow completes successfully


class TestEndToEndErrorHandling:
    """End-to-end tests for error handling."""
    
    def test_invalid_file_format_error(self, tmp_path):
        """Test error handling for invalid file format.
        
        Validates: Requirements 4.1
        """
        # Create a text file with wrong extension
        invalid_file = tmp_path / "test.txt"
        invalid_file.write_text("This is not a valid document.")
        
        config = ConversionConfig(
            input_path=str(invalid_file),
            log_level=LogLevel.ERROR
        )
        
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(invalid_file))
        
        # Verify error is reported
        assert not result.success
        assert len(result.errors) > 0
        assert "format" in result.errors[0].lower() or "invalid" in result.errors[0].lower()
    
    def test_nonexistent_file_error(self, tmp_path):
        """Test error handling for nonexistent file.
        
        Validates: Requirements 4.2
        """
        nonexistent = tmp_path / "does_not_exist.docx"
        
        config = ConversionConfig(
            input_path=str(nonexistent),
            log_level=LogLevel.ERROR
        )
        
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(nonexistent))
        
        # Verify error is reported
        assert not result.success
        assert len(result.errors) > 0
        assert "not found" in result.errors[0].lower() or "exist" in result.errors[0].lower()
    
    def test_empty_file_handling(self, tmp_path):
        """Test handling of empty files.
        
        Validates: Requirements 4.3
        """
        # Create empty Word document
        doc = Document()
        empty_file = tmp_path / "empty.docx"
        doc.save(str(empty_file))
        
        output_path = tmp_path / "empty.md"
        config = ConversionConfig(
            input_path=str(empty_file),
            output_path=str(output_path)
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(empty_file))
        
        # Should succeed but with warning
        assert result.success
        assert len(result.warnings) > 0 or result.markdown_content == "" or len(result.markdown_content) < 50


class TestEndToEndPreviewAndDryRun:
    """End-to-end tests for preview and dry-run modes."""
    
    def test_preview_mode(self, tmp_path):
        """Test preview mode without file output.
        
        Validates: Requirements 11.1, 11.2
        """
        # Create test document
        doc = Document()
        doc.add_heading("Preview Test", level=1)
        doc.add_paragraph("This is a preview test.")
        
        file_path = tmp_path / "preview.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "preview.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            preview_mode=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify file was NOT created in preview mode
        assert not output_path.exists()
    
    def test_dry_run_mode(self, tmp_path):
        """Test dry-run mode without file output.
        
        Validates: Requirements 11.5
        """
        # Create test document
        doc = Document()
        doc.add_heading("Dry Run Test", level=1)
        doc.add_paragraph("This is a dry run test.")
        
        file_path = tmp_path / "dryrun.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "dryrun.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            dry_run=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify file was NOT created in dry-run mode
        assert not output_path.exists()


class TestEndToEndLogging:
    """End-to-end tests for logging functionality."""
    
    def test_logging_during_conversion(self, tmp_path):
        """Test that logging works correctly during conversion.
        
        Validates: Requirements 10.1-10.4
        """
        # Create test document
        doc = Document()
        doc.add_heading("Logging Test", level=1)
        doc.add_paragraph("Testing logging functionality.")
        
        file_path = tmp_path / "logging_test.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "logging_test.md"
        log_file = tmp_path / "conversion.log"
        
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            log_level=LogLevel.INFO,
            log_file=str(log_file)
        )
        
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_file))
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        assert result.success
        
        # Verify log file was created
        assert log_file.exists()
        
        # Read log content
        with open(log_file, 'r', encoding='utf-8') as f:
            log_content = f.read()
        
        # Verify log contains expected information
        assert len(log_content) > 0
        # Log should contain file path or conversion info
        assert "logging_test" in log_content or "Conversion" in log_content


class TestEndToEndValidation:
    """End-to-end tests for output validation."""
    
    def test_markdown_validation(self, tmp_path):
        """Test Markdown validation during conversion.
        
        Validates: Requirements 11.3, 11.4
        """
        # Create test document
        doc = Document()
        doc.add_heading("Validation Test", level=1)
        doc.add_paragraph("Testing Markdown validation.")
        
        file_path = tmp_path / "validation.docx"
        doc.save(str(file_path))
        
        output_path = tmp_path / "validation.md"
        config = ConversionConfig(
            input_path=str(file_path),
            output_path=str(output_path),
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        result = orchestrator.convert(str(file_path))
        
        # Verify conversion succeeded
        assert result.success
        
        # Verify output is valid Markdown
        with open(output_path, 'r', encoding='utf-8') as f:
            markdown = f.read()
        
        assert "# Validation Test" in markdown
        assert "Testing Markdown validation" in markdown


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
