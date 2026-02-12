"""
Property-based tests for Word document structural elements.

This module contains property-based tests using Hypothesis to verify
Word document structural element parsing (headings, lists, tables).

Feature: document-to-markdown-converter
Properties: 2, 3, 4
Validates: Requirements 1.2, 1.3, 1.4, 5.5
"""

import pytest
from hypothesis import given, strategies as st, settings, HealthCheck
from pathlib import Path
import tempfile
import string

from src.parsers import WordParser
from src.internal_representation import (
    InternalDocument, Section, Heading, Table, DocumentList, ListItem
)


# Helper strategies for generating Word-compatible data
def word_safe_text(min_size=1, max_size=100):
    """Generate text that is safe for Word documents."""
    # Word-safe characters (printable ASCII without control characters)
    safe_chars = string.printable.replace('\x0b', '').replace('\x0c', '').replace('\r', '').replace('\n', '')
    return st.text(alphabet=safe_chars, min_size=min_size, max_size=max_size)


class TestWordPropertyHeadingLevelMapping:
    """Property 2: Heading level mapping
    
    For any Word document with headings at various levels (H1-H6), the converter
    should map them to corresponding Markdown heading syntax (# for H1, ## for H2, etc.)
    
    Validates: Requirements 1.2
    """
    
    @given(
        heading_level=st.integers(min_value=1, max_value=6),
        heading_text=word_safe_text(min_size=1, max_size=100)
    )
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_2_heading_level_mapping(self, heading_level, heading_text):
        """
        Feature: document-to-markdown-converter
        Property 2: Heading level mapping
        
        For any Word document with headings at various levels (H1-H6), the converter
        should map them to corresponding Markdown heading syntax.
        
        **Validates: Requirements 1.2**
        """
        from docx import Document
        
        # Skip text that is only whitespace
        if not heading_text.strip():
            return
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with a heading at the specified level
            doc = Document()
            doc.add_heading(heading_text, level=heading_level)
            doc.add_paragraph("Content under the heading")
            
            # Save to temporary file
            file_path = Path(tmp_dir) / f"test_heading_level_{heading_level}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Document should be parsed successfully
            assert isinstance(result, InternalDocument)
            assert len(result.sections) > 0
            
            # Property: First section should have a heading
            first_section = result.sections[0]
            assert first_section.heading is not None, \
                "First section should have a heading"
            
            # Property: Heading level should match the input level
            assert isinstance(first_section.heading, Heading), \
                "Heading should be a Heading object"
            assert first_section.heading.level == heading_level, \
                f"Expected heading level {heading_level}, got {first_section.heading.level}"
            
            # Property: Heading text should match the input text
            # Word normalizes whitespace, so we compare normalized versions
            import re
            normalized_input = re.sub(r'\s+', ' ', heading_text).strip()
            normalized_output = re.sub(r'\s+', ' ', first_section.heading.text).strip()
            
            assert normalized_output == normalized_input, \
                f"Expected heading text '{normalized_input}', got '{normalized_output}'"
    
    @given(
        num_headings=st.integers(min_value=2, max_value=6),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_2_multiple_heading_levels(self, num_headings, data):
        """
        Feature: document-to-markdown-converter
        Property 2: Heading level mapping (multiple headings)
        
        For documents with multiple headings at different levels, all heading
        levels should be correctly mapped.
        
        **Validates: Requirements 1.2**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with multiple headings
            doc = Document()
            
            heading_data = []
            for i in range(num_headings):
                level = data.draw(st.integers(min_value=1, max_value=6))
                text = data.draw(word_safe_text(min_size=1, max_size=50))
                
                # Skip if text is only whitespace
                if not text.strip():
                    continue
                
                doc.add_heading(text, level=level)
                doc.add_paragraph(f"Content under heading {i+1}")
                
                heading_data.append((level, text))
            
            # Skip test if no valid headings were generated
            if not heading_data:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_multiple_headings.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: All headings should be extracted
            assert len(result.sections) == len(heading_data), \
                f"Expected {len(heading_data)} sections, got {len(result.sections)}"
            
            # Property: Each heading should have the correct level and text
            import re
            for i, (expected_level, expected_text) in enumerate(heading_data):
                section = result.sections[i]
                assert section.heading is not None, \
                    f"Section {i} should have a heading"
                
                assert section.heading.level == expected_level, \
                    f"Section {i}: Expected level {expected_level}, got {section.heading.level}"
                
                # Normalize whitespace for comparison
                normalized_expected = re.sub(r'\s+', ' ', expected_text).strip()
                normalized_actual = re.sub(r'\s+', ' ', section.heading.text).strip()
                
                assert normalized_actual == normalized_expected, \
                    f"Section {i}: Expected text '{normalized_expected}', got '{normalized_actual}'"


class TestWordPropertyListStructurePreservation:
    """Property 3: List structure preservation
    
    For any Word document containing ordered or unordered lists, the converter
    should preserve the list structure using appropriate Markdown list syntax.
    
    Validates: Requirements 1.3
    """
    
    @given(
        num_items=st.integers(min_value=1, max_value=10),
        is_ordered=st.booleans(),
        data=st.data()
    )
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_3_list_structure_preservation(self, num_items, is_ordered, data):
        """
        Feature: document-to-markdown-converter
        Property 3: List structure preservation
        
        For any Word document containing ordered or unordered lists, the converter
        should preserve the list structure.
        
        **Validates: Requirements 1.3**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with a list
            doc = Document()
            doc.add_paragraph("Document with a list:")
            
            list_items = []
            for i in range(num_items):
                item_text = data.draw(word_safe_text(min_size=1, max_size=50))
                
                # Skip if text is only whitespace
                if not item_text.strip():
                    continue
                
                # Add list item with appropriate style
                if is_ordered:
                    doc.add_paragraph(item_text, style='List Number')
                else:
                    doc.add_paragraph(item_text, style='List Bullet')
                
                list_items.append(item_text)
            
            # Skip test if no valid list items were generated
            if not list_items:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / f"test_list_{'ordered' if is_ordered else 'unordered'}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Document should be parsed successfully
            assert isinstance(result, InternalDocument)
            assert len(result.sections) > 0
            
            # Property: Should find a DocumentList in the content
            found_list = None
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, DocumentList):
                        found_list = content_item
                        break
                if found_list:
                    break
            
            assert found_list is not None, \
                "Should find a DocumentList in the parsed content"
            
            # Property: List type should match (ordered/unordered)
            assert found_list.ordered == is_ordered, \
                f"Expected list to be {'ordered' if is_ordered else 'unordered'}, got {'ordered' if found_list.ordered else 'unordered'}"
            
            # Property: Number of list items should match
            assert len(found_list.items) == len(list_items), \
                f"Expected {len(list_items)} list items, got {len(found_list.items)}"
            
            # Property: List item texts should match
            import re
            for i, (expected_text, actual_item) in enumerate(zip(list_items, found_list.items)):
                assert isinstance(actual_item, ListItem), \
                    f"Item {i} should be a ListItem"
                
                # Normalize whitespace for comparison
                normalized_expected = re.sub(r'\s+', ' ', expected_text).strip()
                normalized_actual = re.sub(r'\s+', ' ', actual_item.text).strip()
                
                assert normalized_actual == normalized_expected, \
                    f"Item {i}: Expected text '{normalized_expected}', got '{normalized_actual}'"
    
    @given(
        num_items=st.integers(min_value=2, max_value=8),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_3_nested_list_structure(self, num_items, data):
        """
        Feature: document-to-markdown-converter
        Property 3: List structure preservation (nested lists)
        
        For documents with nested lists, the structure and indentation levels
        should be preserved.
        
        **Validates: Requirements 1.3**
        """
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with nested lists
            doc = Document()
            doc.add_paragraph("Document with nested list:")
            
            list_items_with_levels = []
            for i in range(num_items):
                item_text = data.draw(word_safe_text(min_size=1, max_size=40))
                level = data.draw(st.integers(min_value=0, max_value=2))
                
                # Skip if text is only whitespace
                if not item_text.strip():
                    continue
                
                # Add list item
                para = doc.add_paragraph(item_text, style='List Bullet')
                
                # Set indentation level
                if level > 0:
                    # Access the paragraph's numbering properties
                    pPr = para._element.get_or_add_pPr()
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is None:
                        numPr = OxmlElement('w:numPr')
                        pPr.append(numPr)
                    
                    # Set indentation level
                    ilvl = numPr.find(qn('w:ilvl'))
                    if ilvl is None:
                        ilvl = OxmlElement('w:ilvl')
                        numPr.append(ilvl)
                    ilvl.set(qn('w:val'), str(level))
                
                list_items_with_levels.append((item_text, level))
            
            # Skip test if no valid list items were generated
            if not list_items_with_levels:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_nested_list.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Should find a DocumentList in the content
            found_list = None
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, DocumentList):
                        found_list = content_item
                        break
                if found_list:
                    break
            
            assert found_list is not None, \
                "Should find a DocumentList in the parsed content"
            
            # Property: Number of list items should match
            assert len(found_list.items) == len(list_items_with_levels), \
                f"Expected {len(list_items_with_levels)} list items, got {len(found_list.items)}"
            
            # Property: List item levels should be preserved
            import re
            for i, ((expected_text, expected_level), actual_item) in enumerate(zip(list_items_with_levels, found_list.items)):
                assert isinstance(actual_item, ListItem), \
                    f"Item {i} should be a ListItem"
                
                # Check level
                assert actual_item.level == expected_level, \
                    f"Item {i}: Expected level {expected_level}, got {actual_item.level}"
                
                # Check text
                normalized_expected = re.sub(r'\s+', ' ', expected_text).strip()
                normalized_actual = re.sub(r'\s+', ' ', actual_item.text).strip()
                
                assert normalized_actual == normalized_expected, \
                    f"Item {i}: Expected text '{normalized_expected}', got '{normalized_actual}'"


class TestWordPropertyTableConversion:
    """Property 4: Table conversion
    
    For any Word document containing tables, the converter should convert them
    to valid Markdown table format with proper column alignment.
    
    Validates: Requirements 1.4, 5.5
    """
    
    @given(
        num_rows=st.integers(min_value=2, max_value=10),
        num_cols=st.integers(min_value=1, max_value=8),
        data=st.data()
    )
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_4_table_conversion(self, num_rows, num_cols, data):
        """
        Feature: document-to-markdown-converter
        Property 4: Table conversion
        
        For any Word document containing tables, the converter should convert them
        to valid Markdown table format.
        
        **Validates: Requirements 1.4, 5.5**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with a table
            doc = Document()
            doc.add_paragraph("Document with a table:")
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            
            # Generate table data
            table_data = []
            for row_idx in range(num_rows):
                row_data = []
                for col_idx in range(num_cols):
                    cell_text = data.draw(word_safe_text(min_size=0, max_size=30))
                    row_data.append(cell_text)
                    table.rows[row_idx].cells[col_idx].text = cell_text
                table_data.append(row_data)
            
            # Save to temporary file
            file_path = Path(tmp_dir) / f"test_table_{num_rows}x{num_cols}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Document should be parsed successfully
            assert isinstance(result, InternalDocument)
            assert len(result.sections) > 0
            
            # Property: Should find a Table in the content
            found_table = None
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, Table):
                        found_table = content_item
                        break
                if found_table:
                    break
            
            assert found_table is not None, \
                "Should find a Table in the parsed content"
            
            # Property: Table should have correct dimensions
            assert len(found_table.headers) == num_cols, \
                f"Expected {num_cols} columns, got {len(found_table.headers)}"
            assert len(found_table.rows) == num_rows - 1, \
                f"Expected {num_rows - 1} data rows (excluding header), got {len(found_table.rows)}"
            
            # Property: Headers should match first row
            import re
            for col_idx in range(num_cols):
                expected_header = re.sub(r'\s+', ' ', table_data[0][col_idx]).strip()
                actual_header = re.sub(r'\s+', ' ', found_table.headers[col_idx]).strip()
                
                assert actual_header == expected_header, \
                    f"Header {col_idx}: Expected '{expected_header}', got '{actual_header}'"
            
            # Property: Data rows should match remaining rows
            for row_idx in range(1, num_rows):
                for col_idx in range(num_cols):
                    expected_cell = re.sub(r'\s+', ' ', table_data[row_idx][col_idx]).strip()
                    actual_cell = re.sub(r'\s+', ' ', found_table.rows[row_idx - 1][col_idx]).strip()
                    
                    assert actual_cell == expected_cell, \
                        f"Cell ({row_idx}, {col_idx}): Expected '{expected_cell}', got '{actual_cell}'"
            
            # Property: Each row should have the same number of columns
            for row in found_table.rows:
                assert len(row) == num_cols, \
                    f"Each row should have {num_cols} columns, got {len(row)}"
    
    @given(
        num_tables=st.integers(min_value=2, max_value=4),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_4_multiple_tables(self, num_tables, data):
        """
        Feature: document-to-markdown-converter
        Property 4: Table conversion (multiple tables)
        
        For documents with multiple tables, all tables should be correctly
        converted and preserved.
        
        **Validates: Requirements 1.4, 5.5**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with multiple tables
            doc = Document()
            
            tables_data = []
            for i in range(num_tables):
                doc.add_paragraph(f"Table {i+1}:")
                
                num_rows = data.draw(st.integers(min_value=2, max_value=5))
                num_cols = data.draw(st.integers(min_value=2, max_value=5))
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                
                # Generate table data
                table_data = []
                for row_idx in range(num_rows):
                    row_data = []
                    for col_idx in range(num_cols):
                        cell_text = data.draw(word_safe_text(min_size=1, max_size=20))
                        row_data.append(cell_text)
                        table.rows[row_idx].cells[col_idx].text = cell_text
                    table_data.append(row_data)
                
                tables_data.append((num_rows, num_cols, table_data))
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_multiple_tables.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Should find all tables in the content
            found_tables = []
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, Table):
                        found_tables.append(content_item)
            
            assert len(found_tables) == num_tables, \
                f"Expected {num_tables} tables, got {len(found_tables)}"
            
            # Property: Each table should have correct structure
            import re
            for i, (found_table, (num_rows, num_cols, table_data)) in enumerate(zip(found_tables, tables_data)):
                assert len(found_table.headers) == num_cols, \
                    f"Table {i}: Expected {num_cols} columns, got {len(found_table.headers)}"
                assert len(found_table.rows) == num_rows - 1, \
                    f"Table {i}: Expected {num_rows - 1} data rows, got {len(found_table.rows)}"
                
                # Verify headers
                for col_idx in range(num_cols):
                    expected = re.sub(r'\s+', ' ', table_data[0][col_idx]).strip()
                    actual = re.sub(r'\s+', ' ', found_table.headers[col_idx]).strip()
                    assert actual == expected, \
                        f"Table {i}, Header {col_idx}: Expected '{expected}', got '{actual}'"
    
    @given(
        num_rows=st.integers(min_value=2, max_value=6),
        num_cols=st.integers(min_value=2, max_value=6)
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_4_empty_cells(self, num_rows, num_cols):
        """
        Feature: document-to-markdown-converter
        Property 4: Table conversion (empty cells)
        
        For tables with empty cells, the structure should still be preserved
        correctly.
        
        **Validates: Requirements 1.4, 5.5**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with a table containing empty cells
            doc = Document()
            doc.add_paragraph("Table with empty cells:")
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            
            # Fill only some cells (leave others empty)
            for row_idx in range(num_rows):
                for col_idx in range(num_cols):
                    # Fill cells in a checkerboard pattern
                    if (row_idx + col_idx) % 2 == 0:
                        table.rows[row_idx].cells[col_idx].text = f"R{row_idx}C{col_idx}"
                    # else: leave empty
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_table_empty_cells.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Property: Should find a Table in the content
            found_table = None
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, Table):
                        found_table = content_item
                        break
                if found_table:
                    break
            
            assert found_table is not None, \
                "Should find a Table in the parsed content"
            
            # Property: Table dimensions should be preserved
            assert len(found_table.headers) == num_cols, \
                f"Expected {num_cols} columns, got {len(found_table.headers)}"
            assert len(found_table.rows) == num_rows - 1, \
                f"Expected {num_rows - 1} data rows, got {len(found_table.rows)}"
            
            # Property: Each row should have the correct number of columns
            for row in found_table.rows:
                assert len(row) == num_cols, \
                    f"Each row should have {num_cols} columns, got {len(row)}"
            
            # Property: Empty cells should be represented as empty strings
            for row_idx in range(1, num_rows):
                for col_idx in range(num_cols):
                    cell_value = found_table.rows[row_idx - 1][col_idx]
                    
                    if (row_idx + col_idx) % 2 == 0:
                        # Should have content
                        expected = f"R{row_idx}C{col_idx}"
                        assert cell_value == expected, \
                            f"Cell ({row_idx}, {col_idx}): Expected '{expected}', got '{cell_value}'"
                    else:
                        # Should be empty
                        assert cell_value == "", \
                            f"Cell ({row_idx}, {col_idx}): Expected empty string, got '{cell_value}'"


# Run all property tests
if __name__ == "__main__":
    pytest.main([__file__, "-v"])
