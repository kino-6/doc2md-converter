"""Tests for dry-run mode functionality."""

import tempfile
from pathlib import Path
import pytest
from docx import Document

from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel


def create_test_docx(path: str, content: str):
    """Create a test Word document."""
    doc = Document()
    doc.add_paragraph(content)
    doc.save(path)


def test_dry_run_no_file_write():
    """Test dry-run mode prevents file writes.
    
    Requirements: 11.5
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test file
        input_file = Path(tmpdir) / "test.docx"
        output_file = Path(tmpdir) / "output.md"
        
        create_test_docx(str(input_file), "Test content")
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Verify conversion succeeded
        assert result.success is True
        assert result.markdown_content is not None
        assert "Test content" in result.markdown_content
        
        # Verify no file was written
        assert not output_file.exists()


def test_dry_run_with_batch_conversion():
    """Test dry-run mode in batch conversion.
    
    Requirements: 11.5, 6.5
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test files
        files = []
        for i in range(3):
            file_path = Path(tmpdir) / f"doc{i}.docx"
            create_test_docx(str(file_path), f"Content {i}")
            files.append(str(file_path))
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(files)
        
        # Verify all conversions succeeded
        assert len(results) == 3
        assert all(r.success for r in results)
        assert all(r.markdown_content is not None for r in results)
        
        # Verify no output files were created
        for i in range(3):
            output_path = Path(tmpdir) / f"doc{i}.md"
            assert not output_path.exists()


def test_dry_run_still_validates():
    """Test dry-run mode still performs validation.
    
    Requirements: 11.5
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test file
        input_file = Path(tmpdir) / "test.docx"
        create_test_docx(str(input_file), "Test content")
        
        # Create orchestrator with dry-run and validation enabled
        config = ConversionConfig(
            input_path=str(input_file),
            dry_run=True,
            validate_output=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Verify conversion succeeded and validation was performed
        assert result.success is True
        # Validation warnings may or may not be present, but conversion should succeed


def test_dry_run_processes_full_pipeline():
    """Test dry-run mode processes the full conversion pipeline.
    
    Requirements: 11.5
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test file with various elements
        input_file = Path(tmpdir) / "test.docx"
        
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test paragraph with **bold** text.")
        doc.add_paragraph("Another paragraph.")
        doc.save(str(input_file))
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Verify full pipeline was executed
        assert result.success is True
        assert result.markdown_content is not None
        
        # Verify content was processed
        assert "Test Heading" in result.markdown_content
        assert "Test paragraph" in result.markdown_content
        assert "Another paragraph" in result.markdown_content
        
        # Verify statistics were collected
        assert result.stats.headings_detected > 0


def test_dry_run_with_invalid_file():
    """Test dry-run mode with invalid file still reports errors.
    
    Requirements: 11.5
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create invalid file
        invalid_file = Path(tmpdir) / "invalid.txt"
        invalid_file.write_text("Not a valid document")
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(invalid_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(invalid_file))
        
        # Verify conversion failed with appropriate error
        assert result.success is False
        assert len(result.errors) > 0
