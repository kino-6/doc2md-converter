"""
Edge case tests for empty file handling.

Tests cover:
- Empty Word documents (.docx)
- Empty Excel files (.xlsx)
- Empty PDF files (.pdf)
- Zero-byte files

Requirements: 4.3
"""

import pytest
import os
from pathlib import Path
from src.logger import LogLevel


class TestEmptyFileEdgeCases:
    """Test empty file handling across all supported formats."""
    
    def test_empty_word_document(self, tmp_path):
        """Test that empty Word documents are handled correctly.
        
        Validates: Requirements 4.3 - Empty file handling for Word documents
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")
        
        # Create an empty Word document
        docx_path = tmp_path / "empty.docx"
        doc = docx.Document()
        # Don't add any content - leave it completely empty
        doc.save(docx_path)
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(docx_path),
            output_path=None,  # Output to stdout
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the empty document
        result = orchestrator.convert(str(docx_path))
        
        # Verify the result
        assert result.success is True, "Empty Word document conversion should succeed"
        assert result.markdown_content is not None, "Should return Markdown content"
        
        # Check for warning message about empty file
        assert len(result.warnings) > 0 or result.markdown_content.strip() == "", \
            "Should either have warnings or return empty Markdown"
        
        # The output should be minimal (empty or nearly empty)
        assert len(result.markdown_content.strip()) < 100, \
            "Empty document should produce minimal output"
    
    def test_empty_excel_file(self, tmp_path):
        """Test that empty Excel files are handled correctly.
        
        Validates: Requirements 4.3 - Empty file handling for Excel files
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")
        
        # Create an empty Excel file (with one empty sheet)
        excel_path = tmp_path / "empty.xlsx"
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Sheet1"
        # Don't add any data - leave it empty
        workbook.save(excel_path)
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the empty Excel file
        result = orchestrator.convert(str(excel_path))
        
        # Verify the result
        assert result.success is True, "Empty Excel file conversion should succeed"
        assert result.markdown_content is not None, "Should return Markdown content"
        
        # Excel files with empty sheets should indicate empty sheets
        assert "(Empty sheet)" in result.markdown_content or \
               len(result.markdown_content.strip()) < 100, \
            "Should indicate empty sheet or produce minimal output"
    
    def test_empty_pdf_file(self, tmp_path):
        """Test that empty PDF files are handled correctly.
        
        Validates: Requirements 4.3 - Empty file handling for PDF files
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
        except ImportError:
            pytest.skip("reportlab not installed")
        
        # Create an empty PDF file (with one blank page)
        pdf_path = tmp_path / "empty.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        # Add a blank page with no content
        c.showPage()
        c.save()
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(pdf_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the empty PDF
        result = orchestrator.convert(str(pdf_path))
        
        # Verify the result
        assert result.success is True, "Empty PDF conversion should succeed"
        assert result.markdown_content is not None, "Should return Markdown content"
        
        # Empty PDF should produce minimal output
        assert len(result.markdown_content.strip()) < 100, \
            "Empty PDF should produce minimal output"
    
    def test_zero_byte_file_word(self, tmp_path):
        """Test that zero-byte Word files are handled with appropriate error.
        
        Validates: Requirements 4.3 - Zero-byte file handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        # Create a zero-byte file with .docx extension
        docx_path = tmp_path / "zero_byte.docx"
        docx_path.touch()  # Creates empty file
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(docx_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the zero-byte file
        result = orchestrator.convert(str(docx_path))
        
        # Zero-byte files are corrupted and should fail or warn
        # The behavior depends on the parser implementation
        assert result.success is False or len(result.warnings) > 0, \
            "Zero-byte file should fail or produce warnings"
        
        if not result.success:
            assert len(result.errors) > 0, "Should have error messages"
    
    def test_zero_byte_file_excel(self, tmp_path):
        """Test that zero-byte Excel files are handled with appropriate error.
        
        Validates: Requirements 4.3 - Zero-byte file handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        # Create a zero-byte file with .xlsx extension
        excel_path = tmp_path / "zero_byte.xlsx"
        excel_path.touch()
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the zero-byte file
        result = orchestrator.convert(str(excel_path))
        
        # Zero-byte files should fail or warn
        assert result.success is False or len(result.warnings) > 0, \
            "Zero-byte file should fail or produce warnings"
        
        if not result.success:
            assert len(result.errors) > 0, "Should have error messages"
    
    def test_zero_byte_file_pdf(self, tmp_path):
        """Test that zero-byte PDF files are handled with appropriate error.
        
        Validates: Requirements 4.3 - Zero-byte file handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        # Create a zero-byte file with .pdf extension
        pdf_path = tmp_path / "zero_byte.pdf"
        pdf_path.touch()
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(pdf_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the zero-byte file
        result = orchestrator.convert(str(pdf_path))
        
        # Zero-byte files should fail or warn
        assert result.success is False or len(result.warnings) > 0, \
            "Zero-byte file should fail or produce warnings"
        
        if not result.success:
            assert len(result.errors) > 0, "Should have error messages"
    
    def test_word_document_with_only_whitespace(self, tmp_path):
        """Test Word document containing only whitespace.
        
        Validates: Requirements 4.3 - Empty content handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")
        
        # Create a Word document with only whitespace
        docx_path = tmp_path / "whitespace.docx"
        doc = docx.Document()
        doc.add_paragraph("   ")  # Only spaces
        doc.add_paragraph("\n\n")  # Only newlines
        doc.add_paragraph("\t\t")  # Only tabs
        doc.save(docx_path)
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(docx_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the document
        result = orchestrator.convert(str(docx_path))
        
        # Should succeed but produce minimal output
        assert result.success is True, "Whitespace-only document should succeed"
        assert result.markdown_content is not None
        
        # Output should be empty or nearly empty after stripping whitespace
        assert len(result.markdown_content.strip()) < 50, \
            "Whitespace-only document should produce minimal output"
    
    def test_excel_with_all_empty_sheets(self, tmp_path):
        """Test Excel file with multiple empty sheets.
        
        Validates: Requirements 4.3 - Multiple empty sheets handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")
        
        # Create Excel file with multiple empty sheets
        excel_path = tmp_path / "all_empty.xlsx"
        workbook = openpyxl.Workbook()
        
        sheet1 = workbook.active
        sheet1.title = "Empty1"
        
        workbook.create_sheet("Empty2")
        workbook.create_sheet("Empty3")
        
        workbook.save(excel_path)
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the Excel file
        result = orchestrator.convert(str(excel_path))
        
        # Should succeed
        assert result.success is True, "Excel with all empty sheets should succeed"
        assert result.markdown_content is not None
        
        # Should indicate empty sheets
        assert result.markdown_content.count("(Empty sheet)") == 3 or \
               len(result.markdown_content.strip()) < 200, \
            "Should indicate all sheets are empty"
    
    def test_pdf_with_only_blank_pages(self, tmp_path):
        """Test PDF with multiple blank pages.
        
        Validates: Requirements 4.3 - Blank pages handling
        """
        from src.conversion_orchestrator import ConversionOrchestrator
        from src.config import ConversionConfig
        from src.logger import Logger
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
        except ImportError:
            pytest.skip("reportlab not installed")
        
        # Create PDF with multiple blank pages
        pdf_path = tmp_path / "blank_pages.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Add 3 blank pages
        c.showPage()
        c.showPage()
        c.showPage()
        c.save()
        
        # Configure conversion
        config = ConversionConfig(
            input_path=str(pdf_path),
            output_path=None,
            preview_mode=False,
            dry_run=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config, logger)
        
        # Convert the PDF
        result = orchestrator.convert(str(pdf_path))
        
        # Should succeed
        assert result.success is True, "PDF with blank pages should succeed"
        assert result.markdown_content is not None
        
        # Output should be minimal
        assert len(result.markdown_content.strip()) < 200, \
            "Blank pages PDF should produce minimal output"
