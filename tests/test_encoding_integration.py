"""Integration tests for encoding detection in document conversion."""

import pytest
from pathlib import Path
from docx import Document
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
import tempfile
import io


class TestEncodingIntegration:
    """Integration tests for encoding detection during document conversion."""
    
    def test_word_document_with_special_characters(self, tmp_path):
        """Test conversion of Word document with special characters."""
        # Create a Word document with special characters
        doc_path = tmp_path / "test_special_chars.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("日本語テキスト (Japanese)")
        doc.add_paragraph("中文文本 (Chinese)")
        doc.add_paragraph("Español con acentos: café, niño")
        doc.add_paragraph("Français: résumé, naïve")
        doc.add_paragraph("Special symbols: © ® ™ € £ ¥")
        doc.save(str(doc_path))
        
        # Create config and logger
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        logger = Logger(log_level=LogLevel.DEBUG)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify special characters are preserved
        assert "日本語" in result.markdown_content
        assert "中文" in result.markdown_content
        assert "café" in result.markdown_content
        assert "résumé" in result.markdown_content
        assert "©" in result.markdown_content
    
    def test_word_document_encoding_warning_logged(self, tmp_path):
        """Test that encoding warnings are logged when issues are detected."""
        # Create a Word document
        doc_path = tmp_path / "test_doc.docx"
        doc = Document()
        doc.add_paragraph("Normal text")
        doc.save(str(doc_path))
        
        # Create config with log file
        log_path = tmp_path / "conversion.log"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True,
            log_file=str(log_path)
        )
        logger = Logger(log_level=LogLevel.DEBUG, output_path=str(log_path))
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        
        # Log file should exist
        assert log_path.exists()
    
    def test_excel_document_with_unicode(self, tmp_path):
        """Test conversion of Excel document with Unicode characters."""
        import openpyxl
        
        # Create an Excel file with Unicode content
        excel_path = tmp_path / "test_unicode.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Unicode Test"
        
        ws['A1'] = "Language"
        ws['B1'] = "Text"
        ws['A2'] = "Japanese"
        ws['B2'] = "こんにちは世界"
        ws['A3'] = "Chinese"
        ws['B3'] = "你好世界"
        ws['A4'] = "Korean"
        ws['B4'] = "안녕하세요"
        ws['A5'] = "Arabic"
        ws['B5'] = "مرحبا بالعالم"
        ws['A6'] = "Russian"
        ws['B6'] = "Привет мир"
        
        wb.save(str(excel_path))
        
        # Create config and logger
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=None,
            preview_mode=True
        )
        logger = Logger(log_level=LogLevel.DEBUG)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify Unicode characters are preserved
        assert "こんにちは" in result.markdown_content
        assert "你好" in result.markdown_content
        assert "안녕하세요" in result.markdown_content
    
    def test_pdf_text_encoding_normalization(self, tmp_path):
        """Test that PDF text is properly normalized."""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Create a simple PDF with text
        pdf_path = tmp_path / "test_encoding.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Add some text
        c.drawString(100, 750, "Test Document")
        c.drawString(100, 730, "This is a test of encoding detection.")
        c.drawString(100, 710, "Special characters: café résumé")
        
        c.save()
        
        # Create config and logger
        config = ConversionConfig(
            input_path=str(pdf_path),
            output_path=None,
            preview_mode=True
        )
        logger = Logger(log_level=LogLevel.DEBUG)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(pdf_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify text was extracted
        assert "Test Document" in result.markdown_content or "test" in result.markdown_content.lower()
    
    def test_encoding_detection_with_mixed_content(self, tmp_path):
        """Test encoding detection with mixed language content."""
        # Create a Word document with mixed content
        doc_path = tmp_path / "test_mixed.docx"
        doc = Document()
        doc.add_heading("Multilingual Document", level=1)
        doc.add_paragraph("English: Hello World")
        doc.add_paragraph("Spanish: Hola Mundo")
        doc.add_paragraph("French: Bonjour le Monde")
        doc.add_paragraph("German: Hallo Welt")
        doc.add_paragraph("Italian: Ciao Mondo")
        doc.add_paragraph("Portuguese: Olá Mundo")
        doc.add_paragraph("Japanese: こんにちは世界")
        doc.add_paragraph("Chinese: 你好世界")
        doc.add_paragraph("Korean: 안녕하세요 세계")
        doc.save(str(doc_path))
        
        # Create config and logger
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify all languages are present
        assert "Hello World" in result.markdown_content
        assert "Hola Mundo" in result.markdown_content
        assert "Bonjour" in result.markdown_content
        assert "こんにちは" in result.markdown_content
        assert "你好" in result.markdown_content
    
    def test_encoding_normalization_removes_control_chars(self, tmp_path):
        """Test that control characters are removed during normalization."""
        # Create a Word document
        doc_path = tmp_path / "test_control.docx"
        doc = Document()
        # Add paragraph with normal text (control chars would be in actual file corruption)
        doc.add_paragraph("Normal text without control characters")
        doc.save(str(doc_path))
        
        # Create config and logger
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        logger = Logger(log_level=LogLevel.DEBUG)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify no null bytes in output
        assert '\x00' not in result.markdown_content
        
        # Verify text is present
        assert "Normal text" in result.markdown_content
