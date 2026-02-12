"""
Tests for WordParser implementation.

This module contains unit tests for the WordParser class to verify
Word document parsing functionality.
"""

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

from src.parsers import WordParser
from src.internal_representation import (
    InternalDocument, Paragraph, Heading, Table, DocumentList,
    ImageReference, Link, TextFormatting
)


class TestWordParserBasic:
    """Test basic WordParser functionality."""
    
    def test_parse_simple_document(self, tmp_path):
        """Test parsing a simple Word document with text."""
        # Create a simple Word document
        doc = Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another paragraph.")
        
        # Save to temporary file
        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify result
        assert isinstance(result, InternalDocument)
        assert len(result.sections) > 0
        assert result.metadata.source_format == "docx"
    
    def test_parse_document_with_headings(self, tmp_path):
        """Test parsing a document with headings."""
        # Create document with headings
        doc = Document()
        doc.add_heading("Heading 1", level=1)
        doc.add_paragraph("Content under heading 1")
        doc.add_heading("Heading 2", level=2)
        doc.add_paragraph("Content under heading 2")
        
        # Save to temporary file
        file_path = tmp_path / "test_headings.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify headings were extracted
        assert len(result.sections) >= 2
        assert result.sections[0].heading is not None
        assert result.sections[0].heading.level == 1
        assert "Heading 1" in result.sections[0].heading.text
    
    def test_parse_document_with_table(self, tmp_path):
        """Test parsing a document with a table."""
        # Create document with table
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Column 1"
        header_cells[1].text = "Column 2"
        header_cells[2].text = "Column 3"
        
        # Add data
        for i in range(1, 3):
            row_cells = table.rows[i].cells
            row_cells[0].text = f"Row {i} Col 1"
            row_cells[1].text = f"Row {i} Col 2"
            row_cells[2].text = f"Row {i} Col 3"
        
        # Save to temporary file
        file_path = tmp_path / "test_table.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify table was extracted
        assert len(result.sections) > 0
        # Find table in content
        tables = [item for item in result.sections[0].content if isinstance(item, Table)]
        assert len(tables) > 0
        assert len(tables[0].headers) == 3
        assert len(tables[0].rows) == 2
    
    def test_parse_document_with_formatting(self, tmp_path):
        """Test parsing a document with text formatting."""
        # Create document with formatted text
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True
        
        para2 = doc.add_paragraph()
        run2 = para2.add_run("Italic text")
        run2.italic = True
        
        # Save to temporary file
        file_path = tmp_path / "test_formatting.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify formatting was detected
        assert len(result.sections) > 0
        paragraphs = [item for item in result.sections[0].content if isinstance(item, Paragraph)]
        assert len(paragraphs) >= 2
        
        # Check that at least one paragraph has bold or italic formatting
        has_formatting = any(
            p.formatting in [TextFormatting.BOLD, TextFormatting.ITALIC, TextFormatting.BOLD_ITALIC]
            for p in paragraphs
        )
        assert has_formatting
    
    def test_parse_empty_document(self, tmp_path):
        """Test parsing an empty Word document."""
        # Create empty document
        doc = Document()
        
        # Save to temporary file
        file_path = tmp_path / "test_empty.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify result
        assert isinstance(result, InternalDocument)
        assert result.metadata.source_format == "docx"
    
    def test_parse_nonexistent_file(self):
        """Test parsing a nonexistent file raises an error."""
        parser = WordParser()
        
        with pytest.raises(Exception):
            parser.parse("/nonexistent/file.docx")


class TestWordParserMetadata:
    """Test metadata extraction."""
    
    def test_extract_metadata(self, tmp_path):
        """Test that metadata is extracted correctly."""
        # Create document with metadata
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Content")
        
        # Save to temporary file
        file_path = tmp_path / "test_metadata.docx"
        doc.save(str(file_path))
        
        # Parse the document
        parser = WordParser()
        result = parser.parse(str(file_path))
        
        # Verify metadata
        assert result.metadata.title == "Test Document"
        assert result.metadata.author == "Test Author"
        assert result.metadata.source_format == "docx"



class TestWordParserProperties:
    """Property-based tests for WordParser using Hypothesis."""
    
    def test_property_1_text_content_extraction(self, tmp_path):
        """Feature: document-to-markdown-converter, Property 1: Text content extraction
        
        For any valid .docx file with text content, the converter should produce
        Markdown output containing all the text from the source document.
        
        **Validates: Requirements 1.1**
        """
        from hypothesis import given, strategies as st
        from docx import Document
        import string
        
        # Generate XML-compatible text (printable characters only)
        xml_safe_chars = string.printable.replace('\x0b', '').replace('\x0c', '')
        
        @given(st.text(alphabet=xml_safe_chars, min_size=1, max_size=1000))
        def property_test(text_content):
            # Skip text that is only whitespace
            if not text_content.strip():
                return
            
            # Create a Word document with the text content
            doc = Document()
            
            # Split text by line breaks and add each as a separate paragraph
            # This matches how Word actually handles text with line breaks
            lines = text_content.replace('\r\n', '\n').replace('\r', '\n').split('\n')
            for line in lines:
                doc.add_paragraph(line)
            
            # Save to temporary file
            file_path = tmp_path / f"test_property_{abs(hash(text_content))}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Verify the text content is present in the internal representation
            assert isinstance(result, InternalDocument)
            
            # Extract all text from sections
            extracted_text = ""
            for section in result.sections:
                for i, content_item in enumerate(section.content):
                    if isinstance(content_item, Paragraph):
                        extracted_text += content_item.text
                        # Add newline between paragraphs (except for the last one)
                        if i < len(section.content) - 1:
                            extracted_text += "\n"
            
            # Normalize line endings and whitespace for comparison
            normalized_input = text_content.replace('\r\n', '\n').replace('\r', '\n').strip()
            normalized_extracted = extracted_text.replace('\r\n', '\n').replace('\r', '\n').strip()
            
            # Verify that the essential text content is preserved
            # We check that all non-empty lines from input appear in output
            input_lines = [line.strip() for line in normalized_input.split('\n') if line.strip()]
            extracted_lines = [line.strip() for line in normalized_extracted.split('\n') if line.strip()]
            
            for input_line in input_lines:
                assert any(input_line in extracted_line for extracted_line in extracted_lines), \
                    f"Expected line '{input_line}' not found in extracted text"
            
            # Clean up
            if file_path.exists():
                file_path.unlink()
        
        # Run the property test
        property_test()
    
    def test_property_5_text_formatting_preservation(self, tmp_path):
        """Feature: document-to-markdown-converter, Property 5: Text formatting preservation
        
        For any Word document with bold or italic text, the converter should preserve
        the formatting using Markdown syntax (**bold**, *italic*).
        
        **Validates: Requirements 1.5**
        """
        from hypothesis import given, strategies as st
        from docx import Document
        import string
        
        # Generate XML-compatible text (printable characters only)
        xml_safe_chars = string.printable.replace('\x0b', '').replace('\x0c', '')
        
        @given(
            st.text(alphabet=xml_safe_chars, min_size=1, max_size=500),
            st.sampled_from(['bold', 'italic', 'both', 'none'])
        )
        def property_test(text_content, formatting_type):
            # Skip text that is only whitespace
            if not text_content.strip():
                return
            
            # Create a Word document with formatted text
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(text_content)
            
            # Apply formatting based on type
            if formatting_type == 'bold':
                run.bold = True
            elif formatting_type == 'italic':
                run.italic = True
            elif formatting_type == 'both':
                run.bold = True
                run.italic = True
            # 'none' means no formatting
            
            # Save to temporary file
            file_path = tmp_path / f"test_formatting_{abs(hash(text_content + formatting_type))}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Verify the formatting is preserved in the internal representation
            assert isinstance(result, InternalDocument)
            assert len(result.sections) > 0
            
            # Find paragraphs in the content
            paragraphs = []
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, Paragraph):
                        paragraphs.append(content_item)
            
            # Verify at least one paragraph was extracted
            assert len(paragraphs) > 0, "No paragraphs found in parsed document"
            
            # Verify formatting matches what was applied
            found_paragraph = paragraphs[0]
            
            if formatting_type == 'bold':
                assert found_paragraph.formatting in [TextFormatting.BOLD, TextFormatting.BOLD_ITALIC], \
                    f"Expected BOLD formatting, got {found_paragraph.formatting}"
            elif formatting_type == 'italic':
                assert found_paragraph.formatting in [TextFormatting.ITALIC, TextFormatting.BOLD_ITALIC], \
                    f"Expected ITALIC formatting, got {found_paragraph.formatting}"
            elif formatting_type == 'both':
                assert found_paragraph.formatting == TextFormatting.BOLD_ITALIC, \
                    f"Expected BOLD_ITALIC formatting, got {found_paragraph.formatting}"
            elif formatting_type == 'none':
                assert found_paragraph.formatting == TextFormatting.NORMAL, \
                    f"Expected NORMAL formatting, got {found_paragraph.formatting}"
            
            # Clean up
            if file_path.exists():
                file_path.unlink()
        
        # Run the property test
        property_test()
    
    def test_property_6_image_extraction_and_referencing(self, tmp_path):
        """Feature: document-to-markdown-converter, Property 6: Image extraction and referencing
        
        For any Word document containing images, the converter should extract the images
        and include proper Markdown image references in the output.
        
        **Validates: Requirements 1.6**
        """
        from hypothesis import given, strategies as st
        from docx import Document
        from docx.shared import Inches
        from PIL import Image as PILImage
        import io
        
        @given(st.integers(min_value=1, max_value=5))
        def property_test(num_images):
            # Create a Word document with images
            doc = Document()
            doc.add_paragraph("Document with images")
            
            # Create simple test images and add them to the document
            # Use different colors to ensure Word doesn't deduplicate them
            colors = [(255, 0, 0), (0, 255, 0), (0, 0, 255), (255, 255, 0), (255, 0, 255)]
            for i in range(num_images):
                # Create a simple colored image with unique color
                img = PILImage.new('RGB', (100, 100), color=colors[i])
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                
                # Add image to document
                doc.add_picture(img_bytes, width=Inches(1.0))
            
            # Save to temporary file
            file_path = tmp_path / f"test_images_{num_images}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Verify images were extracted
            assert isinstance(result, InternalDocument)
            assert len(result.images) == num_images, \
                f"Expected {num_images} images, found {len(result.images)}"
            
            # Verify each image has required attributes
            for img_ref in result.images:
                assert isinstance(img_ref, ImageReference)
                assert img_ref.source_path is not None, "Image source_path should not be None"
                assert img_ref.alt_text is not None, "Image alt_text should not be None"
            
            # Verify image references are in the content
            image_refs_in_content = []
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, ImageReference):
                        image_refs_in_content.append(content_item)
            
            assert len(image_refs_in_content) == num_images, \
                f"Expected {num_images} image references in content, found {len(image_refs_in_content)}"
            
            # Clean up
            if file_path.exists():
                file_path.unlink()
        
        # Run the property test
        property_test()
    
    def test_property_7_hyperlink_preservation(self, tmp_path):
        """Feature: document-to-markdown-converter, Property 7: Hyperlink preservation
        
        For any Word document containing hyperlinks, the converter should preserve them
        using Markdown link syntax [text](url).
        
        **Validates: Requirements 1.7**
        """
        from hypothesis import given, strategies as st
        from docx import Document
        from docx.oxml.shared import OxmlElement, qn
        import string
        
        # Generate XML-compatible text (printable characters only)
        xml_safe_chars = string.printable.replace('\x0b', '').replace('\x0c', '')
        
        def add_hyperlink(paragraph, text, url):
            """Helper function to add a hyperlink to a paragraph."""
            # This is a workaround since python-docx doesn't have direct hyperlink support
            part = paragraph.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            
            paragraph._p.append(hyperlink)
            
            return hyperlink
        
        @given(
            st.text(alphabet=xml_safe_chars, min_size=1, max_size=100),
            st.text(alphabet='abcdefghijklmnopqrstuvwxyz0123456789-', min_size=5, max_size=50)
        )
        def property_test(link_text, url_path):
            # Skip text that is only whitespace or contains only control characters
            if not link_text.strip():
                return
            
            # Word normalizes whitespace in text, so we should test with normalized input
            # Replace all whitespace characters with spaces and collapse multiple spaces
            import re
            normalized_input = re.sub(r'\s+', ' ', link_text).strip()
            
            # Skip if normalization results in empty string
            if not normalized_input:
                return
            
            # Create a valid URL
            url = f"https://example.com/{url_path}"
            
            # Create a Word document with a hyperlink
            doc = Document()
            para = doc.add_paragraph("This document contains a link: ")
            
            # Add hyperlink using helper function with normalized text
            add_hyperlink(para, normalized_input, url)
            
            # Save to temporary file
            file_path = tmp_path / f"test_link_{abs(hash(link_text + url))}.docx"
            doc.save(str(file_path))
            
            # Parse the document
            parser = WordParser()
            result = parser.parse(str(file_path))
            
            # Verify the hyperlink is preserved in the internal representation
            assert isinstance(result, InternalDocument)
            assert len(result.sections) > 0
            
            # Find links in the content
            links = []
            for section in result.sections:
                for content_item in section.content:
                    if isinstance(content_item, Link):
                        links.append(content_item)
            
            # Verify at least one link was extracted
            assert len(links) > 0, "No links found in parsed document"
            
            # Verify the link has the correct text and URL
            # Word normalizes whitespace, so we compare normalized versions
            found_link = links[0]
            
            # Both should match the normalized input we used
            assert found_link.text == normalized_input, \
                f"Expected link text '{normalized_input}', got '{found_link.text}'"
            assert found_link.url == url, \
                f"Expected link URL '{url}', got '{found_link.url}'"
            
            # Clean up
            if file_path.exists():
                file_path.unlink()
        
        # Run the property test
        property_test()
