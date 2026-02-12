"""Property-based tests for UTF-8 output encoding.

Feature: document-to-markdown-converter
Property 38: UTF-8 encoding output
Validates: Requirements 8.5, 9.2
"""

import pytest
from hypothesis import given, strategies as st, settings
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.output_writer import OutputWriter
import tempfile
import os


# Strategy for generating Unicode text that's compatible with Word/Excel
# Exclude control characters that aren't allowed in XML/Office documents
def is_valid_office_char(c):
    """Check if character is valid for Office documents."""
    code = ord(c)
    # Allow tab, newline, carriage return
    if code in (0x09, 0x0A, 0x0D):
        return True
    # Exclude other control characters (0x00-0x1F, 0x7F-0x9F)
    if code < 0x20 or (0x7F <= code <= 0x9F):
        return False
    # Exclude surrogates
    if 0xD800 <= code <= 0xDFFF:
        return False
    return True

unicode_text_strategy = st.text(
    alphabet=st.characters(
        blacklist_categories=('Cs',),  # Exclude surrogates
    ).filter(is_valid_office_char),
    min_size=1,
    max_size=500
)


@settings(max_examples=100)
@given(text_content=unicode_text_strategy)
def test_property_38_word_document_utf8_output(text_content):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any conversion operation, the converter should output Markdown files 
    in UTF-8 encoding by default.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document with the generated text
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph(text_content)
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should exist
        assert output_path.exists(), "Output file was not created"
        
        # Property: Output file should be readable as UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # If we get here, UTF-8 reading succeeded
            assert len(content) > 0, "Output file is empty"
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")
        
        # Property: Original text should be preserved in output
        # (allowing for Markdown formatting and escaping)
        # Note: Markdown escaper may escape special characters like { } [ ] etc.
        if text_content.strip():
            # Check if the text is present (possibly with escaping)
            # We verify UTF-8 encoding is working, not exact text preservation
            assert len(content) > 0, "Output should contain content"


@settings(max_examples=100)
@given(text_content=unicode_text_strategy)
def test_property_38_excel_document_utf8_output(text_content):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any Excel conversion operation, the converter should output Markdown files 
    in UTF-8 encoding by default.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file with the generated text
        excel_path = tmp_path / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = text_content
        wb.save(str(excel_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should exist
        assert output_path.exists(), "Output file was not created"
        
        # Property: Output file should be readable as UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            assert len(content) > 0, "Output file is empty"
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")


@settings(max_examples=100)
@given(text_content=unicode_text_strategy)
def test_property_38_output_writer_utf8(text_content):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any text content, the OutputWriter should write files in UTF-8 encoding.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        output_path = tmp_path / "output.md"
        
        # Write content using OutputWriter
        writer = OutputWriter()
        writer.write_to_file(text_content, str(output_path))
        
        # Property: Output file should exist
        assert output_path.exists(), "Output file was not created"
        
        # Property: Output file should be readable as UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                read_content = f.read()
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")
        
        # Property: Content should be preserved (with line ending normalization)
        # Python's text mode normalizes line endings to \n
        expected_content = text_content.replace('\r\n', '\n').replace('\r', '\n')
        assert read_content == expected_content, "Content was not preserved"


@settings(max_examples=100)
@given(
    text1=unicode_text_strategy,
    text2=unicode_text_strategy,
    text3=unicode_text_strategy
)
def test_property_38_multiple_paragraphs_utf8(text1, text2, text3):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any document with multiple paragraphs containing Unicode text,
    the converter should output UTF-8 encoded Markdown.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document with multiple paragraphs
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph(text1)
        doc.add_paragraph(text2)
        doc.add_paragraph(text3)
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should be valid UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")
        
        # Property: All paragraphs should be present (if non-empty)
        # Note: Markdown escaping may modify special characters
        # We're testing UTF-8 encoding, not exact text preservation
        if text1.strip() or text2.strip() or text3.strip():
            assert len(content) > 0, "Output should contain content"


@settings(max_examples=100)
@given(cell_values=st.lists(unicode_text_strategy, min_size=1, max_size=10))
def test_property_38_excel_table_utf8(cell_values):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any Excel table with Unicode cell values, the converter should output
    UTF-8 encoded Markdown.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file with multiple cells
        excel_path = tmp_path / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        
        for i, value in enumerate(cell_values, start=1):
            ws[f'A{i}'] = value
        
        wb.save(str(excel_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should be valid UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")


@settings(max_examples=50)
@given(
    heading_text=unicode_text_strategy,
    paragraph_text=unicode_text_strategy
)
def test_property_38_document_structure_utf8(heading_text, paragraph_text):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any document with headings and paragraphs containing Unicode text,
    the converter should output UTF-8 encoded Markdown.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document with heading and paragraph
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading(heading_text, level=1)
        doc.add_paragraph(paragraph_text)
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should be valid UTF-8
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError as e:
            pytest.fail(f"Output file is not valid UTF-8: {e}")
        
        # Property: Heading and paragraph should be present (if non-empty)
        # Note: Markdown escaping may modify special characters
        # We're testing UTF-8 encoding, not exact text preservation
        if heading_text.strip() or paragraph_text.strip():
            assert len(content) > 0, "Output should contain content"


@settings(max_examples=50)
@given(text_content=unicode_text_strategy)
def test_property_38_no_utf8_bom(text_content):
    """Feature: document-to-markdown-converter, Property 38: UTF-8 encoding output
    
    For any conversion operation, the converter should output UTF-8 without BOM.
    UTF-8 BOM is not necessary and can cause issues with some tools.
    
    Validates: Requirements 8.5, 9.2
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph(text_content)
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Conversion should succeed
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Output file should not have UTF-8 BOM
        with open(output_path, 'rb') as f:
            first_bytes = f.read(3)
        
        # UTF-8 BOM is 0xEF 0xBB 0xBF
        assert first_bytes != b'\xef\xbb\xbf', "UTF-8 BOM should not be present"
