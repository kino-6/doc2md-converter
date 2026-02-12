"""Integration tests for CLI and ConversionOrchestrator."""

import pytest
from pathlib import Path
from docx import Document
from click.testing import CliRunner

from src.cli import main
from src.conversion_orchestrator import ConversionOrchestrator, ConversionConfig
from src.logger import Logger, LogLevel


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample Word document for testing."""
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("Content in section 1.")
    doc.save(str(doc_path))
    return doc_path


def test_cli_help():
    """Test CLI help command."""
    runner = CliRunner()
    result = runner.invoke(main, ['--help'])
    
    assert result.exit_code == 0
    assert 'Convert Word, Excel, and PDF documents to Markdown format' in result.output
    assert '--input' in result.output
    assert '--output' in result.output


def test_cli_version():
    """Test CLI version command."""
    runner = CliRunner()
    result = runner.invoke(main, ['--version'])
    
    assert result.exit_code == 0
    assert '0.1.0' in result.output


def test_cli_missing_input():
    """Test CLI with missing required input argument."""
    runner = CliRunner()
    result = runner.invoke(main, [])
    
    assert result.exit_code != 0
    assert 'Missing option' in result.output or 'required' in result.output.lower()


def test_cli_invalid_file():
    """Test CLI with non-existent file."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', 'nonexistent.docx'])
    
    # Click should catch the file existence check
    assert result.exit_code != 0


def test_cli_preview_mode(sample_docx):
    """Test CLI preview mode."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--preview'])
    
    # Should succeed and display content
    assert result.exit_code == 0
    assert 'Test Document' in result.output or 'Preview' in result.output


def test_cli_dry_run_mode(sample_docx):
    """Test CLI dry-run mode."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--dry-run'])
    
    # Should succeed without creating output file
    assert result.exit_code == 0


def test_cli_output_to_file(sample_docx, tmp_path):
    """Test CLI with output file."""
    output_path = tmp_path / "output.md"
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '-o', str(output_path)])
    
    assert result.exit_code == 0
    assert output_path.exists()
    
    # Verify content
    content = output_path.read_text()
    assert '# Test Document' in content
    assert 'This is a test paragraph' in content


def test_orchestrator_basic_conversion(sample_docx, tmp_path):
    """Test ConversionOrchestrator basic conversion."""
    output_path = tmp_path / "output.md"
    
    config = ConversionConfig(
        input_path=str(sample_docx),
        output_path=str(output_path)
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    assert result.markdown_content is not None
    assert '# Test Document' in result.markdown_content
    assert result.duration > 0
    assert len(result.errors) == 0
    assert output_path.exists()


def test_orchestrator_preview_mode(sample_docx):
    """Test ConversionOrchestrator preview mode."""
    config = ConversionConfig(
        input_path=str(sample_docx),
        preview_mode=True
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    assert result.markdown_content is not None
    # In preview mode, no output file should be created
    assert result.output_path is None


def test_orchestrator_dry_run_mode(sample_docx, tmp_path):
    """Test ConversionOrchestrator dry-run mode."""
    output_path = tmp_path / "output.md"
    
    config = ConversionConfig(
        input_path=str(sample_docx),
        output_path=str(output_path),
        dry_run=True
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    assert result.markdown_content is not None
    # In dry-run mode, output file should not be created
    assert not output_path.exists()


def test_orchestrator_invalid_file():
    """Test ConversionOrchestrator with invalid file."""
    config = ConversionConfig(
        input_path="nonexistent.docx"
    )
    logger = Logger(log_level=LogLevel.ERROR)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert("nonexistent.docx")
    
    assert not result.success
    assert len(result.errors) > 0
    assert 'not found' in result.errors[0].lower()


def test_orchestrator_unsupported_format(tmp_path):
    """Test ConversionOrchestrator with unsupported file format."""
    # Create a text file with wrong extension
    invalid_file = tmp_path / "test.txt"
    invalid_file.write_text("test content")
    
    config = ConversionConfig(
        input_path=str(invalid_file)
    )
    logger = Logger(log_level=LogLevel.ERROR)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(invalid_file))
    
    assert not result.success
    assert len(result.errors) > 0
    assert 'format' in result.errors[0].lower()


def test_orchestrator_statistics(sample_docx):
    """Test ConversionOrchestrator statistics collection."""
    config = ConversionConfig(
        input_path=str(sample_docx)
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    assert result.stats.headings_detected >= 2  # We added 2 headings
    assert result.stats.tables_converted >= 0


def test_orchestrator_with_heading_offset(sample_docx, tmp_path):
    """Test ConversionOrchestrator with heading offset."""
    output_path = tmp_path / "output.md"
    
    config = ConversionConfig(
        input_path=str(sample_docx),
        output_path=str(output_path),
        heading_offset=1
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    # With offset of 1, H1 should become H2
    assert '## Test Document' in result.markdown_content


def test_orchestrator_with_metadata(sample_docx, tmp_path):
    """Test ConversionOrchestrator with metadata inclusion."""
    output_path = tmp_path / "output.md"
    
    config = ConversionConfig(
        input_path=str(sample_docx),
        output_path=str(output_path),
        include_metadata=True
    )
    logger = Logger(log_level=LogLevel.INFO)
    
    orchestrator = ConversionOrchestrator(config=config, logger=logger)
    result = orchestrator.convert(str(sample_docx))
    
    assert result.success
    # Metadata should be included in output
    assert result.markdown_content is not None


# Additional CLI argument parsing tests

def test_cli_log_level_parsing(sample_docx):
    """Test CLI log level argument parsing."""
    runner = CliRunner()
    
    # Test each log level
    for level in ['DEBUG', 'INFO', 'WARNING', 'ERROR']:
        result = runner.invoke(main, ['-i', str(sample_docx), '--log-level', level, '--preview'])
        assert result.exit_code == 0


def test_cli_ocr_language_parsing(sample_docx):
    """Test CLI OCR language argument parsing."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--ocr-lang', 'eng', '--preview'])
    
    assert result.exit_code == 0


def test_cli_heading_offset_parsing(sample_docx, tmp_path):
    """Test CLI heading offset argument parsing."""
    output_path = tmp_path / "output.md"
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '-o', str(output_path), '--heading-offset', '2'])
    
    assert result.exit_code == 0
    assert output_path.exists()
    
    # Verify heading offset was applied
    content = output_path.read_text()
    assert '###' in content  # H1 + offset 2 = H3


def test_cli_max_file_size_parsing(sample_docx):
    """Test CLI max file size argument parsing."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--max-file-size', '50', '--preview'])
    
    assert result.exit_code == 0


def test_cli_include_metadata_flag(sample_docx, tmp_path):
    """Test CLI include metadata flag parsing."""
    output_path = tmp_path / "output.md"
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '-o', str(output_path), '--include-metadata'])
    
    assert result.exit_code == 0
    assert output_path.exists()


def test_cli_embed_images_base64_flag(sample_docx):
    """Test CLI embed images base64 flag parsing."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--embed-images-base64', '--preview'])
    
    assert result.exit_code == 0


def test_cli_no_extract_images_flag(sample_docx):
    """Test CLI no extract images flag parsing."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--no-extract-images', '--preview'])
    
    assert result.exit_code == 0


def test_cli_no_validate_flag(sample_docx):
    """Test CLI no validate flag parsing."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--no-validate', '--preview'])
    
    assert result.exit_code == 0


def test_cli_short_options(sample_docx, tmp_path):
    """Test CLI short option aliases."""
    output_path = tmp_path / "output.md"
    runner = CliRunner()
    
    # Test -i and -o short options
    result = runner.invoke(main, ['-i', str(sample_docx), '-o', str(output_path)])
    assert result.exit_code == 0
    assert output_path.exists()


def test_cli_preview_short_option(sample_docx):
    """Test CLI preview short option (-p)."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '-p'])
    
    assert result.exit_code == 0


def test_cli_multiple_input_files(tmp_path):
    """Test CLI with multiple input files (batch mode)."""
    # Create multiple sample documents
    doc1_path = tmp_path / "doc1.docx"
    doc2_path = tmp_path / "doc2.docx"
    
    for doc_path in [doc1_path, doc2_path]:
        doc = Document()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("Content")
        doc.save(str(doc_path))
    
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(doc1_path), '-i', str(doc2_path)])
    
    assert result.exit_code == 0
    assert 'Batch conversion' in result.output


def test_cli_log_file_option(sample_docx, tmp_path):
    """Test CLI log file option."""
    log_path = tmp_path / "conversion.log"
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--log-file', str(log_path), '--preview'])
    
    assert result.exit_code == 0
    assert log_path.exists()


def test_cli_config_file_option(sample_docx, tmp_path):
    """Test CLI config file option."""
    # Create a simple config file
    config_path = tmp_path / "config.yaml"
    config_path.write_text("heading_offset: 1\ninclude_metadata: true\n")
    
    output_path = tmp_path / "output.md"
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '-o', str(output_path), '-c', str(config_path)])
    
    assert result.exit_code == 0
    assert output_path.exists()


def test_cli_help_content():
    """Test CLI help message contains all expected options."""
    runner = CliRunner()
    result = runner.invoke(main, ['--help'])
    
    assert result.exit_code == 0
    
    # Check for all major options
    expected_options = [
        '--input', '-i',
        '--output', '-o',
        '--config', '-c',
        '--preview', '-p',
        '--dry-run',
        '--log-level',
        '--log-file',
        '--extract-images',
        '--no-extract-images',
        '--embed-images-base64',
        '--ocr-lang',
        '--heading-offset',
        '--include-metadata',
        '--validate',
        '--no-validate',
        '--max-file-size',
        '--version',
        '--help'
    ]
    
    for option in expected_options:
        assert option in result.output, f"Option {option} not found in help output"


def test_cli_help_examples():
    """Test CLI help message contains usage examples."""
    runner = CliRunner()
    result = runner.invoke(main, ['--help'])
    
    assert result.exit_code == 0
    assert 'Examples:' in result.output or 'example' in result.output.lower()


def test_cli_invalid_log_level(sample_docx):
    """Test CLI with invalid log level."""
    runner = CliRunner()
    result = runner.invoke(main, ['-i', str(sample_docx), '--log-level', 'INVALID'])
    
    assert result.exit_code != 0
    assert 'Invalid value' in result.output or 'invalid choice' in result.output.lower()
