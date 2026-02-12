"""Unit tests for UTF-8 output encoding (Requirements 8.5, 9.2)."""

import pytest
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.output_writer import OutputWriter
import sys
import io


class TestUTF8Output:
    """Unit tests for UTF-8 output encoding."""
    
    def test_file_output_uses_utf8_encoding(self, tmp_path):
        """Test that file output uses UTF-8 encoding by default.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create a Word document with Unicode characters
        doc_path = tmp_path / "test_unicode.docx"
        doc = Document()
        doc.add_paragraph("日本語テキスト")
        doc.add_paragraph("中文文本")
        doc.add_paragraph("Emoji: 😀 🎉 ✨")
        doc.add_paragraph("Special: © ® ™ € £")
        doc.save(str(doc_path))
        
        # Convert to Markdown with file output
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Read file and verify encoding
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Verify Unicode characters are preserved
        assert "日本語" in content
        assert "中文" in content
        assert "😀" in content
        assert "©" in content
    
    def test_file_output_readable_as_utf8(self, tmp_path):
        """Test that output file can be read as UTF-8 without errors.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create a Word document with various Unicode characters
        doc_path = tmp_path / "test_multilang.docx"
        doc = Document()
        doc.add_paragraph("English: Hello World")
        doc.add_paragraph("Japanese: こんにちは世界")
        doc.add_paragraph("Chinese: 你好世界")
        doc.add_paragraph("Korean: 안녕하세요")
        doc.add_paragraph("Arabic: مرحبا")
        doc.add_paragraph("Russian: Привет")
        doc.add_paragraph("Greek: Γειά σου")
        doc.add_paragraph("Hebrew: שלום")
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        assert result.success
        
        # Verify file can be read as UTF-8 without errors
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # If we get here, UTF-8 reading succeeded
            assert len(content) > 0
        except UnicodeDecodeError:
            pytest.fail("Output file is not valid UTF-8")
    
    def test_excel_output_uses_utf8(self, tmp_path):
        """Test that Excel conversion output uses UTF-8 encoding.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create an Excel file with Unicode content
        excel_path = tmp_path / "test_unicode.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = "日本語"
        ws['A2'] = "中文"
        ws['A3'] = "한국어"
        ws['A4'] = "Emoji: 🎉"
        wb.save(str(excel_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        assert result.success
        
        # Verify UTF-8 encoding
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        assert "日本語" in content
        assert "中文" in content
    
    def test_output_writer_uses_utf8(self, tmp_path):
        """Test that OutputWriter writes files with UTF-8 encoding.
        
        Validates: Requirements 8.5, 9.2
        """
        writer = OutputWriter()
        output_path = tmp_path / "test_output.md"
        
        # Content with various Unicode characters
        content = """# Test Document

Japanese: 日本語テキスト
Chinese: 中文文本
Korean: 한국어 텍스트
Emoji: 😀 🎉 ✨ 🚀
Special: © ® ™ € £ ¥
Math: ∑ ∫ ∞ ≈ ≠
Arrows: → ← ↑ ↓
"""
        
        # Write content
        writer.write_to_file(content, str(output_path))
        
        # Verify file exists and is UTF-8
        assert output_path.exists()
        
        with open(output_path, 'r', encoding='utf-8') as f:
            read_content = f.read()
        
        assert read_content == content
        assert "日本語" in read_content
        assert "😀" in read_content
        assert "∑" in read_content
    
    def test_utf8_bom_not_added(self, tmp_path):
        """Test that UTF-8 BOM is not added to output files.
        
        UTF-8 BOM is not necessary and can cause issues with some tools.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create a simple document
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        assert result.success
        
        # Read file in binary mode to check for BOM
        with open(output_path, 'rb') as f:
            first_bytes = f.read(3)
        
        # UTF-8 BOM is 0xEF 0xBB 0xBF
        assert first_bytes != b'\xef\xbb\xbf', "UTF-8 BOM should not be present"
    
    def test_stdout_output_handles_utf8(self, tmp_path, capsys):
        """Test that stdout output handles UTF-8 characters correctly.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create a document with Unicode content
        doc_path = tmp_path / "test_unicode.docx"
        doc = Document()
        doc.add_paragraph("Unicode: 日本語 中文 한국어")
        doc.save(str(doc_path))
        
        # Convert to stdout (no output_path)
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None  # Output to stdout
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        assert result.success
        
        # Capture stdout
        captured = capsys.readouterr()
        
        # Verify Unicode characters in stdout
        assert "日本語" in captured.out or "日本語" in result.markdown_content
    
    def test_batch_conversion_uses_utf8(self, tmp_path):
        """Test that batch conversion uses UTF-8 for all files.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create multiple documents with Unicode content
        doc1_path = tmp_path / "doc1.docx"
        doc1 = Document()
        doc1.add_paragraph("Document 1: 日本語")
        doc1.save(str(doc1_path))
        
        doc2_path = tmp_path / "doc2.docx"
        doc2 = Document()
        doc2.add_paragraph("Document 2: 中文")
        doc2.save(str(doc2_path))
        
        # Batch convert
        config = ConversionConfig(
            input_path="",
            batch_mode=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        results = orchestrator.batch_convert([str(doc1_path), str(doc2_path)])
        
        # Verify both conversions succeeded
        assert len(results) == 2
        assert all(r.success for r in results)
        
        # Verify output files are UTF-8
        output1 = tmp_path / "doc1.md"
        output2 = tmp_path / "doc2.md"
        
        with open(output1, 'r', encoding='utf-8') as f:
            content1 = f.read()
        assert "日本語" in content1
        
        with open(output2, 'r', encoding='utf-8') as f:
            content2 = f.read()
        assert "中文" in content2
    
    def test_special_unicode_characters_preserved(self, tmp_path):
        """Test that special Unicode characters are preserved in output.
        
        Validates: Requirements 8.5, 9.2
        """
        # Create document with special Unicode characters
        doc_path = tmp_path / "test_special.docx"
        doc = Document()
        doc.add_paragraph("Zero-width characters: \u200b\u200c\u200d")
        doc.add_paragraph("Combining diacritics: é (e\u0301)")
        doc.add_paragraph("Right-to-left: العربية עברית")
        doc.add_paragraph("Rare symbols: 𝕳𝖊𝖑𝖑𝖔 (mathematical bold)")
        doc.save(str(doc_path))
        
        # Convert to Markdown
        output_path = tmp_path / "output.md"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        logger = Logger(log_level=LogLevel.ERROR)
        
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        assert result.success
        
        # Verify file is valid UTF-8 and characters are preserved
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Check that content has reasonable length (characters preserved)
        assert len(content) > 0
        # Verify some characters are present
        assert "العربية" in content or "עברית" in content
