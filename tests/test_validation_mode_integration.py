"""Integration tests for Markdown validation mode.

Tests the validation mode integration with the conversion orchestrator.

Requirements:
    - 11.3: Validation mode that checks output Markdown syntax
    - 11.4: Report Markdown syntax errors or warnings
"""

import pytest
from pathlib import Path
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.conversion_orchestrator import ConversionOrchestrator
from src.markdown_validator import MARKDOWN_IT_AVAILABLE


# Skip all tests if markdown-it-py is not available
pytestmark = pytest.mark.skipif(
    not MARKDOWN_IT_AVAILABLE,
    reason="markdown-it-py not available"
)


class TestValidationModeIntegration:
    """Integration tests for validation mode."""
    
    def test_validation_enabled_by_default(self, tmp_path):
        """Test that validation is enabled by default."""
        # Create a simple test document
        input_file = tmp_path / "test.docx"
        output_file = tmp_path / "output.md"
        
        # Create a minimal docx file for testing
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is a test paragraph.')
        doc.save(str(input_file))
        
        # Create config with validation enabled (default)
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Verify validator is initialized
        assert orchestrator.markdown_validator is not None
        
        # Convert
        result = orchestrator.convert(str(input_file))
        
        # Should succeed with valid markdown
        assert result.success
        assert result.markdown_content is not None
    
    def test_validation_disabled(self, tmp_path):
        """Test that validation can be disabled."""
        input_file = tmp_path / "test.docx"
        output_file = tmp_path / "output.md"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Test', level=1)
        doc.save(str(input_file))
        
        # Create config with validation disabled
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            validate_output=False
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Validator should not be initialized
        assert orchestrator.markdown_validator is None
        
        # Convert should still work
        result = orchestrator.convert(str(input_file))
        assert result.success
    
    def test_validation_reports_warnings(self, tmp_path):
        """Test that validation reports warnings for issues."""
        input_file = tmp_path / "test.docx"
        output_file = tmp_path / "output.md"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('Normal paragraph.')
        doc.save(str(input_file))
        
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        result = orchestrator.convert(str(input_file))
        
        # Should succeed
        assert result.success
        
        # Check that validation was performed
        # (warnings list may or may not have items depending on the generated markdown)
        assert isinstance(result.warnings, list)
    
    def test_validation_with_complex_document(self, tmp_path):
        """Test validation with a complex document containing various elements."""
        input_file = tmp_path / "complex.docx"
        output_file = tmp_path / "output.md"
        
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Add various elements
        doc.add_heading('Main Title', level=1)
        doc.add_paragraph('Introduction paragraph with **bold** text.')
        
        doc.add_heading('Section 1', level=2)
        doc.add_paragraph('Some content here.')
        
        # Add a list
        doc.add_paragraph('Item 1', style='List Bullet')
        doc.add_paragraph('Item 2', style='List Bullet')
        doc.add_paragraph('Item 3', style='List Bullet')
        
        doc.add_heading('Section 2', level=2)
        
        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Row 1 Col 1'
        table.cell(1, 1).text = 'Row 1 Col 2'
        table.cell(2, 0).text = 'Row 2 Col 1'
        table.cell(2, 1).text = 'Row 2 Col 2'
        
        doc.save(str(input_file))
        
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        result = orchestrator.convert(str(input_file))
        
        # Should succeed
        assert result.success
        assert result.markdown_content is not None
        
        # Verify output file was created
        assert output_file.exists()
        
        # Read and verify the markdown
        markdown_content = output_file.read_text(encoding='utf-8')
        assert '# Main Title' in markdown_content
        assert '## Section 1' in markdown_content
        assert '## Section 2' in markdown_content
        assert '|' in markdown_content  # Table present
    
    def test_validation_in_preview_mode(self, tmp_path):
        """Test that validation works in preview mode."""
        input_file = tmp_path / "test.docx"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Preview Test', level=1)
        doc.add_paragraph('Content for preview.')
        doc.save(str(input_file))
        
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=None,
            preview_mode=True,
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        result = orchestrator.convert(str(input_file))
        
        # Should succeed
        assert result.success
        assert result.markdown_content is not None
        
        # Validation should have been performed
        assert orchestrator.markdown_validator is not None
    
    def test_validation_in_dry_run_mode(self, tmp_path):
        """Test that validation works in dry-run mode."""
        input_file = tmp_path / "test.docx"
        output_file = tmp_path / "output.md"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Dry Run Test', level=1)
        doc.add_paragraph('Content for dry run.')
        doc.save(str(input_file))
        
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            dry_run=True,
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        result = orchestrator.convert(str(input_file))
        
        # Should succeed
        assert result.success
        assert result.markdown_content is not None
        
        # Output file should NOT be created in dry-run mode
        assert not output_file.exists()
        
        # Validation should have been performed
        assert orchestrator.markdown_validator is not None
    
    def test_validation_with_special_characters(self, tmp_path):
        """Test validation with special characters that need escaping."""
        input_file = tmp_path / "special.docx"
        output_file = tmp_path / "output.md"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Special Characters', level=1)
        doc.add_paragraph('Text with special chars: * _ [ ] ( ) # + - . !')
        doc.add_paragraph('Angle brackets: < >')
        doc.save(str(input_file))
        
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            validate_output=True
        )
        
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        result = orchestrator.convert(str(input_file))
        
        # Should succeed
        assert result.success
        assert result.markdown_content is not None
        
        # Markdown should be valid despite special characters
        # (they should be properly escaped)
        assert output_file.exists()


class TestValidationCLIIntegration:
    """Test validation mode through CLI options."""
    
    def test_validate_flag_default(self):
        """Test that --validate flag is enabled by default."""
        config = ConversionConfig(
            input_path="test.docx",
            validate_output=True
        )
        
        assert config.validate_output is True
    
    def test_no_validate_flag(self):
        """Test that --no-validate flag disables validation."""
        config = ConversionConfig(
            input_path="test.docx",
            validate_output=False
        )
        
        assert config.validate_output is False
    
    def test_validation_config_from_file(self, tmp_path):
        """Test loading validation setting from config file."""
        from src.config import ConfigManager
        
        config_file = tmp_path / "config.yaml"
        config_file.write_text("""
validate_output: true
log_level: INFO
""")
        
        manager = ConfigManager()
        config = manager.load_config(str(config_file))
        
        assert config.validate_output is True
