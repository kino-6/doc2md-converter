"""Property-based tests for dry-run mode functionality.

Feature: document-to-markdown-converter
Property 49: Dry-run mode
Validates: Requirements 11.5
"""

import pytest
from hypothesis import given, strategies as st, settings
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
import tempfile


# Strategy for generating valid file names
filename_strategy = st.text(
    alphabet=st.characters(
        whitelist_categories=('Lu', 'Ll', 'Nd'),
        min_codepoint=ord('a'),
        max_codepoint=ord('z')
    ),
    min_size=1,
    max_size=20
)

# Strategy for generating text content compatible with Office documents
def is_valid_office_char(c):
    """Check if character is valid for Office documents."""
    code = ord(c)
    # Allow tab, newline, carriage return
    if code in (0x09, 0x0A, 0x0D):
        return True
    # Exclude other control characters (0x00-0x1F, 0x7F-0x9F)
    if code < 0x20 or (0x7F <= code <= 0x9F):
        return False
    # Exclude surrogates
    if 0xD800 <= code <= 0xDFFF:
        return False
    return True

text_content_strategy = st.text(
    alphabet=st.characters(
        blacklist_categories=('Cs',),  # Exclude surrogates
    ).filter(is_valid_office_char),
    min_size=1,
    max_size=200
)


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_no_file_write(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode, the converter should perform the 
    conversion without writing any output files.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test Word document
        input_file = tmp_path / f"{filename}.docx"
        output_file = tmp_path / f"{filename}.md"
        
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(input_file))
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Property: Conversion should succeed
        assert result.success is True, \
            f"Conversion should succeed in dry-run mode: {result.errors}"
        
        # Property: Markdown content should be generated
        assert result.markdown_content is not None, \
            "Dry-run should generate markdown content"
        
        # Property: No output file should be written
        assert not output_file.exists(), \
            "Dry-run mode should not write output file"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy,
    file_type=st.sampled_from(['docx', 'xlsx'])
)
def test_property_49_dry_run_processes_all_formats(filename, content, file_type):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any supported file format in dry-run mode, the converter should 
    perform full conversion without writing output files.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test file based on type
        if file_type == 'docx':
            input_file = tmp_path / f"{filename}.docx"
            doc = Document()
            doc.add_paragraph(content)
            doc.save(str(input_file))
        else:  # xlsx
            input_file = tmp_path / f"{filename}.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws['A1'] = content
            wb.save(str(input_file))
        
        output_file = tmp_path / f"{filename}.md"
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Property: Conversion should succeed for all formats
        assert result.success is True, \
            f"Dry-run should succeed for {file_type}: {result.errors}"
        
        # Property: Content should be converted
        assert result.markdown_content is not None, \
            f"Dry-run should generate content for {file_type}"
        
        # Property: No output file should exist
        assert not output_file.exists(), \
            f"Dry-run should not write file for {file_type}"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_generates_complete_markdown(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode, the converter should generate 
    complete markdown content equivalent to normal mode.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test Word document
        input_file = tmp_path / f"{filename}.docx"
        
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph(content)
        doc.save(str(input_file))
        
        # Perform conversion in dry-run mode
        config_dry = ConversionConfig(
            input_path=str(input_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator_dry = ConversionOrchestrator(config=config_dry, logger=logger)
        result_dry = orchestrator_dry.convert(str(input_file))
        
        # Perform conversion in normal mode (to stdout)
        config_normal = ConversionConfig(
            input_path=str(input_file),
            output_path=None,  # stdout
            dry_run=False
        )
        orchestrator_normal = ConversionOrchestrator(config=config_normal, logger=logger)
        result_normal = orchestrator_normal.convert(str(input_file))
        
        # Property: Both should succeed
        assert result_dry.success is True, \
            f"Dry-run conversion should succeed: {result_dry.errors}"
        assert result_normal.success is True, \
            f"Normal conversion should succeed: {result_normal.errors}"
        
        # Property: Markdown content should be identical
        assert result_dry.markdown_content == result_normal.markdown_content, \
            "Dry-run should generate same content as normal mode"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=8),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=8)
)
def test_property_49_dry_run_batch_no_files_written(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any batch conversion in dry-run mode, the converter should process 
    all files without writing any output files.
    
    **Validates: Requirements 11.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents
        file_paths = []
        expected_outputs = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
            expected_outputs.append(tmp_path / f"{base_filename}_{i}.md")
        
        # Create orchestrator with batch mode and dry-run enabled
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: All conversions should succeed
        assert len(results) == num_files, \
            f"Expected {num_files} results, got {len(results)}"
        assert all(r.success for r in results), \
            f"All dry-run conversions should succeed: {[r.errors for r in results if not r.success]}"
        
        # Property: All results should have markdown content
        assert all(r.markdown_content is not None for r in results), \
            "All dry-run results should have markdown content"
        
        # Property: No output files should be written
        for output_path in expected_outputs:
            assert not output_path.exists(), \
                f"Dry-run batch should not write file: {output_path}"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_collects_statistics(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode, the converter should collect 
    conversion statistics just like normal mode.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test Word document with various elements
        input_file = tmp_path / f"{filename}.docx"
        
        doc = Document()
        doc.add_heading("Heading 1", level=1)
        doc.add_heading("Heading 2", level=2)
        doc.add_paragraph(content)
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        
        doc.save(str(input_file))
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Property: Conversion should succeed
        assert result.success is True, \
            f"Dry-run conversion should succeed: {result.errors}"
        
        # Property: Statistics should be collected
        assert result.stats is not None, \
            "Dry-run should collect statistics"
        
        # Property: Headings should be detected
        assert result.stats.headings_detected >= 2, \
            f"Should detect at least 2 headings, got {result.stats.headings_detected}"
        
        # Property: Tables should be counted
        assert result.stats.tables_converted >= 1, \
            f"Should detect at least 1 table, got {result.stats.tables_converted}"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_with_validation(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode with validation enabled, the converter 
    should perform validation without writing output files.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test Word document
        input_file = tmp_path / f"{filename}.docx"
        output_file = tmp_path / f"{filename}.md"
        
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(input_file))
        
        # Create orchestrator with dry-run and validation enabled
        config = ConversionConfig(
            input_path=str(input_file),
            output_path=str(output_file),
            dry_run=True,
            validate_output=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Property: Conversion should succeed
        assert result.success is True, \
            f"Dry-run with validation should succeed: {result.errors}"
        
        # Property: Markdown content should be generated
        assert result.markdown_content is not None, \
            "Dry-run with validation should generate content"
        
        # Property: No output file should be written
        assert not output_file.exists(), \
            "Dry-run with validation should not write file"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_preserves_errors(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode that encounters errors, the converter 
    should report errors without writing output files.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create invalid file (wrong format)
        invalid_file = tmp_path / f"{filename}.txt"
        invalid_file.write_text(content)
        
        output_file = tmp_path / f"{filename}.md"
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(invalid_file),
            output_path=str(output_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(invalid_file))
        
        # Property: Conversion should fail
        assert result.success is False, \
            "Dry-run should fail for invalid files"
        
        # Property: Errors should be reported
        assert len(result.errors) > 0, \
            "Dry-run should report errors for invalid files"
        
        # Property: No output file should be written
        assert not output_file.exists(), \
            "Dry-run should not write file even on error"


@settings(max_examples=100)
@given(
    filename=filename_strategy,
    content=text_content_strategy
)
def test_property_49_dry_run_duration_tracking(filename, content):
    """Feature: document-to-markdown-converter, Property 49: Dry-run mode
    
    For any conversion in dry-run mode, the converter should track 
    conversion duration just like normal mode.
    
    **Validates: Requirements 11.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create test Word document
        input_file = tmp_path / f"{filename}.docx"
        
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(input_file))
        
        # Create orchestrator with dry-run enabled
        config = ConversionConfig(
            input_path=str(input_file),
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform conversion
        result = orchestrator.convert(str(input_file))
        
        # Property: Conversion should succeed
        assert result.success is True, \
            f"Dry-run conversion should succeed: {result.errors}"
        
        # Property: Duration should be tracked
        assert result.duration >= 0, \
            "Dry-run should track conversion duration"
        
        # Property: Duration should be reasonable (not zero for actual work)
        # Note: Very fast conversions might be close to zero, so we just check >= 0
        assert isinstance(result.duration, float), \
            "Duration should be a float value"
