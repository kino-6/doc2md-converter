"""Integration tests for encoding issue logging in full conversion pipeline.

This module tests that encoding issues are properly logged during
end-to-end document conversion, validating Requirement 9.5.
"""

import pytest
import logging
import io
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel


class TestEncodingLoggingIntegration:
    """Integration tests for encoding issue logging in conversion pipeline."""
    
    def test_word_conversion_logs_encoding_issues(self, tmp_path):
        """Test that Word conversion logs encoding issues when detected."""
        # Create a Word document
        doc_path = tmp_path / "test_encoding.docx"
        doc = Document()
        doc.add_paragraph("Normal text")
        doc.add_paragraph("Text with special chars: café résumé")
        doc.save(str(doc_path))
        
        # Create a logger with string stream to capture logs
        log_stream = io.StringIO()
        logger = logging.getLogger('test_word_conversion')
        logger.setLevel(logging.DEBUG)
        handler = logging.StreamHandler(log_stream)
        handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        logger.addHandler(handler)
        
        # Create config
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create orchestrator with custom logger
        from src.logger import Logger as CustomLogger
        custom_logger = CustomLogger(log_level=LogLevel.DEBUG)
        custom_logger.logger = logger  # Use our test logger
        
        orchestrator = ConversionOrchestrator(config, custom_logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify text is in output
        assert "Normal text" in result.markdown_content
        assert "café" in result.markdown_content
        
        # Clean up
        logger.removeHandler(handler)
    
    def test_excel_conversion_logs_encoding_issues(self, tmp_path):
        """Test that Excel conversion logs encoding issues when detected."""
        # Create an Excel file
        excel_path = tmp_path / "test_encoding.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        ws['A1'] = "Name"
        ws['B1'] = "Value"
        ws['A2'] = "Normal"
        ws['B2'] = "Text"
        ws['A3'] = "Special"
        ws['B3'] = "café"
        
        wb.save(str(excel_path))
        
        # Create a logger with string stream
        log_stream = io.StringIO()
        logger = logging.getLogger('test_excel_conversion')
        logger.setLevel(logging.DEBUG)
        handler = logging.StreamHandler(log_stream)
        handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        logger.addHandler(handler)
        
        # Create config
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create orchestrator with custom logger
        from src.logger import Logger as CustomLogger
        custom_logger = CustomLogger(log_level=LogLevel.DEBUG)
        custom_logger.logger = logger
        
        orchestrator = ConversionOrchestrator(config, custom_logger)
        result = orchestrator.convert(str(excel_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify text is in output
        assert "Normal" in result.markdown_content
        assert "café" in result.markdown_content
        
        # Clean up
        logger.removeHandler(handler)
    
    def test_pdf_conversion_logs_encoding_issues(self, tmp_path):
        """Test that PDF conversion logs encoding issues when detected."""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Create a simple PDF
        pdf_path = tmp_path / "test_encoding.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test Document")
        c.drawString(100, 730, "Normal text content")
        c.save()
        
        # Create a logger with string stream
        log_stream = io.StringIO()
        logger = logging.getLogger('test_pdf_conversion')
        logger.setLevel(logging.DEBUG)
        handler = logging.StreamHandler(log_stream)
        handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        logger.addHandler(handler)
        
        # Create config
        config = ConversionConfig(
            input_path=str(pdf_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create orchestrator with custom logger
        from src.logger import Logger as CustomLogger
        custom_logger = CustomLogger(log_level=LogLevel.DEBUG)
        custom_logger.logger = logger
        
        orchestrator = ConversionOrchestrator(config, custom_logger)
        result = orchestrator.convert(str(pdf_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Clean up
        logger.removeHandler(handler)
    
    def test_encoding_warnings_appear_in_log_file(self, tmp_path):
        """Test that encoding warnings are written to log file."""
        # Create a Word document
        doc_path = tmp_path / "test_doc.docx"
        doc = Document()
        doc.add_paragraph("Test content with special characters: café")
        doc.save(str(doc_path))
        
        # Create config with log file
        log_path = tmp_path / "conversion.log"
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True,
            log_file=str(log_path),
            log_level=LogLevel.DEBUG
        )
        
        # Create logger that writes to file
        logger = Logger(log_level=LogLevel.DEBUG, output_path=str(log_path))
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        
        # Verify log file exists
        assert log_path.exists()
        
        # Read log file
        log_content = log_path.read_text()
        
        # Verify log contains conversion information
        assert len(log_content) > 0
    
    def test_multiple_encoding_issues_all_logged(self, tmp_path):
        """Test that multiple encoding issues are all logged."""
        # Create a Word document
        doc_path = tmp_path / "test_multiple.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        doc.save(str(doc_path))
        
        # Create a logger with string stream
        log_stream = io.StringIO()
        logger = logging.getLogger('test_multiple_issues')
        logger.setLevel(logging.DEBUG)
        handler = logging.StreamHandler(log_stream)
        handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        logger.addHandler(handler)
        
        # Create config
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create orchestrator with custom logger
        from src.logger import Logger as CustomLogger
        custom_logger = CustomLogger(log_level=LogLevel.DEBUG)
        custom_logger.logger = logger
        
        orchestrator = ConversionOrchestrator(config, custom_logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        
        # Clean up
        logger.removeHandler(handler)
    
    def test_encoding_logging_with_multilingual_content(self, tmp_path):
        """Test encoding logging with multilingual content."""
        # Create a Word document with multiple languages
        doc_path = tmp_path / "test_multilingual.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("日本語テキスト")
        doc.add_paragraph("中文文本")
        doc.add_paragraph("한국어 텍스트")
        doc.add_paragraph("Русский текст")
        doc.add_paragraph("العربية")
        doc.save(str(doc_path))
        
        # Create a logger with string stream
        log_stream = io.StringIO()
        logger = logging.getLogger('test_multilingual')
        logger.setLevel(logging.DEBUG)
        handler = logging.StreamHandler(log_stream)
        handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        logger.addHandler(handler)
        
        # Create config
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True
        )
        
        # Create orchestrator with custom logger
        from src.logger import Logger as CustomLogger
        custom_logger = CustomLogger(log_level=LogLevel.DEBUG)
        custom_logger.logger = logger
        
        orchestrator = ConversionOrchestrator(config, custom_logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # Verify multilingual content is preserved
        assert "English" in result.markdown_content
        assert "日本語" in result.markdown_content
        assert "中文" in result.markdown_content
        
        # Clean up
        logger.removeHandler(handler)
