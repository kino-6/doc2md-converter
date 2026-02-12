"""Property-based tests for batch conversion functionality.

Feature: document-to-markdown-converter
Property 27: Batch conversion
Validates: Requirements 6.5
"""

import pytest
from hypothesis import given, strategies as st, settings
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
import tempfile


# Strategy for generating valid file names
filename_strategy = st.text(
    alphabet=st.characters(
        whitelist_categories=('Lu', 'Ll', 'Nd'),
        min_codepoint=ord('a'),
        max_codepoint=ord('z')
    ),
    min_size=1,
    max_size=20
)

# Strategy for generating text content compatible with Office documents
def is_valid_office_char(c):
    """Check if character is valid for Office documents."""
    code = ord(c)
    # Allow tab, newline, carriage return
    if code in (0x09, 0x0A, 0x0D):
        return True
    # Exclude other control characters (0x00-0x1F, 0x7F-0x9F)
    if code < 0x20 or (0x7F <= code <= 0x9F):
        return False
    # Exclude surrogates
    if 0xD800 <= code <= 0xDFFF:
        return False
    return True

text_content_strategy = st.text(
    alphabet=st.characters(
        blacklist_categories=('Cs',),  # Exclude surrogates
    ).filter(is_valid_office_char),
    min_size=1,
    max_size=200
)


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=10),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=10)
)
def test_property_27_batch_conversion_success(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of multiple valid input files, the converter should successfully 
    convert all files in a single batch operation.
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents
        file_paths = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Create orchestrator with batch mode
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True  # Don't write files
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: Number of results should match number of input files
        assert len(results) == num_files, \
            f"Expected {num_files} results, got {len(results)}"
        
        # Property: All conversions should succeed for valid files
        assert all(r.success for r in results), \
            f"Not all conversions succeeded: {[r.errors for r in results if not r.success]}"
        
        # Property: All results should have markdown content
        assert all(r.markdown_content is not None for r in results), \
            "All results should have markdown content"
        
        # Property: Conversion should produce output (may be empty for whitespace-only input)
        # Note: Empty or whitespace-only content may result in empty markdown
        assert all(r.markdown_content is not None for r in results), \
            "All results should have markdown content (even if empty)"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=2, max_value=8),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=2, max_size=8)
)
def test_property_27_batch_conversion_aggregates_results(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of multiple input files, the converter should aggregate all 
    conversion results correctly.
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents
        file_paths = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Create orchestrator with batch mode
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: Each result should have the correct input path
        for i, result in enumerate(results):
            assert result.input_path == file_paths[i], \
                f"Result {i} has incorrect input path"
        
        # Property: All results should be present in order
        assert len(results) == len(file_paths), \
            "Results should match input files count"
        
        # Property: Each result should be independent
        # (changing one file shouldn't affect others)
        for i, result in enumerate(results):
            assert result.success is True, \
                f"Result {i} should succeed independently"


@settings(max_examples=100)
@given(
    num_valid=st.integers(min_value=1, max_value=5),
    num_invalid=st.integers(min_value=1, max_value=3),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=5)
)
def test_property_27_batch_conversion_handles_failures(num_valid, num_invalid, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of input files containing both valid and invalid files, the 
    converter should handle failures gracefully and continue processing remaining files.
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for valid files
    num_valid = min(num_valid, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create valid Word documents
        file_paths = []
        for i in range(num_valid):
            doc_path = tmp_path / f"{base_filename}_valid_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Create invalid files (wrong format)
        for i in range(num_invalid):
            invalid_path = tmp_path / f"{base_filename}_invalid_{i}.txt"
            invalid_path.write_text("Invalid content")
            file_paths.append(str(invalid_path))
        
        # Create orchestrator with batch mode
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: Number of results should match total number of files
        assert len(results) == num_valid + num_invalid, \
            f"Expected {num_valid + num_invalid} results, got {len(results)}"
        
        # Property: Valid files should succeed
        valid_results = results[:num_valid]
        assert all(r.success for r in valid_results), \
            "Valid files should convert successfully"
        
        # Property: Invalid files should fail
        invalid_results = results[num_valid:]
        assert all(not r.success for r in invalid_results), \
            "Invalid files should fail conversion"
        
        # Property: Batch operation should complete despite failures
        # (all files should be processed)
        assert len(results) == len(file_paths), \
            "All files should be processed even with failures"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=8),
    base_filename=filename_strategy,
    file_type=st.sampled_from(['docx', 'xlsx'])
)
def test_property_27_batch_conversion_mixed_formats(num_files, base_filename, file_type):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of input files with different formats, the converter should 
    successfully convert all supported formats in a single batch operation.
    
    **Validates: Requirements 6.5**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create files of specified type
        file_paths = []
        for i in range(num_files):
            if file_type == 'docx':
                file_path = tmp_path / f"{base_filename}_{i}.docx"
                doc = Document()
                doc.add_paragraph(f"Content {i}")
                doc.save(str(file_path))
            else:  # xlsx
                file_path = tmp_path / f"{base_filename}_{i}.xlsx"
                wb = openpyxl.Workbook()
                ws = wb.active
                ws['A1'] = f"Content {i}"
                wb.save(str(file_path))
            
            file_paths.append(str(file_path))
        
        # Create orchestrator with batch mode
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: All conversions should succeed
        assert len(results) == num_files, \
            f"Expected {num_files} results, got {len(results)}"
        
        assert all(r.success for r in results), \
            f"All conversions should succeed: {[r.errors for r in results if not r.success]}"
        
        # Property: Each result should have markdown content
        assert all(r.markdown_content is not None for r in results), \
            "All results should have markdown content"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=10),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=10)
)
def test_property_27_batch_conversion_generates_output_paths(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of input files, the converter should generate appropriate output 
    paths for each file in batch mode.
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents
        file_paths = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Create orchestrator with batch mode (not dry-run to generate paths)
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=False
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: Each result should have an output path
        assert all(r.output_path is not None for r in results), \
            "All results should have output paths"
        
        # Property: Output paths should end with .md
        assert all(r.output_path.endswith('.md') for r in results), \
            "All output paths should end with .md"
        
        # Property: Output paths should be unique
        output_paths = [r.output_path for r in results]
        assert len(output_paths) == len(set(output_paths)), \
            "Output paths should be unique"
        
        # Property: Output paths should correspond to input files
        for i, result in enumerate(results):
            input_stem = Path(file_paths[i]).stem
            assert input_stem in result.output_path, \
                f"Output path should contain input filename stem"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=5),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=5)
)
def test_property_27_batch_conversion_preserves_order(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of input files, the converter should preserve the order of files 
    in the results.
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents with unique identifiable content
        file_paths = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            # Add unique identifier to content
            doc.add_paragraph(f"FILE_ID_{i}: {contents[i]}")
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Create orchestrator with batch mode
        config = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator = ConversionOrchestrator(config=config, logger=logger)
        
        # Perform batch conversion
        results = orchestrator.batch_convert(file_paths)
        
        # Property: Results should be in the same order as input files
        for i, result in enumerate(results):
            assert result.input_path == file_paths[i], \
                f"Result {i} should correspond to input file {i}"
            
            # Verify content matches the expected file
            # Note: Underscores may be escaped in Markdown
            file_id = f"FILE_ID_{i}"
            escaped_file_id = file_id.replace('_', r'\_')
            assert file_id in result.markdown_content or escaped_file_id in result.markdown_content, \
                f"Result {i} should contain content from file {i}"


@settings(max_examples=100)
@given(
    num_files=st.integers(min_value=1, max_value=8),
    base_filename=filename_strategy,
    contents=st.lists(text_content_strategy, min_size=1, max_size=8)
)
def test_property_27_batch_conversion_independent_results(num_files, base_filename, contents):
    """Feature: document-to-markdown-converter, Property 27: Batch conversion
    
    For any set of input files, each conversion result should be independent 
    (one file's conversion should not affect another).
    
    **Validates: Requirements 6.5**
    """
    # Ensure we have enough content for all files
    num_files = min(num_files, len(contents))
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create multiple Word documents
        file_paths = []
        for i in range(num_files):
            doc_path = tmp_path / f"{base_filename}_{i}.docx"
            doc = Document()
            doc.add_paragraph(contents[i])
            doc.save(str(doc_path))
            file_paths.append(str(doc_path))
        
        # Perform batch conversion
        config_batch = ConversionConfig(
            input_path="",
            batch_mode=True,
            dry_run=True
        )
        logger = Logger(log_level=LogLevel.ERROR)
        orchestrator_batch = ConversionOrchestrator(config=config_batch, logger=logger)
        batch_results = orchestrator_batch.batch_convert(file_paths)
        
        # Perform individual conversions
        individual_results = []
        for file_path in file_paths:
            config_single = ConversionConfig(
                input_path=file_path,
                dry_run=True
            )
            orchestrator_single = ConversionOrchestrator(config=config_single, logger=logger)
            result = orchestrator_single.convert(file_path)
            individual_results.append(result)
        
        # Property: Batch results should match individual results
        assert len(batch_results) == len(individual_results), \
            "Batch and individual result counts should match"
        
        for i in range(num_files):
            # Property: Success status should be the same
            assert batch_results[i].success == individual_results[i].success, \
                f"Result {i} success status should match"
            
            # Property: Markdown content should be the same
            if batch_results[i].success and individual_results[i].success:
                assert batch_results[i].markdown_content == individual_results[i].markdown_content, \
                    f"Result {i} content should match between batch and individual conversion"
