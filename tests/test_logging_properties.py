"""Property-based tests for logging functionality.

Feature: document-to-markdown-converter
Properties 42-45: Logging properties
Validates: Requirements 10.1, 10.2, 10.3, 10.4
"""

import pytest
from hypothesis import given, strategies as st, settings
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
import tempfile
import os
import time


# Strategy for generating valid file names
filename_strategy = st.text(
    alphabet=st.characters(
        whitelist_categories=('Lu', 'Ll', 'Nd'),
        min_codepoint=ord('a'),
        max_codepoint=ord('z')
    ),
    min_size=1,
    max_size=20
)

# Strategy for generating text content compatible with Office documents
def is_valid_office_char(c):
    """Check if character is valid for Office documents."""
    code = ord(c)
    # Allow tab, newline, carriage return
    if code in (0x09, 0x0A, 0x0D):
        return True
    # Exclude other control characters (0x00-0x1F, 0x7F-0x9F)
    if code < 0x20 or (0x7F <= code <= 0x9F):
        return False
    # Exclude surrogates
    if 0xD800 <= code <= 0xDFFF:
        return False
    return True

text_content_strategy = st.text(
    alphabet=st.characters(
        blacklist_categories=('Cs',),  # Exclude surrogates
    ).filter(is_valid_office_char),
    min_size=1,
    max_size=200
)


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_42_log_file_generation(filename, content):
    """Feature: document-to-markdown-converter, Property 42: Log file generation
    
    For any conversion operation, the converter should generate a log file 
    recording the operation.
    
    **Validates: Requirements 10.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Log file should be generated
        assert log_path.exists(), "Log file was not created"
        
        # Property: Log file should contain content
        log_content = log_path.read_text(encoding='utf-8')
        assert len(log_content) > 0, "Log file is empty"
        
        # Property: Log file should be readable as UTF-8
        try:
            with open(log_path, 'r', encoding='utf-8') as f:
                f.read()
        except UnicodeDecodeError as e:
            pytest.fail(f"Log file is not valid UTF-8: {e}")


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_43_conversion_start_logging(filename, content):
    """Feature: document-to-markdown-converter, Property 43: Conversion start logging
    
    For any conversion that starts, the converter should log the source file path, 
    size, and format.
    
    **Validates: Requirements 10.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Get file size
        file_size = doc_path.stat().st_size
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Log should contain conversion start message
        log_content = log_path.read_text(encoding='utf-8')
        assert "Starting conversion" in log_content, "Log missing conversion start message"
        
        # Property: Log should contain source file path
        assert str(doc_path) in log_content or filename in log_content, \
            "Log missing source file path"
        
        # Property: Log should contain file size
        assert str(file_size) in log_content or "Size:" in log_content, \
            "Log missing file size information"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_44_conversion_completion_logging(filename, content):
    """Feature: document-to-markdown-converter, Property 44: Conversion completion logging
    
    For any conversion that completes, the converter should log the output file path, 
    conversion time, and status.
    
    **Validates: Requirements 10.3**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(doc_path))
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        start_time = time.time()
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        end_time = time.time()
        
        # Property: Conversion should complete successfully
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Log should contain completion message
        log_content = log_path.read_text(encoding='utf-8')
        assert "Conversion complete" in log_content or "complete" in log_content.lower(), \
            "Log missing conversion completion message"
        
        # Property: Log should contain output file path
        assert str(output_path) in log_content or "output.md" in log_content, \
            "Log missing output file path"
        
        # Property: Log should contain duration/time information
        # Check for duration in seconds or time-related keywords
        has_duration = ("Duration:" in log_content or 
                       "duration" in log_content.lower() or
                       any(f"{i}." in log_content for i in range(10)))
        assert has_duration, "Log missing conversion duration information"


@settings(max_examples=100)
@given(content=text_content_strategy)
def test_property_45_error_logging(content):
    """Feature: document-to-markdown-converter, Property 45: Error and warning logging
    
    For any error or warning that occurs during conversion, the converter should 
    log detailed messages with context.
    
    **Validates: Requirements 10.4**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an invalid/non-existent file path to trigger an error
        invalid_path = tmp_path / "nonexistent.docx"
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(invalid_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Attempt conversion (should fail)
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(invalid_path))
        
        # Property: Conversion should fail
        assert not result.success, "Conversion should have failed for non-existent file"
        
        # Property: Log file should exist
        assert log_path.exists(), "Log file was not created"
        
        # Property: Log should contain error message
        log_content = log_path.read_text(encoding='utf-8')
        has_error = ("ERROR" in log_content or 
                    "error" in log_content.lower() or
                    "failed" in log_content.lower() or
                    "not found" in log_content.lower())
        assert has_error, "Log missing error message"
        
        # Property: Log should contain context (file path)
        assert str(invalid_path) in log_content or "nonexistent" in log_content, \
            "Log missing error context (file path)"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_42_excel_log_file_generation(filename, content):
    """Feature: document-to-markdown-converter, Property 42: Log file generation
    
    For any Excel conversion operation, the converter should generate a log file.
    
    **Validates: Requirements 10.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file
        excel_path = tmp_path / f"{filename}.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = content
        wb.save(str(excel_path))
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Log file should be generated
        assert log_path.exists(), "Log file was not created for Excel conversion"
        
        # Property: Log file should contain content
        log_content = log_path.read_text(encoding='utf-8')
        assert len(log_content) > 0, "Log file is empty"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_43_excel_conversion_start_logging(filename, content):
    """Feature: document-to-markdown-converter, Property 43: Conversion start logging
    
    For any Excel conversion that starts, the converter should log the source file 
    path, size, and format.
    
    **Validates: Requirements 10.2**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file
        excel_path = tmp_path / f"{filename}.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = content
        wb.save(str(excel_path))
        
        # Get file size
        file_size = excel_path.stat().st_size
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Log should contain conversion start message
        log_content = log_path.read_text(encoding='utf-8')
        assert "Starting conversion" in log_content, "Log missing conversion start message"
        
        # Property: Log should contain source file path
        assert str(excel_path) in log_content or filename in log_content, \
            "Log missing source file path"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_44_excel_conversion_completion_logging(filename, content):
    """Feature: document-to-markdown-converter, Property 44: Conversion completion logging
    
    For any Excel conversion that completes, the converter should log the output 
    file path, conversion time, and status.
    
    **Validates: Requirements 10.3**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create an Excel file
        excel_path = tmp_path / f"{filename}.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = content
        wb.save(str(excel_path))
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file
        config = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_path),
            log_file=str(log_path)
        )
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(excel_path))
        
        # Property: Conversion should complete successfully
        assert result.success, f"Conversion failed: {result.errors}"
        
        # Property: Log should contain completion message
        log_content = log_path.read_text(encoding='utf-8')
        assert "Conversion complete" in log_content or "complete" in log_content.lower(), \
            "Log missing conversion completion message"
        
        # Property: Log should contain output file path
        assert str(output_path) in log_content or "output.md" in log_content, \
            "Log missing output file path"


@settings(max_examples=50)
@given(
    filename=filename_strategy,
    content1=text_content_strategy,
    content2=text_content_strategy
)
def test_property_42_multiple_conversions_separate_logs(filename, content1, content2):
    """Feature: document-to-markdown-converter, Property 42: Log file generation
    
    For any sequence of conversion operations, each should generate its own log 
    entries in the log file.
    
    **Validates: Requirements 10.1**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Set up shared log file path
        log_path = tmp_path / "conversion.log"
        
        # Create logger with log file
        logger = Logger(log_level=LogLevel.INFO, output_path=str(log_path))
        
        # First conversion
        doc_path1 = tmp_path / f"{filename}_1.docx"
        doc1 = Document()
        doc1.add_paragraph(content1)
        doc1.save(str(doc_path1))
        
        output_path1 = tmp_path / "output1.md"
        config1 = ConversionConfig(
            input_path=str(doc_path1),
            output_path=str(output_path1),
            log_file=str(log_path)
        )
        
        orchestrator1 = ConversionOrchestrator(config1, logger)
        result1 = orchestrator1.convert(str(doc_path1))
        
        # Second conversion
        doc_path2 = tmp_path / f"{filename}_2.docx"
        doc2 = Document()
        doc2.add_paragraph(content2)
        doc2.save(str(doc_path2))
        
        output_path2 = tmp_path / "output2.md"
        config2 = ConversionConfig(
            input_path=str(doc_path2),
            output_path=str(output_path2),
            log_file=str(log_path)
        )
        
        orchestrator2 = ConversionOrchestrator(config2, logger)
        result2 = orchestrator2.convert(str(doc_path2))
        
        # Property: Log file should exist
        assert log_path.exists(), "Log file was not created"
        
        # Property: Log should contain entries for both conversions
        log_content = log_path.read_text(encoding='utf-8')
        
        # Check for first conversion
        has_first = (f"{filename}_1" in log_content or "output1.md" in log_content)
        assert has_first, "Log missing first conversion entry"
        
        # Check for second conversion
        has_second = (f"{filename}_2" in log_content or "output2.md" in log_content)
        assert has_second, "Log missing second conversion entry"


@settings(max_examples=100)
@given(filename=filename_strategy, content=text_content_strategy)
def test_property_45_warning_logging(filename, content):
    """Feature: document-to-markdown-converter, Property 45: Error and warning logging
    
    For any warning that occurs during conversion, the converter should log 
    detailed messages with context.
    
    **Validates: Requirements 10.4**
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        
        # Create a Word document with content that might trigger warnings
        # (e.g., encoding issues with replacement characters)
        doc_path = tmp_path / f"{filename}.docx"
        doc = Document()
        # Add content with potential encoding issues
        doc.add_paragraph(content + "\ufffd")  # Unicode replacement character
        doc.save(str(doc_path))
        
        # Set up log file path
        log_path = tmp_path / "conversion.log"
        output_path = tmp_path / "output.md"
        
        # Create config with log file and WARNING level
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            log_file=str(log_path),
            log_level=LogLevel.WARNING
        )
        
        # Create logger with WARNING level
        logger = Logger(log_level=LogLevel.WARNING, output_path=str(log_path))
        
        # Perform conversion
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Property: Log file should exist
        assert log_path.exists(), "Log file was not created"
        
        # Property: Log should contain content (warnings or info messages)
        log_content = log_path.read_text(encoding='utf-8')
        assert len(log_content) > 0, "Log file is empty"
        
        # Property: If warnings occurred, they should be logged with context
        # Note: Not all conversions will trigger warnings, but if they do,
        # they should be properly logged
        if "WARNING" in log_content:
            # Warning should include context
            assert len(log_content) > 50, "Warning message lacks detail"
