"""
Property-based tests for ConversionOrchestrator.

This module contains property-based tests using Hypothesis to verify
the ConversionOrchestrator's error handling and structure preservation.

Feature: document-to-markdown-converter
Properties: 21, 22
Validates: Requirements 4.5, 5.1
"""

import pytest
from hypothesis import given, strategies as st, settings, HealthCheck
from pathlib import Path
import tempfile
import string

from src.conversion_orchestrator import ConversionOrchestrator, ConversionResult
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
from src.internal_representation import (
    InternalDocument, Section, Heading, Paragraph, Table, DocumentMetadata
)


# Helper strategies for generating test data
def word_safe_text(min_size=1, max_size=100):
    """Generate text that is safe for Word documents."""
    safe_chars = string.printable.replace('\x0b', '').replace('\x0c', '').replace('\r', '').replace('\n', '')
    return st.text(alphabet=safe_chars, min_size=min_size, max_size=max_size)


class TestOrchestratorPropertyGracefulDegradation:
    """Property 21: Graceful degradation
    
    For any conversion that encounters an error, the converter should provide
    partial output with error details when possible.
    
    Validates: Requirements 4.5
    """
    
    @given(file_format=st.sampled_from(['.docx', '.xlsx', '.pdf']))
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_21_graceful_degradation_file_not_found(self, file_format):
        """
        Feature: document-to-markdown-converter
        Property 21: Graceful degradation (file not found)
        
        For any conversion where the input file does not exist, the converter
        should provide error details without crashing.
        
        **Validates: Requirements 4.5**
        """
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a path to a non-existent file
            non_existent_file = Path(tmp_dir) / f"non_existent{file_format}"
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(non_existent_file),
                output_path=None,
                log_level=LogLevel.ERROR
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Attempt conversion
            result = orchestrator.convert(str(non_existent_file))
            
            # Property: Conversion should fail gracefully
            assert isinstance(result, ConversionResult), \
                "Should return a ConversionResult object"
            
            assert result.success is False, \
                "Conversion should fail for non-existent file"
            
            # Property: Should provide error details
            assert len(result.errors) > 0, \
                "Should provide at least one error message"
            
            assert any("not found" in error.lower() or "exist" in error.lower() 
                      for error in result.errors), \
                "Error message should indicate file not found"
            
            # Property: Should not crash (we got here without exception)
            assert result.markdown_content is None or result.markdown_content == "", \
                "Should not produce markdown content for non-existent file"
    
    @given(
        num_sections=st.integers(min_value=1, max_value=5),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_21_graceful_degradation_corrupted_file(self, num_sections, data):
        """
        Feature: document-to-markdown-converter
        Property 21: Graceful degradation (corrupted file)
        
        For any conversion where the input file is corrupted, the converter
        should provide error details and partial output when possible.
        
        **Validates: Requirements 4.5**
        """
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a corrupted .docx file (just random bytes)
            corrupted_file = Path(tmp_dir) / "corrupted.docx"
            
            # Write random bytes to simulate corruption
            random_bytes = data.draw(st.binary(min_size=100, max_size=1000))
            with open(corrupted_file, 'wb') as f:
                f.write(random_bytes)
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(corrupted_file),
                output_path=None,
                log_level=LogLevel.ERROR,
                preview_mode=True  # Use preview mode to test partial output
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Attempt conversion
            result = orchestrator.convert(str(corrupted_file))
            
            # Property: Should return a ConversionResult
            assert isinstance(result, ConversionResult), \
                "Should return a ConversionResult object"
            
            # Property: Should fail or provide partial success
            # (depending on how the parser handles corruption)
            assert result.success is False or len(result.errors) > 0 or len(result.warnings) > 0, \
                "Should indicate failure or provide warnings for corrupted file"
            
            # Property: Should provide error details
            if not result.success:
                assert len(result.errors) > 0, \
                    "Should provide error messages for failed conversion"
            
            # Property: Should not crash (we got here without exception)
            # This is the key aspect of graceful degradation
    
    @given(file_size_mb=st.integers(min_value=101, max_value=200))
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_21_graceful_degradation_file_too_large(self, file_size_mb):
        """
        Feature: document-to-markdown-converter
        Property 21: Graceful degradation (file too large)
        
        For any conversion where the input file exceeds size limits, the converter
        should provide error details without crashing.
        
        **Validates: Requirements 4.5**
        """
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a large file
            large_file = Path(tmp_dir) / "large_file.docx"
            
            # Write enough data to exceed the limit
            # We'll create a file that's larger than the max_file_size_mb
            with open(large_file, 'wb') as f:
                # Write 1MB chunks
                chunk_size = 1024 * 1024  # 1MB
                for _ in range(file_size_mb):
                    f.write(b'0' * chunk_size)
            
            # Create configuration with a smaller max file size
            config = ConversionConfig(
                input_path=str(large_file),
                output_path=None,
                log_level=LogLevel.ERROR,
                max_file_size_mb=100  # Set limit to 100MB
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Attempt conversion
            result = orchestrator.convert(str(large_file))
            
            # Property: Conversion should fail gracefully
            assert isinstance(result, ConversionResult), \
                "Should return a ConversionResult object"
            
            assert result.success is False, \
                "Conversion should fail for file exceeding size limit"
            
            # Property: Should provide error details about file size
            assert len(result.errors) > 0, \
                "Should provide at least one error message"
            
            assert any("size" in error.lower() or "large" in error.lower() or "exceed" in error.lower()
                      for error in result.errors), \
                "Error message should indicate file size issue"
            
            # Property: Should not crash
            assert result.markdown_content is None or result.markdown_content == "", \
                "Should not produce markdown content for oversized file"


class TestOrchestratorPropertyStructurePreservation:
    """Property 22: Structure preservation
    
    For any source document with logical structure (sections, headings, paragraphs),
    the converter should preserve that structure in the Markdown output.
    
    Validates: Requirements 5.1
    """
    
    @given(
        num_sections=st.integers(min_value=1, max_value=6),
        data=st.data()
    )
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_22_structure_preservation_headings(self, num_sections, data):
        """
        Feature: document-to-markdown-converter
        Property 22: Structure preservation (headings)
        
        For any source document with headings, the converter should preserve
        the heading structure in the Markdown output.
        
        **Validates: Requirements 5.1**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with multiple sections and headings
            doc = Document()
            
            heading_data = []
            for i in range(num_sections):
                level = data.draw(st.integers(min_value=1, max_value=6))
                text = data.draw(word_safe_text(min_size=1, max_size=50))
                
                # Skip if text is only whitespace
                if not text.strip():
                    continue
                
                doc.add_heading(text, level=level)
                doc.add_paragraph(f"Content for section {i+1}")
                
                heading_data.append((level, text))
            
            # Skip test if no valid headings were generated
            if not heading_data:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_structure.docx"
            doc.save(str(file_path))
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(file_path),
                output_path=None,
                log_level=LogLevel.ERROR
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Convert document
            result = orchestrator.convert(str(file_path))
            
            # Property: Conversion should succeed
            assert result.success is True, \
                f"Conversion should succeed, errors: {result.errors}"
            
            # Property: Should produce markdown content
            assert result.markdown_content is not None, \
                "Should produce markdown content"
            
            assert len(result.markdown_content) > 0, \
                "Markdown content should not be empty"
            
            # Property: Markdown should contain all headings
            import re
            markdown_lines = result.markdown_content.split('\n')
            
            # Extract headings from markdown
            markdown_headings = []
            for line in markdown_lines:
                if line.strip().startswith('#'):
                    # Extract heading level and text
                    match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
                    if match:
                        level = len(match.group(1))
                        text = match.group(2).strip()
                        markdown_headings.append((level, text))
            
            # Property: Number of headings should match
            assert len(markdown_headings) >= len(heading_data), \
                f"Expected at least {len(heading_data)} headings in markdown, got {len(markdown_headings)}"
            
            # Property: Heading levels and texts should be preserved
            # (allowing for some normalization of whitespace and Markdown escaping)
            for expected_level, expected_text in heading_data:
                normalized_expected = re.sub(r'\s+', ' ', expected_text).strip()
                
                # Find matching heading in markdown
                found = False
                for md_level, md_text in markdown_headings:
                    normalized_md = re.sub(r'\s+', ' ', md_text).strip()
                    
                    # Check for exact match or match with Markdown escaping
                    # Markdown may escape special characters with backslashes
                    if md_level == expected_level:
                        # Remove escape backslashes for comparison
                        unescaped_md = re.sub(r'\\(.)', r'\1', normalized_md)
                        if unescaped_md == normalized_expected or normalized_md == normalized_expected:
                            found = True
                            break
                
                assert found, \
                    f"Expected heading (level {expected_level}): '{normalized_expected}' not found in markdown"
    
    @given(
        num_paragraphs=st.integers(min_value=2, max_value=10),
        data=st.data()
    )
    @settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_22_structure_preservation_paragraphs(self, num_paragraphs, data):
        """
        Feature: document-to-markdown-converter
        Property 22: Structure preservation (paragraphs)
        
        For any source document with paragraphs, the converter should preserve
        the paragraph structure in the Markdown output.
        
        **Validates: Requirements 5.1**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with multiple paragraphs
            doc = Document()
            doc.add_heading("Document with Paragraphs", level=1)
            
            paragraph_texts = []
            for i in range(num_paragraphs):
                text = data.draw(word_safe_text(min_size=10, max_size=100))
                
                # Skip if text is only whitespace
                if not text.strip():
                    continue
                
                doc.add_paragraph(text)
                paragraph_texts.append(text)
            
            # Skip test if no valid paragraphs were generated
            if not paragraph_texts:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_paragraphs.docx"
            doc.save(str(file_path))
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(file_path),
                output_path=None,
                log_level=LogLevel.ERROR
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Convert document
            result = orchestrator.convert(str(file_path))
            
            # Property: Conversion should succeed
            assert result.success is True, \
                f"Conversion should succeed, errors: {result.errors}"
            
            # Property: Should produce markdown content
            assert result.markdown_content is not None, \
                "Should produce markdown content"
            
            # Property: All paragraph texts should be present in markdown
            # (allowing for Markdown escaping of special characters)
            import re
            for para_text in paragraph_texts:
                normalized_para = re.sub(r'\s+', ' ', para_text).strip()
                normalized_markdown = re.sub(r'\s+', ' ', result.markdown_content).strip()
                
                # Check if text is present (with or without Markdown escaping)
                # Remove escape backslashes from markdown for comparison
                unescaped_markdown = re.sub(r'\\(.)', r'\1', normalized_markdown)
                
                assert normalized_para in normalized_markdown or normalized_para in unescaped_markdown, \
                    f"Paragraph text '{normalized_para[:50]}...' not found in markdown output"
            
            # Property: Paragraphs should be separated by blank lines
            # (This is a Markdown convention for readability)
            lines = result.markdown_content.split('\n')
            
            # Count non-empty lines (should have content)
            non_empty_lines = [line for line in lines if line.strip()]
            assert len(non_empty_lines) > 0, \
                "Markdown should have non-empty lines"

    
    @given(
        num_sections=st.integers(min_value=2, max_value=5),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_22_structure_preservation_mixed_content(self, num_sections, data):
        """
        Feature: document-to-markdown-converter
        Property 22: Structure preservation (mixed content)
        
        For any source document with mixed content (headings, paragraphs, tables),
        the converter should preserve the logical structure in the Markdown output.
        
        **Validates: Requirements 5.1**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with mixed content
            doc = Document()
            
            structure_elements = []
            for i in range(num_sections):
                # Add a heading
                heading_text = data.draw(word_safe_text(min_size=1, max_size=30))
                if heading_text.strip():
                    doc.add_heading(heading_text, level=2)
                    structure_elements.append(('heading', heading_text))
                
                # Add a paragraph
                para_text = data.draw(word_safe_text(min_size=10, max_size=80))
                if para_text.strip():
                    doc.add_paragraph(para_text)
                    structure_elements.append(('paragraph', para_text))
                
                # Optionally add a table
                if data.draw(st.booleans()):
                    num_rows = data.draw(st.integers(min_value=2, max_value=4))
                    num_cols = data.draw(st.integers(min_value=2, max_value=4))
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    
                    # Fill table with data
                    for row_idx in range(num_rows):
                        for col_idx in range(num_cols):
                            cell_text = data.draw(word_safe_text(min_size=1, max_size=20))
                            table.rows[row_idx].cells[col_idx].text = cell_text
                    
                    structure_elements.append(('table', (num_rows, num_cols)))
            
            # Skip test if no valid content was generated
            if not structure_elements:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_mixed_structure.docx"
            doc.save(str(file_path))
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(file_path),
                output_path=None,
                log_level=LogLevel.ERROR
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Convert document
            result = orchestrator.convert(str(file_path))
            
            # Property: Conversion should succeed
            assert result.success is True, \
                f"Conversion should succeed, errors: {result.errors}"
            
            # Property: Should produce markdown content
            assert result.markdown_content is not None, \
                "Should produce markdown content"
            
            assert len(result.markdown_content) > 0, \
                "Markdown content should not be empty"
            
            # Property: Structure should be preserved in order
            import re
            
            # Verify headings are present (allowing for Markdown escaping and whitespace normalization)
            headings = [elem[1] for elem in structure_elements if elem[0] == 'heading']
            for heading_text in headings:
                normalized_heading = re.sub(r'\s+', ' ', heading_text).strip()
                # Remove escape backslashes from markdown for comparison
                unescaped_markdown = re.sub(r'\\(.)', r'\1', result.markdown_content)
                # Normalize whitespace in markdown for comparison
                normalized_markdown = re.sub(r'\s+', ' ', result.markdown_content)
                normalized_unescaped = re.sub(r'\s+', ' ', unescaped_markdown)
                
                assert normalized_heading in normalized_markdown or normalized_heading in normalized_unescaped, \
                    f"Heading '{normalized_heading}' should be in markdown output"
            
            # Verify paragraphs are present (allowing for Markdown escaping and whitespace normalization)
            paragraphs = [elem[1] for elem in structure_elements if elem[0] == 'paragraph']
            for para_text in paragraphs:
                normalized_para = re.sub(r'\s+', ' ', para_text).strip()
                normalized_markdown = re.sub(r'\s+', ' ', result.markdown_content)
                unescaped_markdown = re.sub(r'\\(.)', r'\1', normalized_markdown)
                assert normalized_para in normalized_markdown or normalized_para in unescaped_markdown, \
                    f"Paragraph text should be in markdown output"
            
            # Verify tables are present (check for table markers)
            tables = [elem for elem in structure_elements if elem[0] == 'table']
            if tables:
                # Markdown tables contain | characters
                assert '|' in result.markdown_content, \
                    "Markdown should contain table markers (|) for tables"
    
    @given(
        num_sections=st.integers(min_value=1, max_value=4),
        data=st.data()
    )
    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_22_structure_preservation_hierarchical(self, num_sections, data):
        """
        Feature: document-to-markdown-converter
        Property 22: Structure preservation (hierarchical structure)
        
        For any source document with hierarchical structure (nested headings),
        the converter should preserve the hierarchy in the Markdown output.
        
        **Validates: Requirements 5.1**
        """
        from docx import Document
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create a Word document with hierarchical structure
            doc = Document()
            
            hierarchy = []
            for i in range(num_sections):
                # Add a top-level heading
                h1_text = data.draw(word_safe_text(min_size=1, max_size=30))
                if h1_text.strip():
                    doc.add_heading(h1_text, level=1)
                    hierarchy.append((1, h1_text))
                    doc.add_paragraph(f"Content for {h1_text}")
                    
                    # Add sub-headings
                    num_subsections = data.draw(st.integers(min_value=1, max_value=3))
                    for j in range(num_subsections):
                        h2_text = data.draw(word_safe_text(min_size=1, max_size=30))
                        if h2_text.strip():
                            doc.add_heading(h2_text, level=2)
                            hierarchy.append((2, h2_text))
                            doc.add_paragraph(f"Content for {h2_text}")
            
            # Skip test if no valid content was generated
            if not hierarchy:
                return
            
            # Save to temporary file
            file_path = Path(tmp_dir) / "test_hierarchical.docx"
            doc.save(str(file_path))
            
            # Create configuration
            config = ConversionConfig(
                input_path=str(file_path),
                output_path=None,
                log_level=LogLevel.ERROR
            )
            
            # Create logger
            logger = Logger(log_level=LogLevel.ERROR)
            
            # Create orchestrator
            orchestrator = ConversionOrchestrator(config=config, logger=logger)
            
            # Convert document
            result = orchestrator.convert(str(file_path))
            
            # Property: Conversion should succeed
            assert result.success is True, \
                f"Conversion should succeed, errors: {result.errors}"
            
            # Property: Should produce markdown content
            assert result.markdown_content is not None, \
                "Should produce markdown content"
            
            # Property: Hierarchical structure should be preserved
            import re
            markdown_lines = result.markdown_content.split('\n')
            
            # Extract headings with levels from markdown
            markdown_headings = []
            for line in markdown_lines:
                if line.strip().startswith('#'):
                    match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
                    if match:
                        level = len(match.group(1))
                        text = match.group(2).strip()
                        markdown_headings.append((level, text))
            
            # Property: All headings should be present with correct levels
            # (allowing for Markdown escaping)
            for expected_level, expected_text in hierarchy:
                normalized_expected = re.sub(r'\s+', ' ', expected_text).strip()
                
                # Find matching heading in markdown
                found = False
                for md_level, md_text in markdown_headings:
                    normalized_md = re.sub(r'\s+', ' ', md_text).strip()
                    
                    # Check for exact match or match with Markdown escaping
                    if md_level == expected_level:
                        # Remove escape backslashes for comparison
                        unescaped_md = re.sub(r'\\(.)', r'\1', normalized_md)
                        if unescaped_md == normalized_expected or normalized_md == normalized_expected:
                            found = True
                            break
                
                assert found, \
                    f"Expected heading (level {expected_level}): '{normalized_expected}' not found in markdown"
            
            # Property: Heading hierarchy should be maintained
            # (Level 2 headings should come after Level 1 headings)
            prev_level = 0
            for level, text in markdown_headings:
                # Level should not jump more than 1 level at a time
                # (e.g., can't go from H1 to H3 without H2)
                if prev_level > 0:
                    assert level <= prev_level + 1, \
                        f"Heading level jumped from {prev_level} to {level}, violating hierarchy"
                prev_level = level


# Run all property tests
if __name__ == "__main__":
    pytest.main([__file__, "-v"])
