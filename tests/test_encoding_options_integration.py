"""Integration tests for output encoding options.

This module contains integration tests to verify that output encoding
options work correctly when converting documents to Markdown.

Validates: Requirements 9.4
"""

import pytest
from pathlib import Path
from docx import Document
import openpyxl
from src.conversion_orchestrator import ConversionOrchestrator
from src.config import ConversionConfig
from src.logger import Logger, LogLevel
import tempfile


class TestEncodingOptionsIntegration:
    """Integration tests for output encoding selection.
    
    Validates: Requirements 9.4 - WHERE encoding options are specified,
    THE Converter SHALL allow users to choose the output encoding format
    """
    
    def test_utf8_encoding_output(self, tmp_path):
        """Test that UTF-8 encoding is used when specified."""
        # Create a Word document with multilingual content
        doc_path = tmp_path / "test_utf8.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("日本語テキスト")
        doc.add_paragraph("中文文本")
        doc.add_paragraph("Español: café")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_utf8.md"
        
        # Create config with UTF-8 encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            output_encoding="utf-8"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Read the file with UTF-8 encoding
        content = output_path.read_text(encoding='utf-8')
        
        # Verify all characters are preserved correctly
        assert "English text" in content
        assert "日本語テキスト" in content
        assert "中文文本" in content
        assert "café" in content
    
    def test_utf16_encoding_output(self, tmp_path):
        """Test that UTF-16 encoding is used when specified."""
        # Create a Word document with multilingual content
        doc_path = tmp_path / "test_utf16.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("日本語テキスト")
        doc.add_paragraph("中文文本")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_utf16.md"
        
        # Create config with UTF-16 encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            output_encoding="utf-16"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Read the file with UTF-16 encoding
        content = output_path.read_text(encoding='utf-16')
        
        # Verify all characters are preserved correctly
        assert "English text" in content
        assert "日本語テキスト" in content
        assert "中文文本" in content
    
    def test_latin1_encoding_output(self, tmp_path):
        """Test that Latin-1 encoding is used when specified."""
        # Create a Word document with Latin-1 compatible content
        doc_path = tmp_path / "test_latin1.docx"
        doc = Document()
        doc.add_paragraph("English text")
        doc.add_paragraph("Español: café, niño")
        doc.add_paragraph("Français: résumé, naïve")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_latin1.md"
        
        # Create config with Latin-1 encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            output_encoding="latin-1"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Read the file with Latin-1 encoding
        content = output_path.read_text(encoding='latin-1')
        
        # Verify Latin-1 compatible characters are preserved
        assert "English text" in content
        assert "café" in content
        assert "résumé" in content
    
    def test_ascii_encoding_with_replacement(self, tmp_path):
        """Test that ASCII encoding handles non-ASCII characters."""
        # Create a Word document with mixed content
        doc_path = tmp_path / "test_ascii.docx"
        doc = Document()
        doc.add_paragraph("English text only")
        doc.add_paragraph("Special: café")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_ascii.md"
        
        # Create config with ASCII encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            output_encoding="ascii"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Read the file with ASCII encoding (with error handling)
        content = output_path.read_text(encoding='ascii', errors='replace')
        
        # Verify ASCII-compatible text is preserved
        assert "English text only" in content
    
    def test_encoding_option_from_config_file(self, tmp_path):
        """Test that encoding option is loaded from config file."""
        # Create a Word document
        doc_path = tmp_path / "test_config.docx"
        doc = Document()
        doc.add_paragraph("Test content with special chars: café")
        doc.save(str(doc_path))
        
        # Create a config file with UTF-16 encoding
        config_path = tmp_path / "config.yaml"
        config_path.write_text("""
input_path: test.docx
output_encoding: utf-16
""")
        
        # Create output path
        output_path = tmp_path / "output_config.md"
        
        # Load config from file
        from src.config import ConfigManager
        config_manager = ConfigManager()
        config = config_manager.load_config(str(config_path))
        
        # Override paths
        config.input_path = str(doc_path)
        config.output_path = str(output_path)
        
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Verify file can be read with UTF-16
        content = output_path.read_text(encoding='utf-16')
        assert "café" in content
    
    def test_excel_with_different_encodings(self, tmp_path):
        """Test Excel conversion with different output encodings."""
        # Create an Excel file with Unicode content
        excel_path = tmp_path / "test_excel.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = "Language"
        ws['B1'] = "Greeting"
        ws['A2'] = "English"
        ws['B2'] = "Hello"
        ws['A3'] = "Japanese"
        ws['B3'] = "こんにちは"
        ws['A4'] = "Spanish"
        ws['B4'] = "Hola"
        wb.save(str(excel_path))
        
        # Test with UTF-8
        output_utf8 = tmp_path / "excel_utf8.md"
        config_utf8 = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_utf8),
            output_encoding="utf-8"
        )
        logger = Logger(log_level=LogLevel.INFO)
        orchestrator = ConversionOrchestrator(config_utf8, logger)
        result = orchestrator.convert(str(excel_path))
        
        assert result.success
        assert output_utf8.exists()
        content_utf8 = output_utf8.read_text(encoding='utf-8')
        assert "こんにちは" in content_utf8
        
        # Test with UTF-16
        output_utf16 = tmp_path / "excel_utf16.md"
        config_utf16 = ConversionConfig(
            input_path=str(excel_path),
            output_path=str(output_utf16),
            output_encoding="utf-16"
        )
        orchestrator = ConversionOrchestrator(config_utf16, logger)
        result = orchestrator.convert(str(excel_path))
        
        assert result.success
        assert output_utf16.exists()
        content_utf16 = output_utf16.read_text(encoding='utf-16')
        assert "こんにちは" in content_utf16
    
    def test_default_encoding_is_utf8(self, tmp_path):
        """Test that default encoding is UTF-8 when not specified."""
        # Create a Word document
        doc_path = tmp_path / "test_default.docx"
        doc = Document()
        doc.add_paragraph("Test with Unicode: 日本語")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_default.md"
        
        # Create config without specifying encoding (should default to UTF-8)
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path)
        )
        
        # Verify default encoding is UTF-8
        assert config.output_encoding == "utf-8"
        
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert output_path.exists()
        
        # Verify file is UTF-8 encoded
        content = output_path.read_text(encoding='utf-8')
        assert "日本語" in content
    
    def test_encoding_option_with_preview_mode(self, tmp_path):
        """Test that encoding option works with preview mode."""
        # Create a Word document
        doc_path = tmp_path / "test_preview.docx"
        doc = Document()
        doc.add_paragraph("Preview test: café")
        doc.save(str(doc_path))
        
        # Create config with preview mode and UTF-16 encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=None,
            preview_mode=True,
            output_encoding="utf-16"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Verify conversion succeeded
        assert result.success
        assert result.markdown_content is not None
        
        # In preview mode, content should still be in memory as string
        assert "café" in result.markdown_content
    
    def test_encoding_option_with_batch_conversion(self, tmp_path):
        """Test that encoding option works with batch conversion."""
        # Create multiple Word documents
        doc1_path = tmp_path / "test1.docx"
        doc1 = Document()
        doc1.add_paragraph("Document 1: café")
        doc1.save(str(doc1_path))
        
        doc2_path = tmp_path / "test2.docx"
        doc2 = Document()
        doc2.add_paragraph("Document 2: résumé")
        doc2.save(str(doc2_path))
        
        # Create config with UTF-16 encoding
        config = ConversionConfig(
            input_path=str(tmp_path),
            output_path=None,  # Will be auto-generated
            output_encoding="utf-16",
            batch_mode=True
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert documents
        orchestrator = ConversionOrchestrator(config, logger)
        results = orchestrator.batch_convert([str(doc1_path), str(doc2_path)])
        
        # Verify both conversions succeeded
        assert len(results) == 2
        assert all(r.success for r in results)
        
        # Verify output files exist in same directory as input (batch mode behavior)
        output1 = tmp_path / "test1.md"
        output2 = tmp_path / "test2.md"
        
        assert output1.exists()
        assert output2.exists()
        
        content1 = output1.read_text(encoding='utf-16')
        content2 = output2.read_text(encoding='utf-16')
        
        assert "café" in content1
        assert "résumé" in content2
    
    def test_unsupported_encoding_fallback(self, tmp_path):
        """Test that unsupported encoding falls back gracefully."""
        # Create a Word document
        doc_path = tmp_path / "test_unsupported.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(doc_path))
        
        # Create output path
        output_path = tmp_path / "output_unsupported.md"
        
        # Create config with an invalid encoding
        config = ConversionConfig(
            input_path=str(doc_path),
            output_path=str(output_path),
            output_encoding="invalid-encoding-xyz"
        )
        logger = Logger(log_level=LogLevel.INFO)
        
        # Convert document - should handle gracefully
        orchestrator = ConversionOrchestrator(config, logger)
        result = orchestrator.convert(str(doc_path))
        
        # Conversion might fail or fall back to UTF-8
        # Either behavior is acceptable for invalid encoding
        if result.success:
            # If it succeeded, it should have fallen back to UTF-8
            assert output_path.exists()
            content = output_path.read_text(encoding='utf-8')
            assert "Test content" in content
        else:
            # If it failed, there should be an error message
            assert len(result.errors) > 0


# Run tests
if __name__ == "__main__":
    pytest.main([__file__, "-v"])
