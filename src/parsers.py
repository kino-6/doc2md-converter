"""Document parser interfaces and base classes.

This module defines the abstract base class for document parsers
and will contain concrete implementations for Word, Excel, and PDF parsers.
"""

from abc import ABC, abstractmethod
from src.internal_representation import InternalDocument


class DocumentParser(ABC):
    """Abstract base class for document parsers.
    
    All format-specific parsers (Word, Excel, PDF) must inherit from this class
    and implement the parse method.
    """
    
    @abstractmethod
    def parse(self, file_path: str) -> InternalDocument:
        """Parse a document file into internal representation.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            InternalDocument representation of the parsed document
            
        Raises:
            NotImplementedError: If the parser is not implemented
        """
        pass


class WordParser(DocumentParser):
    """Parser for Word (.docx) documents.

    Uses python-docx library to extract text content, headings, lists, tables,
    formatting, links, and image references from Word documents.
    """

    def __init__(self, logger=None):
        """Initialize Word parser with optional logger.

        Args:
            logger: Optional logger for encoding warnings
        """
        from src.encoding_detector import EncodingDetector
        self.encoding_detector = EncodingDetector(logger=logger)

    def parse(self, file_path: str) -> InternalDocument:
        """Parse a Word document into internal representation.

        Args:
            file_path: Path to the .docx file

        Returns:
            InternalDocument representation with extracted content

        Raises:
            FileNotFoundError: If the file does not exist
            Exception: If the file cannot be parsed
        """
        from docx import Document
        from src.internal_representation import (
            InternalDocument, DocumentMetadata, Section, Paragraph,
            Heading, Table, DocumentList, ListItem, ImageReference, Link,
            TextFormatting
        )

        try:
            doc = Document(file_path)
        except Exception as e:
            raise Exception(f"Failed to parse Word document: {str(e)}")

        # Extract metadata
        metadata = self._extract_metadata(doc, file_path)

        # Create internal document
        internal_doc = InternalDocument(metadata=metadata)

        # Create a single section for basic text extraction
        section = Section()

        # Extract text content from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                # Validate and normalize text encoding
                normalized_text = self._process_text_encoding(para.text)

                # Check if this is a heading
                if para.style.name.startswith('Heading'):
                    heading = self._extract_heading(para, normalized_text)
                    if heading:
                        # Start a new section with this heading
                        if section.content or section.heading:
                            internal_doc.sections.append(section)
                        section = Section(heading=heading)
                else:
                    # Check if this is a list item
                    list_item = self._extract_list_item(para, normalized_text)
                    if list_item:
                        # Group consecutive list items into DocumentList
                        # Check if the last content item is a list
                        if section.content and isinstance(section.content[-1], DocumentList):
                            # Add to existing list
                            section.content[-1].items.append(list_item)
                        else:
                            # Create new list
                            # Determine if ordered or unordered based on style
                            is_ordered = 'List Number' in para.style.name or 'Ordered' in para.style.name
                            doc_list = DocumentList(ordered=is_ordered, items=[list_item])
                            section.content.append(doc_list)
                    else:
                        # Extract hyperlinks from paragraph
                        links = self._extract_links(para)
                        for link in links:
                            section.content.append(link)

                        # Regular paragraph (only add if no links were extracted or text remains)
                        if not links or normalized_text.strip():
                            paragraph = self._extract_paragraph(para, normalized_text)
                            section.content.append(paragraph)

        # Extract tables
        for table in doc.tables:
            table_obj = self._extract_table(table)
            if table_obj:
                section.content.append(table_obj)

        # Extract images
        images = self._extract_images(doc)
        internal_doc.images.extend(images)

        # Add image references to content
        for img in images:
            section.content.append(img)

        # Add the last section if it has content
        if section.content or section.heading:
            internal_doc.sections.append(section)

        return internal_doc

    def _process_text_encoding(self, text: str) -> str:
        """Process and validate text encoding.

        Args:
            text: Extracted text to process

        Returns:
            Normalized text with encoding issues resolved
        """
        # Validate text encoding and detect issues
        validation_result = self.encoding_detector.validate_text_encoding(text)

        # Log any encoding issues found
        if validation_result.has_issues:
            for issue in validation_result.issues:
                if self.encoding_detector.logger:
                    self.encoding_detector.logger.warning(
                        f"Encoding issue in Word document: {issue}"
                    )

        # Normalize the text
        normalized_text = self.encoding_detector.normalize_text(text)

        return normalized_text

    def _extract_metadata(self, doc, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from Word document.

        Args:
            doc: python-docx Document object
            file_path: Path to the source file

        Returns:
            DocumentMetadata object
        """
        from src.internal_representation import DocumentMetadata

        core_props = doc.core_properties

        return DocumentMetadata(
            title=core_props.title if core_props.title else None,
            author=core_props.author if core_props.author else None,
            created_date=str(core_props.created) if core_props.created else None,
            modified_date=str(core_props.modified) if core_props.modified else None,
            source_format="docx"
        )

    def _extract_heading(self, para, text: str = None) -> 'Heading':
        """Extract heading from a paragraph with heading style.

        Args:
            para: python-docx Paragraph object with heading style
            text: Optional pre-processed text (if None, uses para.text)

        Returns:
            Heading object or None if not a valid heading
        """
        from src.internal_representation import Heading

        style_name = para.style.name
        heading_text = text if text is not None else para.text

        # Extract heading level from style name (e.g., "Heading 1" -> 1)
        if style_name.startswith('Heading'):
            try:
                level_str = style_name.replace('Heading', '').strip()
                level = int(level_str) if level_str else 1
                return Heading(level=level, text=heading_text)
            except (ValueError, AttributeError):
                return None

        return None

    def _extract_paragraph(self, para, text: str = None) -> 'Paragraph':
        """Extract paragraph with text formatting detection.

        Args:
            para: python-docx Paragraph object
            text: Optional pre-processed text (if None, uses para.text)

        Returns:
            Paragraph object with detected formatting
        """
        from src.internal_representation import Paragraph, TextFormatting

        paragraph_text = text if text is not None else para.text

        # Detect formatting by checking runs
        has_bold = False
        has_italic = False

        for run in para.runs:
            if run.bold:
                has_bold = True
            if run.italic:
                has_italic = True

        # Determine overall formatting
        if has_bold and has_italic:
            formatting = TextFormatting.BOLD_ITALIC
        elif has_bold:
            formatting = TextFormatting.BOLD
        elif has_italic:
            formatting = TextFormatting.ITALIC
        else:
            formatting = TextFormatting.NORMAL

        return Paragraph(text=paragraph_text, formatting=formatting)

    def _extract_links(self, para) -> list:
        """Extract hyperlinks from a paragraph.

        Args:
            para: python-docx Paragraph object

        Returns:
            List of Link objects
        """
        from src.internal_representation import Link

        links = []

        # Access hyperlinks through the paragraph's XML element
        if hasattr(para, '_element'):
            # Find all hyperlink elements
            hyperlinks = para._element.xpath('.//w:hyperlink')

            for hyperlink in hyperlinks:
                # Get the relationship ID
                r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

                if r_id:
                    try:
                        # Get the URL from the document relationships
                        url = para.part.rels[r_id].target_ref

                        # Get the link text
                        text_elements = hyperlink.xpath('.//w:t')
                        text = ''.join([t.text for t in text_elements if t.text])

                        # Normalize link text encoding
                        if text:
                            text = self._process_text_encoding(text)

                        if text and url:
                            links.append(Link(text=text, url=url))
                    except (KeyError, AttributeError):
                        # Skip if relationship not found
                        pass

        return links

    def _extract_images(self, doc) -> list:
        """Extract image references from Word document.

        Args:
            doc: python-docx Document object

        Returns:
            List of ImageReference objects
        """
        from src.internal_representation import ImageReference
        import base64

        images = []
        image_counter = 1

        # Iterate through all relationships to find images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # Get image bytes
                image_part = rel.target_part
                image_bytes = image_part.blob

                # Detect MIME type from content type
                content_type = image_part.content_type
                # content_type is like 'image/png', 'image/jpeg', etc.
                mime_type = content_type if content_type else "image/png"

                # Convert to base64
                base64_data = base64.b64encode(image_bytes).decode('utf-8')

                # Create image reference
                img_ref = ImageReference(
                    source_path=rel.target_ref,
                    alt_text=f"Image {image_counter}",
                    base64_data=base64_data,
                    mime_type=mime_type
                )
                images.append(img_ref)
                image_counter += 1

        return images

    def _extract_list_item(self, para, text: str = None) -> 'ListItem':
        """Extract list item from a paragraph with list style.

        Args:
            para: python-docx Paragraph object
            text: Optional pre-processed text (if None, uses para.text)

        Returns:
            ListItem object or None if not a list item
        """
        from src.internal_representation import ListItem

        style_name = para.style.name
        list_text = text if text is not None else para.text

        # Check if this is a list style
        if 'List' in style_name or para._element.xpath('.//w:numPr'):
            # Determine indentation level
            level = 0
            if hasattr(para, '_element'):
                num_pr = para._element.xpath('.//w:numPr/w:ilvl')
                if num_pr:
                    try:
                        level = int(num_pr[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0))
                    except (ValueError, AttributeError):
                        level = 0

            return ListItem(text=list_text, level=level)

        return None

    def _extract_table(self, table) -> 'Table':
        """Extract table structure from Word table.

        Args:
            table: python-docx Table object

        Returns:
            Table object or None if table is empty
        """
        from src.internal_representation import Table

        if not table.rows:
            return None

        # Extract headers from first row
        headers = []
        if table.rows:
            first_row = table.rows[0]
            headers = [
                self._process_text_encoding(cell.text.strip())
                for cell in first_row.cells
            ]

        # Extract data rows (skip first row if it's the header)
        rows = []
        for row in table.rows[1:]:
            row_data = [
                self._process_text_encoding(cell.text.strip())
                for cell in row.cells
            ]
            rows.append(row_data)

        # If no data rows, treat first row as data
        if not rows and headers:
            rows = [headers]
            headers = [f"Column {i+1}" for i in range(len(headers))]

        return Table(headers=headers, rows=rows)



class ExcelParser(DocumentParser):
    """Parser for Excel (.xlsx) documents.

    Extracts all sheets, handles formulas, merged cells, hyperlinks,
    date/time formatting, and hidden rows/columns.
    """

    def __init__(self, logger=None):
        """Initialize Excel parser with optional logger.

        Args:
            logger: Optional logger for encoding warnings
        """
        from src.encoding_detector import EncodingDetector
        self.encoding_detector = EncodingDetector(logger=logger)

    def parse(self, file_path: str) -> InternalDocument:
        """Parse an Excel document.

        Args:
            file_path: Path to the .xlsx file

        Returns:
            InternalDocument representation
        """
        try:
            import openpyxl
            from openpyxl.utils import get_column_letter
            from openpyxl.utils.exceptions import InvalidFileException
        except ImportError:
            raise ImportError("openpyxl is required for Excel parsing. Install it with: pip install openpyxl")

        try:
            # Load workbook with formulas
            workbook = openpyxl.load_workbook(file_path, data_only=False)
            # Load workbook with calculated values for formula evaluation
            workbook_data = openpyxl.load_workbook(file_path, data_only=True)
        except InvalidFileException as e:
            raise ValueError(f"Invalid Excel file: {e}")
        except Exception as e:
            raise ValueError(f"Failed to load Excel file: {e}")

        # Extract metadata
        metadata = self._extract_metadata(workbook, file_path)

        # Extract all sheets
        sections = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = workbook_data[sheet_name]
            section = self._extract_sheet(sheet, sheet_data)
            sections.append(section)

        return InternalDocument(
            metadata=metadata,
            sections=sections,
            images=[]
        )

    def _extract_metadata(self, workbook, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from Excel workbook.

        Args:
            workbook: openpyxl Workbook object
            file_path: Path to the Excel file

        Returns:
            DocumentMetadata object
        """
        from src.internal_representation import DocumentMetadata
        import os

        props = workbook.properties

        return DocumentMetadata(
            title=props.title if props.title else os.path.basename(file_path),
            author=props.creator if props.creator else None,
            created_date=props.created.isoformat() if props.created else None,
            modified_date=props.modified.isoformat() if props.modified else None,
            source_format="xlsx"
        )

    def _extract_sheet(self, sheet, sheet_data) -> 'Section':
        """Extract a single sheet as a Section.

        Args:
            sheet: openpyxl Worksheet object (with formulas)
            sheet_data: openpyxl Worksheet object (with calculated values)

        Returns:
            Section object containing the sheet data
        """
        from src.internal_representation import Section, Heading, Table, Paragraph, TextFormatting

        # Create heading for sheet name
        heading = Heading(level=2, text=sheet.title)

        # Check if sheet has any rows/columns at all
        if not sheet.max_row or not sheet.max_column:
            return Section(
                heading=heading,
                content=[Paragraph(text="(Empty sheet)", formatting=TextFormatting.ITALIC)]
            )

        # Check if sheet is truly empty (all cells are None)
        has_any_value = False
        for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row):
            for cell in row:
                if cell.value is not None:  # Even empty string '' is a value
                    has_any_value = True
                    break
            if has_any_value:
                break

        if not has_any_value:
            return Section(
                heading=heading,
                content=[Paragraph(text="(Empty sheet)", formatting=TextFormatting.ITALIC)]
            )

        # Extract table data from sheet
        table = self._extract_table_from_sheet(sheet, sheet_data)

        # If table has no structure at all, indicate empty
        if not table.headers and not table.rows:
            return Section(
                heading=heading,
                content=[Paragraph(text="(Empty sheet)", formatting=TextFormatting.ITALIC)]
            )

        return Section(
            heading=heading,
            content=[table]
        )

    def _extract_table_from_sheet(self, sheet, sheet_data) -> 'Table':
        """Extract table data from a worksheet.

        Handles:
        - Formula evaluation (calculated values)
        - Merged cells (value expansion)
        - Formula errors (#DIV/0!, #VALUE!, #REF!, etc.)
        - Hyperlinks (converted to Markdown format)
        - Date/time formatting
        - Hidden rows/columns (included by default)

        Args:
            sheet: openpyxl Worksheet object (with formulas)
            sheet_data: openpyxl Worksheet object (with calculated values)

        Returns:
            Table object
        """
        from src.internal_representation import Table
        import datetime

        # Get all data including hidden rows/columns
        rows_data = []

        # Process merged cells - create a map of merged cell ranges
        merged_cells_map = {}
        for merged_range in sheet.merged_cells.ranges:
            # Get the top-left cell value
            min_row, min_col = merged_range.min_row, merged_range.min_col
            top_left_cell = sheet.cell(min_row, min_col)
            top_left_cell_data = sheet_data.cell(min_row, min_col)
            value = self._get_cell_value(top_left_cell, top_left_cell_data)

            # Map all cells in the merged range to this value
            for row in range(merged_range.min_row, merged_range.max_row + 1):
                for col in range(merged_range.min_col, merged_range.max_col + 1):
                    merged_cells_map[(row, col)] = value

        # Extract all rows
        for row_idx, row in enumerate(sheet.iter_rows(min_row=1, max_row=sheet.max_row), start=1):
            row_data = []
            for col_idx, cell in enumerate(row, start=1):
                # Check if this cell is part of a merged range
                if (row_idx, col_idx) in merged_cells_map:
                    value = merged_cells_map[(row_idx, col_idx)]
                else:
                    cell_data = sheet_data.cell(row_idx, col_idx)
                    value = self._get_cell_value(cell, cell_data)

                row_data.append(value)

            rows_data.append(row_data)

        # If no data, return empty table
        if not rows_data:
            return Table(headers=[], rows=[])

        # Use first row as headers - validate encoding
        headers = []
        for cell in rows_data[0]:
            cell_str = str(cell) if cell is not None else ""
            if cell_str:
                cell_str = self._process_text_encoding(cell_str)
            headers.append(cell_str)

        # Rest are data rows - validate encoding
        data_rows = []
        for row in rows_data[1:]:
            processed_row = []
            for cell in row:
                cell_str = str(cell) if cell is not None else ""
                if cell_str:
                    cell_str = self._process_text_encoding(cell_str)
                processed_row.append(cell_str)
            data_rows.append(processed_row)

        return Table(headers=headers, rows=data_rows)

    def _process_text_encoding(self, text: str) -> str:
        """Process and validate text encoding.

        Args:
            text: Extracted text to process

        Returns:
            Normalized text with encoding issues resolved
        """
        # Validate text encoding and detect issues
        validation_result = self.encoding_detector.validate_text_encoding(text)

        # Log any encoding issues found (only for significant issues)
        if validation_result.has_issues:
            # Only log if there are serious issues (not just minor ones)
            serious_issues = [
                issue for issue in validation_result.issues
                if 'replacement character' in issue.lower() or 'mojibake' in issue.lower()
            ]
            if serious_issues and self.encoding_detector.logger:
                for issue in serious_issues:
                    self.encoding_detector.logger.warning(
                        f"Encoding issue in Excel cell: {issue}"
                    )

        # Normalize the text
        normalized_text = self.encoding_detector.normalize_text(text)

        return normalized_text

    def _get_cell_value(self, cell, cell_data):
        """Get the value from a cell, handling formulas, errors, dates, and hyperlinks.

        Args:
            cell: openpyxl Cell object (with formulas)
            cell_data: openpyxl Cell object (with calculated values)

        Returns:
            Cell value as string, with special handling for formulas, errors, dates, and hyperlinks
        """
        import datetime

        # Handle hyperlinks
        if cell.hyperlink:
            link_text = str(cell.value) if cell.value is not None else cell.hyperlink.target
            link_url = cell.hyperlink.target
            # Return Markdown link format
            return f"[{link_text}]({link_url})"

        # Handle formulas - use the calculated value from cell_data
        if cell.data_type == 'f':  # Formula
            # Get the calculated value from the data_only workbook
            calculated_value = cell_data.value
            if calculated_value is not None:
                return self._format_value(calculated_value)
            else:
                # No cached value available, try to evaluate simple formulas
                formula = cell.value
                if formula:
                    try:
                        evaluated = self._evaluate_simple_formula(formula, cell.parent)
                        if evaluated is not None:
                            return self._format_value(evaluated)
                    except:
                        pass
                # If evaluation fails, return the formula string
                return formula if formula is not None else ""

        # Handle errors (formula errors like #DIV/0!, #VALUE!, #REF!, etc.)
        if cell.data_type == 'e':  # Error
            return str(cell.value) if cell.value is not None else "#ERROR!"

        # Handle dates and times
        if isinstance(cell.value, datetime.datetime):
            return cell.value.strftime("%Y-%m-%d %H:%M:%S")
        elif isinstance(cell.value, datetime.date):
            return cell.value.strftime("%Y-%m-%d")
        elif isinstance(cell.value, datetime.time):
            return cell.value.strftime("%H:%M:%S")

        # Handle regular values
        return self._format_value(cell.value)

    def _format_value(self, value):
        """Format a cell value for display.

        Args:
            value: Cell value

        Returns:
            Formatted string representation
        """
        import datetime

        if value is None:
            return ""
        elif isinstance(value, bool):
            return "TRUE" if value else "FALSE"
        elif isinstance(value, (int, float)):
            # Format numbers, removing unnecessary decimal places
            if isinstance(value, float) and value.is_integer():
                return str(int(value))
            return str(value)
        elif isinstance(value, datetime.datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        elif isinstance(value, datetime.date):
            return value.strftime("%Y-%m-%d")
        elif isinstance(value, datetime.time):
            return value.strftime("%H:%M:%S")
        else:
            return str(value)

    def _evaluate_simple_formula(self, formula, sheet):
        """Evaluate simple Excel formulas when cached values are not available.

        This is a basic evaluator for common formulas like arithmetic operations
        and simple functions (SUM, AVERAGE, MAX, MIN). It's used as a fallback
        when openpyxl's data_only mode doesn't have cached values.

        Args:
            formula: Formula string (e.g., "=A2+B2", "=SUM(A2:A5)")
            sheet: openpyxl Worksheet object

        Returns:
            Evaluated result or None if evaluation fails
        """
        import re

        if not formula or not formula.startswith('='):
            return None

        # Remove the leading '='
        expr = formula[1:].strip()

        # Handle SUM function
        sum_match = re.match(r'SUM\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)', expr, re.IGNORECASE)
        if sum_match:
            col1, row1, col2, row2 = sum_match.groups()
            try:
                total = 0
                for row in range(int(row1), int(row2) + 1):
                    cell = sheet[f"{col1}{row}"]
                    if cell.value is not None and isinstance(cell.value, (int, float)):
                        total += cell.value
                return total
            except:
                return None

        # Handle AVERAGE function
        avg_match = re.match(r'AVERAGE\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)', expr, re.IGNORECASE)
        if avg_match:
            col1, row1, col2, row2 = avg_match.groups()
            try:
                values = []
                for row in range(int(row1), int(row2) + 1):
                    cell = sheet[f"{col1}{row}"]
                    if cell.value is not None and isinstance(cell.value, (int, float)):
                        values.append(cell.value)
                return sum(values) / len(values) if values else 0
            except:
                return None

        # Handle MAX function
        max_match = re.match(r'MAX\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)', expr, re.IGNORECASE)
        if max_match:
            col1, row1, col2, row2 = max_match.groups()
            try:
                values = []
                for row in range(int(row1), int(row2) + 1):
                    cell = sheet[f"{col1}{row}"]
                    if cell.value is not None and isinstance(cell.value, (int, float)):
                        values.append(cell.value)
                return max(values) if values else 0
            except:
                return None

        # Handle MIN function
        min_match = re.match(r'MIN\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)', expr, re.IGNORECASE)
        if min_match:
            col1, row1, col2, row2 = min_match.groups()
            try:
                values = []
                for row in range(int(row1), int(row2) + 1):
                    cell = sheet[f"{col1}{row}"]
                    if cell.value is not None and isinstance(cell.value, (int, float)):
                        values.append(cell.value)
                return min(values) if values else 0
            except:
                return None

        # Handle simple arithmetic operations (e.g., A2+B2, A2-B2, A2*B2, A2/B2)
        # Replace cell references with their values
        cell_refs = re.findall(r'([A-Z]+)(\d+)', expr)
        eval_expr = expr

        for col, row in cell_refs:
            try:
                cell = sheet[f"{col}{row}"]
                value = cell.value
                if value is None:
                    value = 0
                elif not isinstance(value, (int, float)):
                    # Can't evaluate non-numeric values
                    return None
                eval_expr = eval_expr.replace(f"{col}{row}", str(value))
            except:
                return None

        # Evaluate the expression safely
        try:
            # Only allow basic arithmetic operations
            if re.match(r'^[\d\s\+\-\*\/\(\)\.]+$', eval_expr):
                result = eval(eval_expr)
                return result
        except:
            pass

        return None



class PDFParser(DocumentParser):
    """Parser for PDF documents using PyPDF2 and pdfplumber."""

    def __init__(self, logger=None):
        """Initialize PDF parser with optional logger.

        Args:
            logger: Optional logger for encoding warnings
        """
        from src.encoding_detector import EncodingDetector
        from src.text_cleaner import TextCleaner
        self.encoding_detector = EncodingDetector(logger=logger)
        self.text_cleaner = TextCleaner()

    def parse(self, file_path: str) -> InternalDocument:
        """Parse a PDF document.

        Args:
            file_path: Path to the .pdf file

        Returns:
            InternalDocument representation
        """
        import PyPDF2
        import pdfplumber
        from pathlib import Path
        from src.internal_representation import Section, Heading, ImageReference
        from tqdm import tqdm

        # Extract metadata
        metadata = self._extract_metadata(file_path)

        # Extract images using PyMuPDF (most reliable method)
        images_with_data = self.extract_images_with_pymupdf(file_path)
        images = [img_ref for img_ref, _ in images_with_data]

        # Store image data for later extraction
        self._image_data = images_with_data

        # Extract content using both libraries
        sections = []

        try:
            # Use pdfplumber for better text extraction and table detection
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                # 進捗バーを表示
                with tqdm(total=total_pages, desc="PDFページ処理中", unit="ページ") as pbar:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        # Get images for this page
                        page_images = [img for img in images if img.page_number == page_num]

                        # Extract tables from page
                        tables = self._extract_tables(page)

                        # Extract text from page
                        text = page.extract_text()

                        content_blocks = []

                        if text and text.strip():
                            # Validate and normalize text encoding
                            text = self._process_text_encoding(text)
                            
                            # Clean text (remove orphan lines, fix line breaks)
                            text = self.text_cleaner.clean_text(text)

                            # Detect structure in text
                            content_blocks = self._detect_structure(text, page)

                            # Add tables to content blocks
                            content_blocks.extend(tables)

                        # Add image references to content blocks (even if no text)
                        for img in page_images:
                            content_blocks.append(img)

                        # Create section for this page if there's any content (text or images)
                        if content_blocks:
                            section = Section(
                                heading=Heading(level=2, text=f"Page {page_num}"),
                                content=content_blocks
                            )
                            sections.append(section)
                        
                        pbar.update(1)

        except Exception as e:
            # If pdfplumber fails, try PyPDF2 as fallback
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)

                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]

                        # Get images for this page
                        page_images = [img for img in images if img.page_number == page_num + 1]

                        text = page.extract_text()

                        content_blocks = []

                        if text and text.strip():
                            # Validate and normalize text encoding
                            text = self._process_text_encoding(text)
                            
                            # Clean text (remove orphan lines, fix line breaks)
                            text = self.text_cleaner.clean_text(text)

                            # Basic structure detection without pdfplumber
                            content_blocks = self._detect_structure_basic(text)

                        # Add image references to content blocks (even if no text)
                        for img in page_images:
                            content_blocks.append(img)

                        # Create section for this page if there's any content (text or images)
                        if content_blocks:
                            section = Section(
                                heading=Heading(level=2, text=f"Page {page_num + 1}"),
                                content=content_blocks
                            )
                            sections.append(section)

            except Exception as fallback_error:
                raise ValueError(f"Failed to parse PDF: {fallback_error}")

        return InternalDocument(
            metadata=metadata,
            sections=sections,
            images=images
        )

    def _process_text_encoding(self, text: str) -> str:
        """Process and validate text encoding.

        Args:
            text: Extracted text to process

        Returns:
            Normalized text with encoding issues resolved
        """
        # Validate text encoding and detect issues
        validation_result = self.encoding_detector.validate_text_encoding(text)

        # Log any encoding issues found
        if validation_result.has_issues:
            for issue in validation_result.issues:
                if self.encoding_detector.logger:
                    self.encoding_detector.logger.warning(
                        f"Encoding issue in PDF text: {issue}"
                    )

        # Normalize the text
        normalized_text = self.encoding_detector.normalize_text(text)

        return normalized_text

    def _extract_metadata(self, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            DocumentMetadata object
        """
        import PyPDF2
        from pathlib import Path

        metadata_dict = {}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract PDF metadata
                if pdf_reader.metadata:
                    metadata_dict = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    }

                # Add page count
                metadata_dict['pages'] = len(pdf_reader.pages)

        except Exception:
            # If metadata extraction fails, use defaults
            pass

        from .internal_representation import DocumentMetadata
        return DocumentMetadata(
            title=metadata_dict.get('title', Path(file_path).stem) or Path(file_path).stem,
            author=metadata_dict.get('author', ''),
            created_date=None,
            modified_date=None,
            source_format='pdf'
        )

    def _detect_structure(self, text: str, page) -> list:
        """Detect structure in PDF text (headings, paragraphs).

        Uses heuristics to identify headings based on:
        - Text in all caps
        - Short lines (< 80 chars)
        - Lines ending without punctuation

        Args:
            text: Extracted text from page
            page: pdfplumber page object (for font analysis if available)

        Returns:
            List of content blocks (Heading or Paragraph objects)
        """
        from src.internal_representation import Heading, Paragraph, TextFormatting

        content_blocks = []
        lines = text.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Heuristic: Detect potential headings
            is_heading = False
            heading_level = 3  # Default to H3 for detected headings

            # Check if line is all caps (likely a heading)
            if line.isupper() and len(line) < 80:
                is_heading = True
                heading_level = 2
            # Check if line is short and doesn't end with punctuation
            elif len(line) < 80 and not line[-1] in '.!?,;:':
                # Check if next line is empty or starts a new paragraph
                if i + 1 < len(lines) and (not lines[i + 1].strip() or lines[i + 1].strip()[0].isupper()):
                    is_heading = True
                    heading_level = 3

            if is_heading:
                content_blocks.append(Heading(level=heading_level, text=line))
            else:
                # Accumulate paragraph text
                para_lines = [line]
                i += 1

                # Continue accumulating lines until we hit an empty line or potential heading
                while i < len(lines):
                    next_line = lines[i].strip()

                    if not next_line:
                        break

                    # Check if next line might be a heading
                    if (next_line.isupper() and len(next_line) < 80) or \
                       (len(next_line) < 80 and not next_line[-1] in '.!?,;:'):
                        break

                    para_lines.append(next_line)
                    i += 1

                # Create paragraph
                para_text = ' '.join(para_lines)
                if para_text:
                    content_blocks.append(Paragraph(
                        text=para_text,
                        formatting=TextFormatting.NORMAL
                    ))
                continue

            i += 1

        return content_blocks

    def _detect_structure_basic(self, text: str) -> list:
        """Basic structure detection without pdfplumber page object.

        Args:
            text: Extracted text

        Returns:
            List of content blocks
        """
        from src.internal_representation import Heading, Paragraph, TextFormatting

        content_blocks = []

        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Check if it might be a heading (short, all caps, or no ending punctuation)
                if len(para_text) < 80 and (para_text.isupper() or not para_text[-1] in '.!?,;:'):
                    content_blocks.append(Heading(level=3, text=para_text))
                else:
                    # Create paragraph
                    paragraph = Paragraph(
                        text=para_text,
                        formatting=TextFormatting.NORMAL
                    )
                    content_blocks.append(paragraph)

        return content_blocks

    def _extract_tables(self, page) -> list:
        """Extract tables from a PDF page using pdfplumber.

        Args:
            page: pdfplumber page object

        Returns:
            List of Table objects
        """
        from src.internal_representation import Table

        tables = []

        try:
            # Extract tables from page
            page_tables = page.extract_tables()

            if page_tables:
                for table_data in page_tables:
                    if table_data and len(table_data) > 0:
                        # First row is typically headers
                        headers = []
                        rows = []

                        # Check if first row looks like headers
                        if len(table_data) > 1:
                            headers = [str(cell) if cell else '' for cell in table_data[0]]
                            rows = [[str(cell) if cell else '' for cell in row] for row in table_data[1:]]
                        else:
                            # No headers, treat all as data
                            rows = [[str(cell) if cell else '' for cell in row] for row in table_data]

                        # Create Table object
                        table = Table(
                            headers=headers,
                            rows=rows
                        )
                        tables.append(table)

        except Exception:
            # If table extraction fails, continue without tables
            pass

        return tables
    def _extract_images_from_page(self, page, page_num: int) -> list:
        """Extract images from a PDF page using pdfplumber.

        Args:
            page: pdfplumber page object
            page_num: Page number (1-indexed)

        Returns:
            List of tuples (ImageReference, image_bytes)
        """
        from src.internal_representation import ImageReference

        images = []

        try:
            # pdfplumber doesn't provide easy image extraction
            # We'll use PyMuPDF for actual image extraction
            # This method just creates placeholders
            if hasattr(page, 'images') and page.images:
                for img_idx, img_info in enumerate(page.images, start=1):
                    try:
                        # Create image reference
                        img_ref = ImageReference(
                            source_path=f"page_{page_num}_image_{img_idx}",
                            alt_text=f"Image {img_idx} from page {page_num}",
                            page_number=page_num
                        )
                        images.append((img_ref, None))  # No bytes yet
                    except Exception:
                        pass

        except Exception:
            pass

        return images

    def _extract_images_pypdf2(self, page, page_num: int) -> list:
        """Extract images from a PDF page using PyPDF2.

        Args:
            page: PyPDF2 page object
            page_num: Page number (1-indexed)

        Returns:
            List of tuples (ImageReference, image_bytes)
        """
        from src.internal_representation import ImageReference

        images = []

        try:
            # Try to extract images from page resources
            if '/XObject' in page['/Resources']:
                xobjects = page['/Resources']['/XObject'].get_object()

                for obj_idx, obj_name in enumerate(xobjects, start=1):
                    obj = xobjects[obj_name]

                    if obj['/Subtype'] == '/Image':
                        try:
                            # Create image reference
                            img_ref = ImageReference(
                                source_path=f"page_{page_num}_image_{obj_idx}",
                                alt_text=f"Image {obj_idx} from page {page_num}",
                                page_number=page_num
                            )
                            images.append((img_ref, None))  # No bytes yet
                        except Exception:
                            pass

        except Exception:
            pass

        return images
    
    def extract_images_with_pymupdf(self, file_path: str) -> list:
        """Extract actual image data from PDF using PyMuPDF.
        
        This method extracts both embedded bitmap images and vector graphics.
        Vector graphics (diagrams, schematics) are rendered as PNG images.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of tuples (ImageReference, image_bytes)
        """
        import fitz  # PyMuPDF
        from src.internal_representation import ImageReference
        from tqdm import tqdm
        
        images_with_data = []
        
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            with tqdm(total=total_pages, desc="画像抽出中", unit="ページ") as pbar:
                for page_num in range(total_pages):
                    page = doc[page_num]
                    
                    # Extract embedded bitmap images
                    image_list = page.get_images(full=True)
                    
                    for img_idx, img_info in enumerate(image_list, start=1):
                        try:
                            xref = img_info[0]
                            
                            # Extract image
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]
                            
                            # Create image reference
                            img_ref = ImageReference(
                                source_path=f"page_{page_num + 1}_image_{img_idx}.{image_ext}",
                                alt_text=f"Image {img_idx} from page {page_num + 1}",
                                page_number=page_num + 1,
                                mime_type=f"image/{image_ext}"
                            )
                            
                            images_with_data.append((img_ref, image_bytes))
                            
                        except Exception as e:
                            print(f"Warning: Failed to extract image {img_idx} from page {page_num + 1}: {e}")
                            continue
                    
                    # Extract vector graphics (drawings)
                    # Check if page has vector graphics by analyzing drawing commands
                    drawings = page.get_drawings()
                    
                    if drawings and len(drawings) > 10:  # Threshold to identify pages with significant vector content
                        # Check if this looks like a diagram/schematic (has many drawing objects)
                        # and doesn't have much text (to avoid extracting text-heavy pages)
                        text_blocks = page.get_text("blocks")
                        text_ratio = sum(len(block[4]) for block in text_blocks) / max(len(drawings), 1)
                        
                        # If there are many drawings and relatively little text, extract as vector graphic
                        if text_ratio < 5:  # Less than 5 characters per drawing object on average
                            try:
                                # Render the page as a high-resolution PNG
                                # Use 144 DPI for good quality without excessive file size
                                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom = 144 DPI
                                pix = page.get_pixmap(matrix=mat, alpha=False)
                                image_bytes = pix.tobytes("png")
                                
                                # Create image reference for vector graphic
                                vector_img_ref = ImageReference(
                                    source_path=f"page_{page_num + 1}_vector_graphic.png",
                                    alt_text=f"Vector graphic from page {page_num + 1}",
                                    page_number=page_num + 1,
                                    mime_type="image/png"
                                )
                                
                                images_with_data.append((vector_img_ref, image_bytes))
                                
                            except Exception as e:
                                print(f"Warning: Failed to extract vector graphic from page {page_num + 1}: {e}")
                    
                    pbar.update(1)
            
            doc.close()
            
        except Exception as e:
            print(f"Warning: PyMuPDF image extraction failed: {e}")
        
        return images_with_data

